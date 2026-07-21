import os
from dotenv import load_dotenv
import litellm
from django.conf import settings
from collections import defaultdict
from datetime import datetime, timedelta, date
import itertools
import re
from difflib import SequenceMatcher

import html as html_module
import markdown

from django.http import JsonResponse, HttpResponse, FileResponse
from django.urls import reverse
from fpdf import FPDF
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods
import json
from django.contrib.auth.decorators import login_required, user_passes_test
from .models import Course, User, Department, Section, DynamicForm, FormQuestion, DynamicFormSubmission, FormAnswer, CourseFaculty, CourseOutline
from .notifications import (
    notify_faculty_form_revision,
    notify_faculty_outline_revision,
)
from django.db.models import Count, Q, F
from datetime import datetime, timedelta

from litellm import completion

        
# Load environment variables
load_dotenv()


def is_admin(user):
    return user.is_authenticated and user.role == User.ROLE_ADMIN

def is_admin_or_crc(user):
    return user.is_authenticated and user.role in [User.ROLE_ADMIN, User.ROLE_CRC_MEMBER]


def course_catalogue_file_payload(course):
    """Return catalogue file metadata for a course."""
    if not course.catalogue_file:
        return {
            'has_catalogue_file': False,
            'catalogue_file_name': None,
            'catalogue_download_url': None,
        }
    return {
        'has_catalogue_file': True,
        'catalogue_file_name': os.path.basename(course.catalogue_file.name),
        'catalogue_download_url': reverse('api_course_catalogue_download', args=[course.id]),
    }


def course_faculty_section_payload(assignment):
    """Serialize assigned sections for a CourseFaculty row."""
    sections = list(assignment.sections.all())
    return {
        'section_ids': [section.id for section in sections],
        'sections': [
            {
                'id': section.id,
                'code': section.code,
                'name': section.name,
            }
            for section in sections
        ],
        'section': assignment.section_display(),
    }


def normalize_section_ids(raw_section_ids):
    """Normalize section id values from API payloads into a list of ints."""
    if raw_section_ids is None:
        return []
    if isinstance(raw_section_ids, str):
        raw_section_ids = [
            part.strip()
            for part in raw_section_ids.split(',')
            if part.strip()
        ]
    if not isinstance(raw_section_ids, (list, tuple)):
        raw_section_ids = [raw_section_ids]

    section_ids = []
    for raw_id in raw_section_ids:
        try:
            section_ids.append(int(raw_id))
        except (TypeError, ValueError):
            continue
    return section_ids


def set_course_faculty_sections(assignment, raw_section_ids):
    """Attach Section rows to a CourseFaculty assignment."""
    section_ids = normalize_section_ids(raw_section_ids)
    if not section_ids:
        assignment.sections.clear()
        return
    sections = Section.objects.filter(id__in=section_ids)
    assignment.sections.set(sections)


def parse_course_request_payload(request):
    """Parse course create/update payload from JSON or multipart form data."""
    content_type = (request.content_type or '').lower()
    if 'multipart/form-data' in content_type:
        data = request.POST.dict()
        if 'credits' in data and data['credits'] != '':
            try:
                data['credits'] = int(data['credits'])
            except (TypeError, ValueError):
                data['credits'] = 3
        if 'department_id' in data and data['department_id'] != '':
            try:
                data['department_id'] = int(data['department_id'])
            except (TypeError, ValueError):
                pass
        return data, request.FILES.get('catalogue_file')
    return json.loads(request.body or '{}'), None


def course_outline_to_dict(outline):
    """Serialize a CourseOutline for JSON APIs (faculty / CRC)."""
    return {
        'id': outline.id,
        'course': {
            'id': outline.course.id,
            'code': outline.course.code,
            'title': outline.course.title,
        },
        'version': outline.version,
        'title': outline.title,
        'description': outline.description,
        'content': outline.content,
        'status': outline.status,
        'notes': outline.notes,
        'is_current': outline.is_current,
        'faculty_author_id': outline.faculty_id,
        'faculty_author_username': outline.faculty.username,
        'created_at': outline.created_at.isoformat() if outline.created_at else None,
        'submitted_at': outline.submitted_at.isoformat() if outline.submitted_at else None,
        'approved_at': outline.approved_at.isoformat() if outline.approved_at else None,
    }


# Department API Views (list: admin + CRC for course management UI)
@login_required
@user_passes_test(is_admin_or_crc)
@require_http_methods(["GET"])
def api_departments(request):
    departments = list(Department.objects.values('id', 'name', 'code', 'description'))
    return JsonResponse(departments, safe=False)

@login_required
@user_passes_test(is_admin)
@csrf_exempt
@require_http_methods(["POST"])
def api_departments_create(request):
    try:
        data = json.loads(request.body)
        department = Department.objects.create(
            name=data.get('name'),
            code=data.get('code'),
            description=data.get('description', '')
        )
        return JsonResponse({
            'id': department.id,
            'name': department.name,
            'code': department.code,
            'description': department.description
        }, status=201)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)


# Section API Views
@login_required
@user_passes_test(is_admin_or_crc)
@require_http_methods(["GET"])
def api_sections(request):
    sections = list(Section.objects.values('id', 'name', 'code', 'description', 'created_at'))
    for section in sections:
        if section.get('created_at'):
            section['created_at'] = section['created_at'].isoformat()
    return JsonResponse(sections, safe=False)


@login_required
@user_passes_test(is_admin)
@csrf_exempt
@require_http_methods(["POST"])
def api_sections_create(request):
    try:
        data = json.loads(request.body)
        code = (data.get('code') or '').strip()
        name = (data.get('name') or '').strip()
        if not name or not code:
            return JsonResponse({'error': 'Section name and code are required'}, status=400)
        if Section.objects.filter(code__iexact=code).exists():
            return JsonResponse({'error': 'Section code already exists'}, status=400)
        section = Section.objects.create(
            name=name,
            code=code,
            description=data.get('description', '') or '',
        )
        return JsonResponse({
            'id': section.id,
            'name': section.name,
            'code': section.code,
            'description': section.description,
        }, status=201)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)


@login_required
@user_passes_test(is_admin)
@csrf_exempt
@require_http_methods(["GET", "PUT", "DELETE"])
def api_section_detail(request, section_id):
    try:
        section = Section.objects.get(id=section_id)

        if request.method == 'GET':
            return JsonResponse({
                'section': {
                    'id': section.id,
                    'name': section.name,
                    'code': section.code,
                    'description': section.description,
                }
            })

        if request.method == 'PUT':
            data = json.loads(request.body)
            name = (data.get('name') or section.name or '').strip()
            code = (data.get('code') or section.code or '').strip()
            if not name or not code:
                return JsonResponse({'error': 'Section name and code are required'}, status=400)
            if Section.objects.filter(code__iexact=code).exclude(id=section.id).exists():
                return JsonResponse({'error': 'Section code already exists'}, status=400)
            section.name = name
            section.code = code
            if 'description' in data:
                section.description = data.get('description') or ''
            section.save()
            return JsonResponse({
                'id': section.id,
                'name': section.name,
                'code': section.code,
                'description': section.description,
            })

        assignment_count = section.faculty_assignments.count()
        section.delete()
        return JsonResponse({
            'message': 'Section deleted successfully',
            'cleared_assignments': assignment_count,
        })
    except Section.DoesNotExist:
        return JsonResponse({'error': 'Section not found'}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)


@login_required
@user_passes_test(is_admin)
@csrf_exempt
@require_http_methods(["GET", "PUT", "DELETE"])
def api_department_detail(request, department_id):
    """Get, update, or delete department details"""
    try:
        department = Department.objects.get(id=department_id)
        
        if request.method == 'GET':
            courses = Course.objects.filter(department=department).values(
                'id', 'title', 'code', 'description', 'credits'
            )
            
            courses_list = []
            for course in courses:
                course_data = dict(course)
                course_faculty = CourseFaculty.objects.filter(course_id=course['id']).select_related('faculty')
                faculty_data = []
                for cf in course_faculty:
                    faculty_payload = {
                        'id': cf.faculty.id,
                        'username': cf.faculty.username,
                        'email': cf.faculty.email,
                        'is_coordinator': cf.is_coordinator,
                    }
                    faculty_payload.update(course_faculty_section_payload(cf))
                    faculty_data.append(faculty_payload)
                course_data['faculty'] = faculty_data
                courses_list.append(course_data)
            
            return JsonResponse({
                'department': {
                    'id': department.id,
                    'name': department.name,
                    'code': department.code,
                    'description': department.description
                },
                'courses': courses_list,
                'total_courses': len(courses_list),
                'total_faculty': CourseFaculty.objects.filter(course__department=department).values('faculty').distinct().count()
            })
        
        elif request.method == 'PUT':
            # Update department
            data = json.loads(request.body)
            
            # Update fields
            if 'name' in data:
                department.name = data['name']
            if 'code' in data:
                department.code = data['code']
            if 'description' in data:
                department.description = data['description']
            
            department.save()
            
            return JsonResponse({
                'id': department.id,
                'name': department.name,
                'code': department.code,
                'description': department.description
            })

        elif request.method == 'DELETE':
            course_count = Course.objects.filter(department=department).count()
            department_name = department.name
            department.delete()
            return JsonResponse({
                'message': f'Department "{department_name}" deleted successfully',
                'deleted_courses': course_count,
            })
            
    except Department.DoesNotExist:
        return JsonResponse({'error': 'Department not found'}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)


# Course API Views (admin + CRC)
@login_required
@user_passes_test(is_admin_or_crc)
@require_http_methods(["GET"])
@login_required
@user_passes_test(is_admin_or_crc)
def api_courses(request):
    courses = list(Course.objects.select_related('department').all())
    courses_data = []
    for course in courses:
        course_data = {
            'id': course.id,
            'title': course.title,
            'code': course.code,
            'description': course.description,
            'department_id': course.department_id,
            'credits': course.credits,
            'department_code': course.department.code if course.department else '',
            'department_name': course.department.name if course.department else '',
        }
        course_data.update(course_catalogue_file_payload(course))
        coordinator = CourseFaculty.objects.filter(
            course_id=course.id,
            is_coordinator=True
        ).select_related('faculty').prefetch_related('sections').first()

        if coordinator:
            course_data['course_coordinator_id'] = coordinator.faculty.id
            course_data['course_coordinator_name'] = coordinator.faculty.username
            course_data['coordinator_section'] = coordinator.section_display()
        else:
            course_data['course_coordinator_id'] = None
            course_data['course_coordinator_name'] = None
            course_data['coordinator_section'] = None
        courses_data.append(course_data)

    return JsonResponse(courses_data, safe=False)

@csrf_exempt
@login_required
@user_passes_test(is_admin_or_crc)
@require_http_methods(["POST"])
def api_courses_create(request):
    try:
        data, catalogue_file = parse_course_request_payload(request)
        department = Department.objects.get(id=data.get('department_id'))
        course = Course(
            title=data.get('title'),
            code=data.get('code'),
            description=data.get('description', ''),
            department=department,
            credits=data.get('credits', 3) or 3,
        )
        if catalogue_file:
            course.catalogue_file = catalogue_file
        course.save()

        faculty_assignments = data.get('faculty_assignments', [])
        if isinstance(faculty_assignments, str):
            try:
                faculty_assignments = json.loads(faculty_assignments)
            except json.JSONDecodeError:
                faculty_assignments = []

        for assignment in faculty_assignments:
            faculty_id = assignment.get('faculty_id')
            is_coordinator = assignment.get('is_coordinator', False)
            section_ids = assignment.get('section_ids', assignment.get('section', []))

            if faculty_id:
                faculty = User.objects.get(id=faculty_id, role=User.ROLE_FACULTY)
                course_faculty = CourseFaculty.objects.create(
                    course=course,
                    faculty=faculty,
                    is_coordinator=is_coordinator,
                )
                set_course_faculty_sections(course_faculty, section_ids)

        response_data = {
            'id': course.id,
            'title': course.title,
            'code': course.code,
            'description': course.description,
            'department_id': course.department.id,
            'credits': course.credits,
        }
        response_data.update(course_catalogue_file_payload(course))
        return JsonResponse(response_data, status=201)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)

@csrf_exempt
@login_required
@user_passes_test(is_admin_or_crc)
@require_http_methods(["PUT", "POST"])
def api_course_update(request, course_id):
    try:
        data, catalogue_file = parse_course_request_payload(request)
        course = Course.objects.get(id=course_id)
        course.title = data.get('title', course.title)
        course.code = data.get('code', course.code)
        course.description = data.get('description', course.description)
        if 'credits' in data and data.get('credits') not in (None, ''):
            course.credits = data.get('credits')

        if 'department_id' in data and data.get('department_id') not in (None, ''):
            department = Department.objects.get(id=data['department_id'])
            course.department = department

        if catalogue_file:
            if course.catalogue_file:
                course.catalogue_file.delete(save=False)
            course.catalogue_file = catalogue_file

        course.save()

        response_data = {
            'id': course.id,
            'title': course.title,
            'code': course.code,
            'description': course.description,
            'department_id': course.department.id,
            'credits': course.credits,
        }
        response_data.update(course_catalogue_file_payload(course))
        return JsonResponse(response_data)
    except Course.DoesNotExist:
        return JsonResponse({'error': 'Course not found'}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)


@login_required
@require_http_methods(["GET"])
def api_course_catalogue_download(request, course_id):
    """Download the uploaded course catalogue file for a course."""
    try:
        course = Course.objects.get(id=course_id)
        if not course.catalogue_file:
            return JsonResponse({'error': 'No catalogue file uploaded for this course'}, status=404)
        return FileResponse(
            course.catalogue_file.open('rb'),
            as_attachment=True,
            filename=os.path.basename(course.catalogue_file.name),
        )
    except Course.DoesNotExist:
        return JsonResponse({'error': 'Course not found'}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)


@login_required
@require_http_methods(["GET"])
def api_course_catalogues_list(request):
    """List all courses with catalogue download info (faculty/CRC/admin)."""
    courses = Course.objects.select_related('department').order_by('code')
    courses_data = []
    for course in courses:
        course_data = {
            'id': course.id,
            'title': course.title,
            'code': course.code,
            'department_name': course.department.name if course.department else '',
            'department_code': course.department.code if course.department else '',
        }
        course_data.update(course_catalogue_file_payload(course))
        courses_data.append(course_data)
    return JsonResponse(courses_data, safe=False)

# Assign/Update Course Faculty
@csrf_exempt
@login_required
@user_passes_test(is_admin_or_crc)
@require_http_methods(["POST"])
def api_assign_course_faculty(request, course_id):
    try:
        data = json.loads(request.body)
        course = Course.objects.get(id=course_id)
        
        # Create new assignments
        faculty_assignments = data.get('faculty_assignments', [])

        coordinator_count = sum(
            1 for assignment in faculty_assignments
            if assignment.get('faculty_id') and assignment.get('is_coordinator', False)
        )
        if coordinator_count > 1:
            return JsonResponse({
                'error': 'A course can have only one coordinator',
                'success': False,
            }, status=400)
        
        # Clear existing assignments
        CourseFaculty.objects.filter(course=course).delete()
        
        assignments_created = 0
        for assignment in faculty_assignments:
            faculty_id = assignment.get('faculty_id')
            is_coordinator = assignment.get('is_coordinator', False)
            section_ids = assignment.get('section_ids', assignment.get('section', []))
            
            if faculty_id:
                try:
                    faculty = User.objects.get(id=faculty_id, role=User.ROLE_FACULTY)
                    course_faculty = CourseFaculty.objects.create(
                        course=course,
                        faculty=faculty,
                        is_coordinator=is_coordinator,
                    )
                    set_course_faculty_sections(course_faculty, section_ids)
                    assignments_created += 1
                except User.DoesNotExist:
                    continue
        
        return JsonResponse({
            'message': f'Successfully assigned {assignments_created} faculty to {course.title}',
            'total_assigned': assignments_created,
            'success': True
        })
    except Course.DoesNotExist:
        return JsonResponse({'error': 'Course not found', 'success': False}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e), 'success': False}, status=400)

@csrf_exempt
@login_required
@user_passes_test(is_admin_or_crc)
@require_http_methods(["DELETE"])
def api_course_delete(request, course_id):
    try:
        course = Course.objects.get(id=course_id)
        course.delete()
        return JsonResponse({'message': 'Course deleted successfully'})
    except Course.DoesNotExist:
        return JsonResponse({'error': 'Course not found'}, status=404)

# Faculty Assignment API
@login_required
@user_passes_test(is_admin)
@csrf_exempt
@require_http_methods(["POST"])
def api_assign_courses_to_faculty(request, user_id):
    try:
        data = json.loads(request.body)
        faculty_member = User.objects.get(id=user_id, role=User.ROLE_FACULTY)
        course_ids = data.get('course_ids', [])
        coordinator_info = data.get('coordinator_info', {}) or {}
        section_info = data.get('section_info', {}) or {}
        section_ids_info = data.get('section_ids_info', {}) or {}

        def mapping_get(mapping, course_id, default=None):
            if not isinstance(mapping, dict):
                return default
            for key in (course_id, str(course_id)):
                if key in mapping:
                    return mapping[key]
            try:
                as_int = int(course_id)
            except (TypeError, ValueError):
                return default
            if as_int in mapping:
                return mapping[as_int]
            return default

        # Validate coordinator conflicts before changing assignments
        for course_id in course_ids:
            if not bool(mapping_get(coordinator_info, course_id, False)):
                continue
            try:
                course = Course.objects.get(id=course_id)
            except Course.DoesNotExist:
                continue
            existing_coordinator = CourseFaculty.objects.filter(
                course=course,
                is_coordinator=True,
            ).exclude(faculty=faculty_member).select_related('faculty').first()
            if existing_coordinator:
                return JsonResponse({
                    'error': (
                        f'Course {course.code} already has coordinator '
                        f'{existing_coordinator.faculty.username}'
                    ),
                    'success': False,
                }, status=400)

        # Replace this faculty's course assignments with the submitted selection
        CourseFaculty.objects.filter(faculty=faculty_member).delete()

        assignments_created = 0
        for course_id in course_ids:
            try:
                course = Course.objects.get(id=course_id)
            except Course.DoesNotExist:
                continue
            is_coordinator = bool(mapping_get(coordinator_info, course_id, False))
            course_faculty = CourseFaculty.objects.create(
                course=course,
                faculty=faculty_member,
                is_coordinator=is_coordinator,
            )
            section_ids = mapping_get(section_ids_info, course_id, None)
            if section_ids is None:
                section_ids = mapping_get(section_info, course_id, [])
            set_course_faculty_sections(course_faculty, section_ids)
            assignments_created += 1

        return JsonResponse({
            'message': f'Successfully assigned {assignments_created} courses to {faculty_member.username}',
            'total_assigned': assignments_created,
            'success': True,
        })
    except User.DoesNotExist:
        return JsonResponse({'error': 'Faculty member not found'}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)

# Faculty Dashboard API
@login_required
def api_faculty_courses(request):
    if request.user.role != User.ROLE_FACULTY:
        return JsonResponse({'error': 'Access denied'}, status=403)
    
    course_assignments = CourseFaculty.objects.filter(faculty=request.user).select_related(
        'course', 'course__department'
    ).prefetch_related('sections')
    
    courses_data = []
    for assignment in course_assignments:
        course_data = {
            'id': assignment.course.id,
            'title': assignment.course.title,
            'code': assignment.course.code,
            'description': assignment.course.description,
            'credits': assignment.course.credits,
            'department_name': assignment.course.department.name if assignment.course.department else '',
            'department_code': assignment.course.department.code if assignment.course.department else '',
            'is_coordinator': assignment.is_coordinator,
        }
        course_data.update(course_faculty_section_payload(assignment))
        courses_data.append(course_data)
    
    return JsonResponse(courses_data, safe=False)

# Form API Views - Updated for Admin and CRC Member
@login_required
@user_passes_test(is_admin_or_crc)
@require_http_methods(["GET"])
def api_forms(request):
    forms = list(DynamicForm.objects.values('id', 'name', 'description', 'form_type', 'status'))
    return JsonResponse(forms, safe=False)



# api_views.py - Update api_publish_form
@login_required
@user_passes_test(is_admin_or_crc)
@csrf_exempt
@require_http_methods(["POST"])
def api_publish_form(request, form_id):
    """Publish a form (set status to active) - Allow multiple forms to be active"""
    try:
        form = DynamicForm.objects.get(id=form_id)
        
        # Check if form type is valid (only CCR or CRR)
        if form.form_type not in ['ccr', 'crr']:
            return JsonResponse({'error': 'Cannot publish non-CCR/CRR forms'}, status=400)
        
        # REMOVED: No longer deactivating other forms of the same type
        # Set this form to active
        form.status = 'active'
        form.save()
        
        return JsonResponse({
            'message': f'{form.get_form_type_display()} published successfully',
            'form_id': form.id,
            'form_type': form.form_type,
            'status': form.status
        })
    except DynamicForm.DoesNotExist:
        return JsonResponse({'error': 'Form not found'}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)




@login_required
@user_passes_test(is_admin_or_crc)
@csrf_exempt
@require_http_methods(["POST"])
def api_unpublish_form(request, form_id):
    """Unpublish a form (set status to inactive) - Only for CCR/CRR forms"""
    try:
        form = DynamicForm.objects.get(id=form_id)
        
        # Check if form type is valid (only CCR or CRR)
        if form.form_type not in ['ccr', 'crr']:
            return JsonResponse({'error': 'Cannot unpublish non-CCR/CRR forms'}, status=400)
        
        form.status = 'inactive'
        form.save()
        
        return JsonResponse({
            'message': f'{form.get_form_type_display()} unpublished',
            'form_id': form.id,
            'form_type': form.form_type,
            'status': form.status
        })
    except DynamicForm.DoesNotExist:
        return JsonResponse({'error': 'Form not found'}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)

@login_required
@user_passes_test(is_admin_or_crc)
@csrf_exempt
@require_http_methods(["PUT"])
def api_form_update(request, form_id):
    try:
        data = json.loads(request.body)
        form = DynamicForm.objects.get(id=form_id)
        form.name = data.get('name', form.name)
        form.description = data.get('description', form.description)
        form.form_type = data.get('form_type', form.form_type)
        form.status = data.get('status', form.status)
        form.save()
        return JsonResponse({
            'id': form.id,
            'name': form.name,
            'description': form.description,
            'form_type': form.form_type,
            'status': form.status
        })
    except DynamicForm.DoesNotExist:
        return JsonResponse({'error': 'Form not found'}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)

@login_required
@user_passes_test(is_admin_or_crc)
@csrf_exempt
@require_http_methods(["DELETE"])
def api_form_delete(request, form_id):
    try:
        form = DynamicForm.objects.get(id=form_id)
        form.delete()
        return JsonResponse({'message': 'Form deleted successfully'})
    except DynamicForm.DoesNotExist:
        return JsonResponse({'error': 'Form not found'}, status=404)

# Dynamic Form API Views
@login_required
@user_passes_test(is_admin_or_crc)
@require_http_methods(["GET"])
def api_dynamic_forms(request):
    """Get all dynamic forms (ONLY UNIVERSAL CCR/CRR)"""
    forms = list(DynamicForm.objects.filter(
        form_type__in=['ccr', 'crr']  # Only universal forms
    ).values('id', 'name', 'description', 'form_type', 'status', 'created_at'))
    return JsonResponse(forms, safe=False)




@login_required
@user_passes_test(is_admin_or_crc)
@csrf_exempt
@require_http_methods(["POST"])
def api_dynamic_forms_create(request):
    try:
        data = json.loads(request.body)
        form = DynamicForm.objects.create(
            name=data.get('name', 'Dynamic Form'),
            description=data.get('description', ''),
            form_type=data.get('form_type', 'crr'),
            status=data.get('status', DynamicForm.STATUS_INACTIVE)
        )
        return JsonResponse({
            'id': form.id,
            'name': form.name,
            'description': form.description,
            'form_type': form.form_type,
            'status': form.status
        }, status=201)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)

@login_required
@user_passes_test(is_admin_or_crc)
@csrf_exempt
@require_http_methods(["PUT"])
def api_dynamic_form_update(request, form_id):
    try:
        data = json.loads(request.body)
        form = DynamicForm.objects.get(id=form_id)
        form.name = data.get('name', form.name)
        form.description = data.get('description', form.description)
        form.form_type = data.get('form_type', form.form_type)
        form.status = data.get('status', form.status)
        form.save()
        return JsonResponse({
            'id': form.id,
            'name': form.name,
            'description': form.description,
            'form_type': form.form_type,
            'status': form.status
        })
    except DynamicForm.DoesNotExist:
        return JsonResponse({'error': 'Dynamic Form not found'}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)

@login_required
@user_passes_test(is_admin_or_crc)
@csrf_exempt
@require_http_methods(["DELETE"])
def api_dynamic_form_delete(request, form_id):
    try:
        form = DynamicForm.objects.get(id=form_id)
        form.delete()
        return JsonResponse({'message': 'Dynamic Form deleted successfully'})
    except DynamicForm.DoesNotExist:
        return JsonResponse({'error': 'Dynamic Form not found'}, status=404)

# Form Questions API
@login_required
@user_passes_test(is_admin_or_crc)
@require_http_methods(["GET"])
def api_form_questions(request, form_id):
    questions = list(FormQuestion.objects.filter(form_id=form_id).values(
        'id', 'question_text', 'question_type', 'order', 'required', 'options', 'config', 'help_text'
    ))
    return JsonResponse(questions, safe=False)

@login_required
@user_passes_test(is_admin_or_crc)
@csrf_exempt
@require_http_methods(["POST"])
def api_form_questions_create(request, form_id):
    try:
        data = json.loads(request.body)
        form = DynamicForm.objects.get(id=form_id)
        
        config = data.get('config')
        if config and isinstance(config, str):
            try:
                config = json.loads(config)
            except:
                config = None
        
        # For CLO percentage questions, ensure proper config
        if data.get('question_type') == 'clo_percentage':
            if not config:
                config = {
                    'clo_fields': ['clo1', 'clo2', 'clo3', 'clo4'],
                    'min_value': 0,
                    'max_value': 100,
                    'step': 0.1,
                    'suffix': '%'
                }
        
        question = FormQuestion.objects.create(
            form=form,
            question_text=data.get('question_text'),
            question_type=data.get('question_type', 'text'),
            order=data.get('order', 0),
            required=data.get('required', True),
            options=data.get('options', ''),
            config=config,
            help_text=data.get('help_text', '')
        )
        return JsonResponse({
            'id': question.id,
            'question_text': question.question_text,
            'question_type': question.question_type,
            'order': question.order,
            'required': question.required,
            'options': question.options,
            'config': question.config,
            'help_text': question.help_text
        }, status=201)
    except Exception as e:
        print(f"Error creating question: {str(e)}")  # Debug logging
        return JsonResponse({'error': str(e)}, status=400)


@login_required
@user_passes_test(is_admin_or_crc)
@csrf_exempt
@require_http_methods(["PUT"])
def api_form_question_update(request, question_id):
    try:
        data = json.loads(request.body)
        question = FormQuestion.objects.get(id=question_id)
        question.question_text = data.get('question_text', question.question_text)
        question.question_type = data.get('question_type', question.question_type)
        question.order = data.get('order', question.order)
        question.required = data.get('required', question.required)
        question.options = data.get('options', question.options)
        question.help_text = data.get('help_text', question.help_text)
        
        config = data.get('config')
        if config is not None:
            if isinstance(config, str):
                try:
                    config = json.loads(config)
                except:
                    config = None
            question.config = config
        
        question.save()
        return JsonResponse({
            'id': question.id,
            'question_text': question.question_text,
            'question_type': question.question_type,
            'order': question.order,
            'required': question.required,
            'options': question.options,
            'config': question.config,
            'help_text': question.help_text
        })
    except FormQuestion.DoesNotExist:
        return JsonResponse({'error': 'Question not found'}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)

@login_required
@user_passes_test(is_admin_or_crc)
@csrf_exempt
@require_http_methods(["DELETE"])
def api_form_question_delete(request, question_id):
    try:
        question = FormQuestion.objects.get(id=question_id)
        question.delete()
        return JsonResponse({'message': 'Question deleted successfully'})
    except FormQuestion.DoesNotExist:
        return JsonResponse({'error': 'Question not found'}, status=404)

# Dynamic Form Submissions API
@login_required
@user_passes_test(is_admin_or_crc)
@require_http_methods(["GET"])
def api_dynamic_submissions(request):
    form_type = request.GET.get('form_type')
    faculty_id = request.GET.get('faculty_id')
    course_id = request.GET.get('course_id')
    
    submissions = DynamicFormSubmission.objects.all().select_related(
        'faculty', 'course', 'dynamic_form'
    )
    
    # Apply filters
    if form_type:
        submissions = submissions.filter(dynamic_form__form_type=form_type)
    if faculty_id:
        submissions = submissions.filter(faculty_id=faculty_id)
    if course_id:
        submissions = submissions.filter(course_id=course_id)
    
    submissions_list = list(submissions.order_by('-submission_date').values(
        'id', 'faculty__username', 'faculty__email', 'course__title', 'course__code',
        'course_coordinator', 'is_coordinator', 'status', 'submission_date', 
        'dynamic_form__name', 'dynamic_form__form_type', 'section'
    ))
    
    return JsonResponse(submissions_list, safe=False)

# Get active form and questions for faculty
@login_required
@require_http_methods(["GET"])
def api_faculty_dynamic_forms(request):
    """Get active forms for faculty - UNIVERSAL FORMS"""
    if request.user.role != User.ROLE_FACULTY:
        return JsonResponse({'error': 'Access denied'}, status=403)
    
    try:
        form_type = request.GET.get('form_type')  # 'ccr' or 'crr'
        course_id = request.GET.get('course_id')
        
        # If course_id is provided but form_type is not, try to determine form_type
        if course_id and not form_type:
            try:
                course_assignment = CourseFaculty.objects.get(
                    faculty=request.user,
                    course_id=course_id
                )
                # If user is coordinator, default to CCR, otherwise CRR
                form_type = 'ccr' if course_assignment.is_coordinator else 'crr'
            except CourseFaculty.DoesNotExist:
                return JsonResponse({'error': 'Course not assigned to you'}, status=400)
        
        if not form_type:
            return JsonResponse({'error': 'form_type is required'}, status=400)
        
        # Validate form type
        if form_type not in ['ccr', 'crr']:
            return JsonResponse({'error': 'Invalid form type. Must be "ccr" or "crr"'}, status=400)
        
        # Check if user can access this form type
        if form_type == 'ccr':
            is_coordinator_for_any = CourseFaculty.objects.filter(
                faculty=request.user,
                is_coordinator=True
            ).exists()
            if not is_coordinator_for_any:
                return JsonResponse({'error': 'Only course coordinators can access CCR forms'}, status=403)
        
        # Get ALL active forms of the appropriate type
        active_forms = DynamicForm.objects.filter(
            status=DynamicForm.STATUS_ACTIVE,
            form_type=form_type
        )
        
        if not active_forms.exists():
            return JsonResponse({
                'active': False, 
                'message': f'No active {form_type.upper()} form available'
            })
        
        # Get all courses assigned to this faculty
        course_assignments = CourseFaculty.objects.filter(
            faculty=request.user
        ).select_related('course').prefetch_related('sections').order_by('course__code')
        
        assigned_courses = []
        for assignment in course_assignments:
            # For CCR form, only show courses where user is coordinator
            if form_type == 'ccr' and not assignment.is_coordinator:
                continue
                
            course_payload = {
                'id': assignment.course.id,
                'code': assignment.course.code,
                'title': assignment.course.title,
                'is_coordinator': assignment.is_coordinator,
            }
            course_payload.update(course_faculty_section_payload(assignment))
            assigned_courses.append(course_payload)
        
        # Get questions for ALL active forms
        forms_with_questions = []
        for form in active_forms:
            questions = list(FormQuestion.objects.filter(form=form).order_by('order').values(
                'id', 'question_text', 'question_type', 'required', 'options', 'config', 'help_text'
            ))
            
            forms_with_questions.append({
                'id': form.id,
                'name': form.name,
                'description': form.description,
                'form_type': form.form_type,
                'questions': questions
            })
        
        return JsonResponse({
            'active': True,
            'forms': forms_with_questions,  # Now returns array of forms
            'assigned_courses': assigned_courses
        })
        
    except Exception as e:
        print(f"Error in api_faculty_dynamic_forms: {str(e)}")  # For debugging
        return JsonResponse({'error': str(e)}, status=400)

# Check form availability for faculty

@login_required
@require_http_methods(["GET"])
def api_form_availability(request):
    """Check which universal forms (CCR/CRR) are available for faculty"""
    if request.user.role != User.ROLE_FACULTY:
        return JsonResponse({'error': 'Access denied'}, status=403)
    
    try:
        # Get ALL active forms for CCR and CRR
        ccr_forms = list(DynamicForm.objects.filter(
            status=DynamicForm.STATUS_ACTIVE,
            form_type='ccr'
        ).values('id', 'name', 'description', 'created_at').order_by('-created_at'))
        
        crr_forms = list(DynamicForm.objects.filter(
            status=DynamicForm.STATUS_ACTIVE,
            form_type='crr'
        ).values('id', 'name', 'description', 'created_at').order_by('-created_at'))
        
        # Check if user is coordinator for ANY course
        is_coordinator_for_any = CourseFaculty.objects.filter(
            faculty=request.user,
            is_coordinator=True
        ).exists()
        
        # Get assigned courses
        course_assignments = CourseFaculty.objects.filter(
            faculty=request.user
        ).select_related('course')
        
        courses_data = []
        for assignment in course_assignments:
            courses_data.append({
                'course_id': assignment.course.id,
                'course_code': assignment.course.code,
                'course_title': assignment.course.title,
                'is_coordinator': assignment.is_coordinator,
                'section': assignment.section_display(),
                'forms_available': {
                    'ccr': ccr_forms if assignment.is_coordinator else [],
                    'crr': crr_forms
                }
            })
        
        return JsonResponse({
            'active_forms': {
                'ccr': ccr_forms,
                'crr': crr_forms
            },
            'user_can_submit_ccr': len(ccr_forms) > 0 and is_coordinator_for_any,
            'user_can_submit_crr': len(crr_forms) > 0,
            'courses': courses_data,
            'status': 'success'
        })
    except Exception as e:
        print(f"Error in api_form_availability: {str(e)}")
        return JsonResponse({
            'error': str(e),
            'status': 'error',
            'active_forms': {'ccr': [], 'crr': []},
            'user_can_submit_ccr': False,
            'user_can_submit_crr': False,
            'courses': []
        }, status=400)

# Submit universal form
@login_required
@csrf_exempt
@require_http_methods(["POST"])
def api_submit_dynamic_form(request):
    """Submit UNIVERSAL form (CCR/CRR) for a specific course section"""
    if request.user.role != User.ROLE_FACULTY:
        return JsonResponse({'error': 'Access denied'}, status=403)
    
    try:
        data = json.loads(request.body)
        course_id = data.get('course_id')
        form_id = data.get('form_id')  # Now form_id is required since multiple forms
        section_id = data.get('section_id')
        answers = data.get('answers', {})
        status = data.get('status', 'draft')  # 'draft' or 'submitted'
        
        print(f"Form submission attempt: user={request.user.username}, course_id={course_id}, form_id={form_id}, section_id={section_id}, status={status}")
        
        if not form_id:
            return JsonResponse({'error': 'form_id is required. Multiple forms may be active.'}, status=400)
        
        # Validate course assignment
        try:
            course_assignment = CourseFaculty.objects.prefetch_related('sections').get(
                faculty=request.user,
                course_id=course_id
            )
        except CourseFaculty.DoesNotExist:
            return JsonResponse({'error': 'Course not assigned to you'}, status=400)
        
        # Get the SPECIFIC UNIVERSAL form
        try:
            form = DynamicForm.objects.get(id=form_id, form_type__in=['ccr', 'crr'])
        except DynamicForm.DoesNotExist:
            return JsonResponse({'error': 'Form not found or not a universal form'}, status=404)
        
        # Check if form is active
        if form.status != 'active':
            return JsonResponse({'error': 'Form is not active'}, status=400)
        
        # Check if form type matches coordinator status
        if form.form_type == 'ccr' and not course_assignment.is_coordinator:
            return JsonResponse({'error': 'Only course coordinators can submit CCR forms'}, status=403)
        
        # Get course
        course = Course.objects.get(id=course_id)

        assigned_section = None
        assigned_sections = list(course_assignment.sections.all())
        if assigned_sections:
            if not section_id:
                return JsonResponse({
                    'error': 'section_id is required. Submit this form for one assigned section.',
                }, status=400)
            try:
                section_id = int(section_id)
            except (TypeError, ValueError):
                return JsonResponse({'error': 'Invalid section_id'}, status=400)
            assigned_section = next(
                (section for section in assigned_sections if section.id == section_id),
                None,
            )
            if not assigned_section:
                return JsonResponse({
                    'error': 'Selected section is not assigned to you for this course.',
                }, status=400)
        elif section_id:
            # Faculty has no sections on assignment; ignore unexpected section_id
            section_id = None
        
        # Check if already submitted for this course + section
        existing_submission = DynamicFormSubmission.objects.filter(
            faculty=request.user,
            course_id=course_id,
            dynamic_form=form,
            assigned_section=assigned_section,
        ).first()
        
        # Allow resubmission if in draft or revision requested status
        # Only block if trying to submit when already submitted
        if existing_submission and existing_submission.status == 'submitted' and status == 'submitted':
            section_label = assigned_section.code if assigned_section else 'this course'
            return JsonResponse({
                'error': f'You have already submitted this form for section {section_label}.',
                'submission_id': existing_submission.id,
                'status': existing_submission.status
            }, status=400)
        
        section_code = assigned_section.code if assigned_section else ''

        # Create or update submission
        if existing_submission:
            submission = existing_submission
            submission.status = status
            submission.assigned_section = assigned_section
            submission.section = section_code
            
            # Update submission date if submitting (not draft)
            if status == 'submitted':
                submission.submission_date = datetime.now()
            
            submission.save()
            
            # Delete existing answers
            FormAnswer.objects.filter(submission=submission).delete()
        else:
            submission = DynamicFormSubmission.objects.create(
                dynamic_form=form,
                faculty=request.user,
                course=course,
                assigned_section=assigned_section,
                course_code_title=f"{course.code} - {course.title}",
                course_coordinator=request.user.username if course_assignment.is_coordinator else "",
                is_coordinator=course_assignment.is_coordinator,
                section=section_code,
                status=status,
                submission_date=datetime.now() if status == 'submitted' else None
            )
        
        # Create answers
        for question_id, answer_value in answers.items():
            try:
                question = FormQuestion.objects.get(id=question_id, form=form)
                
                # Handle different answer types
                if isinstance(answer_value, (list, dict)):
                    FormAnswer.objects.create(
                        submission=submission,
                        question=question,
                        answer_text="",
                        answer_data=answer_value
                    )
                else:
                    FormAnswer.objects.create(
                        submission=submission,
                        question=question,
                        answer_text=str(answer_value) if answer_value is not None else "",
                        answer_data=None
                    )
            except FormQuestion.DoesNotExist:
                print(f"Warning: Question {question_id} not found in form {form_id}")
                continue
        
        return JsonResponse({
            'message': f'Form {"submitted" if status == "submitted" else "saved as draft"} successfully!',
            'submission_id': submission.id,
            'status': submission.status,
            'section_id': assigned_section.id if assigned_section else None,
            'section': section_code,
            'success': True
        })
        
    except Course.DoesNotExist:
        return JsonResponse({'error': 'Course not found'}, status=404)
    except Exception as e:
        print(f"Error in form submission: {str(e)}")
        import traceback
        traceback.print_exc()
        return JsonResponse({'error': f'Submission failed: {str(e)}'}, status=400)

# Faculty Users API
@login_required
@user_passes_test(is_admin_or_crc)
@require_http_methods(["GET"])
def api_faculty_users(request):
    try:
        faculty_users = list(User.objects.filter(role=User.ROLE_FACULTY).values(
            'id', 'username', 'email', 'department', 'designation'
        ))
        return JsonResponse({'results': faculty_users})
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)

# Course Faculty Assignments API
@login_required
@user_passes_test(is_admin_or_crc)
@require_http_methods(["GET"])
def api_course_faculty_assignments(request, course_id):
    try:
        assignments = CourseFaculty.objects.filter(course_id=course_id).prefetch_related('sections')
        payload = []
        for assignment in assignments:
            row = {
                'faculty_id': assignment.faculty_id,
                'is_coordinator': assignment.is_coordinator,
            }
            row.update(course_faculty_section_payload(assignment))
            payload.append(row)
        return JsonResponse(payload, safe=False)
    except Exception as e:
        return JsonResponse([], safe=False)


@login_required
@user_passes_test(is_admin_or_crc)
@require_http_methods(["GET"])
def api_faculty_course_assignments(request, user_id):
    """Return courses currently assigned to a faculty member (for Assign Courses UI)."""
    try:
        faculty_member = User.objects.get(id=user_id, role=User.ROLE_FACULTY)
        assignments = CourseFaculty.objects.filter(faculty=faculty_member).select_related(
            'course', 'course__department'
        ).prefetch_related('sections')
        payload = []
        for assignment in assignments:
            row = {
                'course_id': assignment.course_id,
                'course_code': assignment.course.code,
                'course_title': assignment.course.title,
                'is_coordinator': assignment.is_coordinator,
                'department_name': (
                    assignment.course.department.name if assignment.course.department else ''
                ),
                'department_code': (
                    assignment.course.department.code if assignment.course.department else ''
                ),
            }
            row.update(course_faculty_section_payload(assignment))
            payload.append(row)
        return JsonResponse(payload, safe=False)
    except User.DoesNotExist:
        return JsonResponse({'error': 'Faculty member not found'}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)

# User API Views
@login_required
@user_passes_test(is_admin)
@require_http_methods(["GET"])
def api_users(request):
    users = list(User.objects.values('id', 'username', 'email', 'role', 'department', 'designation'))
    return JsonResponse({'results': users})

@login_required
@user_passes_test(is_admin)
@csrf_exempt
@require_http_methods(["POST"])
def api_users_create(request):
    try:
        data = json.loads(request.body)
        user = User.objects.create_user(
            username=data.get('username'),
            email=data.get('email'),
            password=data.get('password'),
            role=data.get('role', User.ROLE_FACULTY),
            department=data.get('department', ''),
            designation=data.get('designation', '')
        )
        return JsonResponse({
            'id': user.id,
            'username': user.username,
            'email': user.email,
            'role': user.role,
            'department': user.department,
            'designation': user.designation
        }, status=201)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)

@login_required
@user_passes_test(is_admin)
@csrf_exempt
@require_http_methods(["PUT"])
def api_user_update(request, user_id):
    try:
        data = json.loads(request.body)
        user = User.objects.get(id=user_id)
        user.username = data.get('username', user.username)
        user.email = data.get('email', user.email)
        if 'password' in data and data['password']:
            user.set_password(data['password'])
        user.role = data.get('role', user.role)
        user.department = data.get('department', user.department)
        user.designation = data.get('designation', user.designation)
        user.save()
        return JsonResponse({
            'id': user.id,
            'username': user.username,
            'email': user.email,
            'role': user.role,
            'department': user.department,
            'designation': user.designation
        })
    except User.DoesNotExist:
        return JsonResponse({'error': 'User not found'}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)

@login_required
@user_passes_test(is_admin)
@csrf_exempt
@require_http_methods(["DELETE"])
def api_user_delete(request, user_id):
    try:
        user = User.objects.get(id=user_id)
        if user.id == request.user.id:
            return JsonResponse({'error': 'Cannot delete your own account'}, status=400)
        user.delete()
        return JsonResponse({'message': 'User deleted successfully'})
    except User.DoesNotExist:
        return JsonResponse({'error': 'User not found'}, status=404)

# CRC Specific APIs
@login_required
@user_passes_test(is_admin_or_crc)
@require_http_methods(["GET"])
def api_crc_faculty_list(request):
    """Get all faculty with their assigned courses and coordinator status"""
    try:
        faculty_list = []
        
        # Get all faculty users
        faculty_users = User.objects.filter(role=User.ROLE_FACULTY).prefetch_related('coursefaculty_set__course')
        
        for faculty in faculty_users:
            # Get course assignments
            assignments = CourseFaculty.objects.filter(faculty=faculty).select_related('course')
            
            assigned_courses = []
            coordinator_courses = []
            
            for assignment in assignments:
                course_data = {
                    'id': assignment.course.id,
                    'code': assignment.course.code,
                    'title': assignment.course.title,
                    'section': assignment.section_display(),
                    'is_coordinator': assignment.is_coordinator
                }
                
                assigned_courses.append(course_data)
                if assignment.is_coordinator:
                    coordinator_courses.append(course_data)
            
            faculty_list.append({
                'id': faculty.id,
                'username': faculty.username,
                'email': faculty.email,
                'department': faculty.department,
                'designation': faculty.designation,
                'total_courses': len(assigned_courses),
                'coordinator_courses_count': len(coordinator_courses),
                'assigned_courses': assigned_courses,
                'coordinator_courses': coordinator_courses
            })
        
        return JsonResponse({
            'total_faculty': len(faculty_list),
            'faculty': faculty_list
        })
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)

@login_required
@user_passes_test(is_admin_or_crc)
@require_http_methods(["GET"])
def api_crc_course_catalogue(request):
    """Get all course outlines (both old and new)"""
    try:
        # Get all courses with their outlines
        courses = Course.objects.prefetch_related('outlines').all()
        
        catalogue = []
        for course in courses:
            # Get current official outline
            current_outline = course.outlines.filter(is_current=True, status='approved').first()
            
            # Get all outlines for this course
            all_outlines = list(course.outlines.order_by('-version').values(
                'id', 'version', 'title', 'status', 'is_current', 
                'faculty__username', 'created_at', 'submitted_at', 'approved_at'
            ))
            
            catalogue.append({
                'course_id': course.id,
                'course_code': course.code,
                'course_title': course.title,
                'department': course.department.name if course.department else '',
                'current_outline': {
                    'id': current_outline.id if current_outline else None,
                    'version': current_outline.version if current_outline else None,
                    'title': current_outline.title if current_outline else '',
                    'faculty': current_outline.faculty.username if current_outline else ''
                } if current_outline else None,
                'total_outlines': len(all_outlines),
                'outlines': all_outlines,
                **course_catalogue_file_payload(course),
            })
        
        return JsonResponse(catalogue, safe=False)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)

@login_required
@user_passes_test(is_admin_or_crc)
@csrf_exempt
@require_http_methods(["POST"])
def api_crc_update_course_outline(request):
    """Update course outline (edit/upload new version)"""
    try:
        data = json.loads(request.body)
        outline_id = data.get('outline_id')
        action = data.get('action')  # 'approve', 'request_revision', 'update'
        notes = data.get('notes', '')
        
        if action == 'approve':
            outline = CourseOutline.objects.get(id=outline_id)
            
            # Set this outline as approved and current
            outline.status = 'approved'
            outline.approved_at = datetime.now()
            outline.is_current = True
            outline.notes = notes
            
            # Mark other outlines as not current
            CourseOutline.objects.filter(course=outline.course).exclude(id=outline_id).update(is_current=False)
            
            outline.save()
            
            return JsonResponse({
                'message': 'Course outline approved and set as current',
                'outline_id': outline.id,
                'status': outline.status,
                'approved_at': outline.approved_at.isoformat()
            })
            
        elif action == 'request_revision':
            outline = CourseOutline.objects.select_related(
                'faculty', 'course'
            ).get(id=outline_id)
            
            # Only request revision on submitted outlines
            if outline.status != 'submitted':
                return JsonResponse({
                    'error': f'Cannot request revision on outline with status: {outline.status}'
                }, status=400)
            
            outline.status = 'revision_requested'
            outline.notes = notes
            outline.save()

            email_sent = notify_faculty_outline_revision(
                outline=outline,
                notes=notes,
                request=request,
                requested_by=request.user,
            )
            
            return JsonResponse({
                'message': 'Revision requested for course outline',
                'outline_id': outline.id,
                'status': outline.status,
                'notes': outline.notes,
                'faculty_notified': email_sent,
            })
            
        elif action == 'update':
            outline = CourseOutline.objects.get(id=outline_id)
            outline.title = data.get('title', outline.title)
            outline.description = data.get('description', outline.description)
            outline.content = data.get('content', outline.content)
            outline.notes = data.get('notes', outline.notes)
            outline.save()
            
            return JsonResponse({
                'message': 'Course outline updated',
                'outline_id': outline.id
            })
            
        else:
            return JsonResponse({'error': 'Invalid action'}, status=400)
            
    except CourseOutline.DoesNotExist:
        return JsonResponse({'error': 'Course outline not found'}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)

@login_required
@require_http_methods(["GET"])
def api_check_outline_permissions(request, outline_id):
    """Check if faculty can edit a course outline"""
    if request.user.role != User.ROLE_FACULTY:
        return JsonResponse({'error': 'Access denied'}, status=403)
    
    try:
        outline = CourseOutline.objects.get(id=outline_id, faculty=request.user)
        
        can_edit = outline.status in ['draft', 'revision_requested']
        
        return JsonResponse({
            'can_edit': can_edit,
            'status': outline.status,
            'notes': outline.notes if not can_edit else '',
            'message': 'Can edit' if can_edit else f'Cannot edit: Outline is {outline.status}'
        })
    except CourseOutline.DoesNotExist:
        return JsonResponse({'error': 'Outline not found or access denied'}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)

@login_required
@user_passes_test(is_admin_or_crc)
@require_http_methods(["GET"])
def api_crc_course_outline_submissions(request):
    """Get all course outline submissions (new outlines waiting for approval)"""
    try:
        submissions = CourseOutline.objects.filter(
            status__in=['submitted', 'revision_requested']
        ).select_related('course', 'faculty').order_by('-submitted_at')
        
        submissions_list = list(submissions.values(
            'id', 'course__code', 'course__title', 'faculty__username',
            'version', 'title', 'status', 'submitted_at', 'notes'
        ))
        
        return JsonResponse(submissions_list, safe=False)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)


@login_required
@user_passes_test(is_admin_or_crc)
@require_http_methods(["GET"])
def api_crc_form_submissions(request):
    """Get all form submissions with filtering options (UNIVERSAL FORMS ONLY)"""
    try:
        faculty_id = request.GET.get('faculty_id')
        course_id = request.GET.get('course_id')
        department_id = request.GET.get('department_id')
        form_type = request.GET.get('form_type')
        status = request.GET.get('status')
        
        submissions = DynamicFormSubmission.objects.filter(
            dynamic_form__form_type__in=['ccr', 'crr']  # Only universal forms
        ).select_related(
            'faculty', 'course', 'dynamic_form', 'course__department'
        )
        
        # Apply filters (AND when multiple are set)
        if faculty_id:
            submissions = submissions.filter(faculty_id=faculty_id)
        if course_id:
            submissions = submissions.filter(course_id=course_id)
        if department_id:
            submissions = submissions.filter(course__department_id=department_id)
        if form_type:
            submissions = submissions.filter(dynamic_form__form_type=form_type)
        if status:
            submissions = submissions.filter(status=status)
        
        submissions_list = list(submissions.order_by('-submission_date').values(
            'id', 'faculty__username', 'faculty__email', 'faculty__department',
            'course__code', 'course__title', 'course__department__name',
            'dynamic_form__name', 'dynamic_form__form_type',
            'status', 'is_coordinator', 'submission_date', 'section'
        ))
        
        return JsonResponse(submissions_list, safe=False)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)

@login_required
@csrf_exempt
@require_http_methods(["POST"])
def api_save_course_outline(request):
    """Save or update course outline (faculty coordinators or CRC members)."""
    if request.user.role not in (User.ROLE_FACULTY, User.ROLE_CRC_MEMBER):
        return JsonResponse({'error': 'Access denied'}, status=403)
    
    try:
        data = json.loads(request.body)
        course_id = data.get('course_id')
        outline_id = data.get('outline_id')
        content = data.get('content')
        title = data.get('title', 'Course Outline')
        description = data.get('description', '')
        status = data.get('status', 'draft')  # 'draft' or 'submitted'
        
        if not course_id:
            return JsonResponse({'error': 'course_id is required'}, status=400)
        
        if not content:
            return JsonResponse({'error': 'content is required'}, status=400)
        
        if request.user.role == User.ROLE_FACULTY:
            try:
                CourseFaculty.objects.get(
                    faculty=request.user,
                    course_id=course_id,
                    is_coordinator=True
                )
            except CourseFaculty.DoesNotExist:
                return JsonResponse({'error': 'Only course coordinators can save course outlines'}, status=403)
        
        course = Course.objects.get(id=course_id)
        
        if outline_id:
            # Update existing outline
            outline = CourseOutline.objects.get(id=outline_id, faculty=request.user)
            
            # WORKFLOW RESTRICTIONS
            if outline.status == 'approved':
                return JsonResponse({
                    'error': 'Cannot edit an approved outline. Create a new version instead.',
                    'approved': True
                }, status=400)
            
            if outline.status == 'submitted' and status == 'draft':
                return JsonResponse({
                    'error': 'Cannot change submitted outline back to draft. Wait for CRC review.',
                    'submitted': True
                }, status=400)
            
            # Allow editing only in draft or revision_requested status
            if outline.status not in ['draft', 'revision_requested'] and status == 'draft':
                return JsonResponse({
                    'error': f'Cannot edit outline in {outline.status} status',
                    'current_status': outline.status
                }, status=400)
            
            # If CRC requested revision, allow saving as draft or submitted
            if outline.status == 'revision_requested':
                # Clear notes when faculty starts editing (only if saving as draft)
                if status == 'draft':
                    outline.notes = ''
            
            outline.title = title
            outline.description = description
            outline.content = content
            outline.status = status
            
            if status == 'submitted' and outline.status != 'submitted':
                outline.submitted_at = datetime.now()
            
            outline.save()
        else:
            # Create new outline
            # Get latest version
            latest_version = CourseOutline.objects.filter(
                course=course, faculty=request.user
            ).order_by('-version').first()
            
            new_version = (latest_version.version + 1) if latest_version else 1
            
            outline = CourseOutline.objects.create(
                course=course,
                faculty=request.user,
                version=new_version,
                title=title,
                description=description,
                content=content,
                status=status
            )
            
            if status == 'submitted':
                outline.submitted_at = datetime.now()
                outline.save()
        
        return JsonResponse({
            'message': 'Course outline saved successfully',
            'outline_id': outline.id,
            'version': outline.version,
            'status': outline.status,
            'submitted_at': outline.submitted_at.isoformat() if outline.submitted_at else None,
            'notes': outline.notes,
            'can_edit': outline.status in ['draft', 'revision_requested']
        })
        
    except Course.DoesNotExist:
        return JsonResponse({'error': 'Course not found'}, status=404)
    except CourseOutline.DoesNotExist:
        return JsonResponse({'error': 'Course outline not found'}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)


@login_required
@user_passes_test(is_admin_or_crc)
@require_http_methods(["GET"])
def api_crc_view_outline_content(request, outline_id):
    """CRC member can view the actual content of a course outline"""
    try:
        outline = CourseOutline.objects.get(id=outline_id)
        
        # Parse content if it's JSON, otherwise return as-is
        content = outline.content
        try:
            if content and isinstance(content, str):
                content = json.loads(content)
        except:
            pass  # Keep as string if not JSON
        
        return JsonResponse({
            'id': outline.id,
            'course': {
                'id': outline.course.id,
                'code': outline.course.code,
                'title': outline.course.title,
                'credits': outline.course.credits,
                'department': outline.course.department.name if outline.course.department else ''
            },
            'faculty': {
                'id': outline.faculty.id,
                'username': outline.faculty.username,
                'email': outline.faculty.email,
                'department': outline.faculty.department,
                'designation': outline.faculty.designation
            },
            'version': outline.version,
            'title': outline.title,
            'description': outline.description,
            'content': content,
            'status': outline.status,
            'is_current': outline.is_current,
            'notes': outline.notes,
            'created_at': outline.created_at.isoformat() if outline.created_at else None,
            'submitted_at': outline.submitted_at.isoformat() if outline.submitted_at else None,
            'approved_at': outline.approved_at.isoformat() if outline.approved_at else None,
            'can_edit': outline.status in ['draft', 'revision_requested']  # For faculty reference
        })
    except CourseOutline.DoesNotExist:
        return JsonResponse({'error': 'Course outline not found'}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)

@login_required
@require_http_methods(["GET"])
def api_get_course_outline(request):
    """Get course outline: faculty (assigned/coordinator rules); CRC (any outline by id, own drafts by course)."""
    if request.user.role not in (User.ROLE_FACULTY, User.ROLE_CRC_MEMBER):
        return JsonResponse({'error': 'Access denied'}, status=403)

    outline_id = request.GET.get('outline_id')
    course_id = request.GET.get('course_id')

    try:
        if request.user.role == User.ROLE_CRC_MEMBER:
            if outline_id:
                outline = CourseOutline.objects.select_related('course', 'faculty').get(
                    id=outline_id
                )
                return JsonResponse(course_outline_to_dict(outline))
            if not course_id:
                return JsonResponse(
                    {'error': 'course_id or outline_id is required'},
                    status=400,
                )
            Course.objects.get(id=course_id)
            outline = CourseOutline.objects.filter(
                course_id=course_id,
                faculty=request.user,
            ).select_related('course', 'faculty').order_by('-version').first()
            if outline:
                return JsonResponse(course_outline_to_dict(outline))
            return JsonResponse({'exists': False})

        def faculty_assigned_to_course(course_pk):
            return CourseFaculty.objects.filter(
                faculty=request.user,
                course_id=course_pk,
            ).exists()

        def faculty_is_coordinator_for_course(course_pk):
            return CourseFaculty.objects.filter(
                faculty=request.user,
                course_id=course_pk,
                is_coordinator=True,
            ).exists()

        if outline_id:
            outline = CourseOutline.objects.select_related('course', 'faculty').get(
                id=outline_id
            )
            if not faculty_assigned_to_course(outline.course_id):
                return JsonResponse({'error': 'Access denied'}, status=403)
            return JsonResponse(course_outline_to_dict(outline))

        if not course_id:
            return JsonResponse(
                {'error': 'course_id or outline_id is required'},
                status=400,
            )

        if not faculty_assigned_to_course(course_id):
            return JsonResponse({'error': 'Access denied'}, status=403)

        if faculty_is_coordinator_for_course(course_id):
            outline = CourseOutline.objects.filter(
                course_id=course_id,
                faculty=request.user,
            ).select_related('course', 'faculty').order_by('-version').first()
        else:
            outline = (
                CourseOutline.objects.filter(
                    course_id=course_id,
                    is_current=True,
                )
                .select_related('course', 'faculty')
                .order_by('-version')
                .first()
            )
            if not outline:
                outline = (
                    CourseOutline.objects.filter(
                        course_id=course_id,
                        status=CourseOutline.STATUS_APPROVED,
                    )
                    .select_related('course', 'faculty')
                    .order_by('-version')
                    .first()
                )
            if not outline:
                outline = (
                    CourseOutline.objects.filter(course_id=course_id)
                    .select_related('course', 'faculty')
                    .order_by('-version')
                    .first()
                )

        if outline:
            return JsonResponse(course_outline_to_dict(outline))
        return JsonResponse({'exists': False})

    except Course.DoesNotExist:
        return JsonResponse({'error': 'Course not found'}, status=404)
    except CourseOutline.DoesNotExist:
        return JsonResponse({'error': 'Course outline not found'}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)



# Submission Details API
@login_required
@require_http_methods(["GET"])
def api_submission_details(request, submission_id):
    """Get detailed information about a specific submission"""
    try:
        submission = DynamicFormSubmission.objects.get(id=submission_id)
        
        # Check permissions
        if request.user.role == User.ROLE_FACULTY and submission.faculty != request.user:
            return JsonResponse({'error': 'Access denied'}, status=403)
        
        # Get all answers for this submission
        answers = FormAnswer.objects.filter(submission=submission).select_related('question')
        
        answers_data = []
        for answer in answers:
            answers_data.append({
                'question_id': answer.question.id,
                'question_text': answer.question.question_text,
                'question_type': answer.question.question_type,
                'answer_text': answer.answer_text,
                'answer_data': answer.answer_data
            })
        
        return JsonResponse({
            'id': submission.id,
            'course_code': submission.course.code,
            'course_title': submission.course.title,
            'form_name': submission.dynamic_form.name,
            'form_type': submission.dynamic_form.form_type,
            'faculty_name': submission.faculty.username,
            'faculty_email': submission.faculty.email,
            'faculty_department': submission.faculty.department,
            'is_coordinator': submission.is_coordinator,
            'status': submission.status,
            'section': submission.section,
            'submission_date': submission.submission_date.strftime('%B %d, %Y at %I:%M %p'),
            'answers': answers_data
        })
        
    except DynamicFormSubmission.DoesNotExist:
        return JsonResponse({'error': 'Submission not found'}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)


def _format_submission_answer_for_pdf(answer):
    """Flatten a FormAnswer into plain text for PDF output."""
    question_type = answer.question.question_type
    answer_text = answer.answer_text or ""
    answer_data = answer.answer_data

    if question_type == "clo_percentage":
        clo_rows = _extract_clo_rows_for_pdf(answer)
        if not clo_rows:
            return "No CLO data provided"
        return clo_rows

    if isinstance(answer_data, list):
        return ", ".join(str(item) for item in answer_data) or "No answer provided"
    if isinstance(answer_data, dict):
        return ", ".join(f"{key}: {value}" for key, value in answer_data.items()) or (
            answer_text or "No answer provided"
        )
    if answer_text:
        return answer_text
    if answer_data is not None:
        return str(answer_data)
    return "No answer provided"


def _extract_clo_rows_for_pdf(answer):
    """Return sorted CLO rows as [(label, percentage), ...] plus optional average."""
    clo_data = {}
    answer_data = answer.answer_data
    answer_text = answer.answer_text or ""

    if isinstance(answer_data, dict):
        clo_data = answer_data
    elif answer_text:
        try:
            parsed = json.loads(answer_text)
            if isinstance(parsed, dict):
                clo_data = parsed
        except (TypeError, ValueError, json.JSONDecodeError):
            pairs = answer_text.split(",")
            for pair in pairs:
                if ":" not in pair:
                    continue
                key, value = pair.split(":", 1)
                try:
                    clo_data[key.strip()] = float(value.strip().replace("%", ""))
                except (TypeError, ValueError):
                    continue

    if not clo_data:
        return []

    def clo_sort_key(key):
        digits = "".join(ch for ch in str(key) if ch.isdigit())
        return int(digits) if digits else 0

    rows = []
    numeric_values = []
    for key in sorted(clo_data.keys(), key=clo_sort_key):
        raw_value = clo_data[key]
        try:
            if isinstance(raw_value, str):
                percentage = float(raw_value.replace("%", "").strip())
            else:
                percentage = float(raw_value)
        except (TypeError, ValueError):
            rows.append((str(key), str(raw_value)))
            continue
        numeric_values.append(percentage)
        rows.append((str(key), f"{percentage}%"))

    if len(numeric_values) > 1:
        average = round(sum(numeric_values) / len(numeric_values), 1)
        rows.append(("Average", f"{average}%"))

    return rows


def _write_clo_table_to_pdf(pdf_document, clo_rows, font_name):
    """Render CLO / Percentage table matching the view popup layout."""
    if not clo_rows:
        return

    col_widths = (45, 45)
    row_height = 8
    pdf_document.set_x(pdf_document.l_margin)
    pdf_document.set_font(font_name, style="B", size=10)
    pdf_document.set_fill_color(249, 250, 251)
    pdf_document.cell(col_widths[0], row_height, "CLO", border=1, fill=True)
    pdf_document.cell(col_widths[1], row_height, "Percentage", border=1, fill=True, new_x="LMARGIN", new_y="NEXT")

    for label, percentage in clo_rows:
        is_average = str(label).lower() == "average"
        pdf_document.set_font(font_name, style="B" if is_average else "", size=10)
        if is_average:
            pdf_document.set_fill_color(249, 250, 251)
            fill = True
        else:
            fill = False
        pdf_document.set_x(pdf_document.l_margin)
        pdf_document.cell(col_widths[0], row_height, str(label), border=1, fill=fill)
        pdf_document.cell(
            col_widths[1],
            row_height,
            str(percentage),
            border=1,
            fill=fill,
            new_x="LMARGIN",
            new_y="NEXT",
        )
    pdf_document.ln(2)


@login_required
@require_http_methods(["GET"])
def api_submission_pdf(request, submission_id):
    """Download a form submission as a PDF (faculty: own only; CRC/admin: any)."""
    try:
        submission = DynamicFormSubmission.objects.select_related(
            "course", "dynamic_form", "faculty"
        ).get(id=submission_id)
    except DynamicFormSubmission.DoesNotExist:
        return JsonResponse({"error": "Submission not found"}, status=404)

    if request.user.role == User.ROLE_FACULTY and submission.faculty_id != request.user.id:
        return JsonResponse({"error": "Access denied"}, status=403)
    if request.user.role not in (
        User.ROLE_FACULTY,
        User.ROLE_ADMIN,
        User.ROLE_CRC_MEMBER,
    ):
        return JsonResponse({"error": "Access denied"}, status=403)

    answers = (
        FormAnswer.objects.filter(submission=submission)
        .select_related("question")
        .order_by("question__order", "question__id")
    )

    pdf_document = FPDF(orientation="P", unit="mm", format="A4")
    pdf_document.set_auto_page_break(auto=True, margin=15)
    pdf_document.set_left_margin(14)
    pdf_document.set_right_margin(14)
    pdf_document.add_page()

    use_unicode_font = _register_dejavu_fonts_if_available(pdf_document)
    title_font = "DejaVuSans" if use_unicode_font else "Helvetica"
    body_font = "DejaVuSans" if use_unicode_font else "Helvetica"

    def write_line(text, size=11, bold=False, gap=6):
        style = "B" if bold else ""
        raw_text = "" if text is None else str(text)
        safe_text = (
            raw_text if use_unicode_font else _normalize_html_for_core_pdf_fonts(raw_text)
        )
        pdf_document.set_x(pdf_document.l_margin)
        pdf_document.set_font(title_font if bold else body_font, style=style, size=size)
        # new_x/new_y keep the cursor on the left margin for the next line
        pdf_document.multi_cell(
            0,
            gap,
            safe_text,
            new_x="LMARGIN",
            new_y="NEXT",
        )

    try:
        write_line("Form Submission Details", size=16, bold=True, gap=8)
        pdf_document.ln(2)
        write_line(
            f"Form: {submission.dynamic_form.name} ({submission.dynamic_form.form_type.upper()})",
            bold=True,
        )
        write_line(f"Course: {submission.course.code} - {submission.course.title}")
        write_line(f"Status: {submission.status.replace('_', ' ')}")
        write_line(
            f"Date: {submission.submission_date.strftime('%B %d, %Y at %I:%M %p') if submission.submission_date else 'N/A'}"
        )
        write_line(f"Faculty: {submission.faculty.username}")
        if submission.section:
            write_line(f"Section: {submission.section}")
        pdf_document.ln(4)
        write_line("Answers", size=14, bold=True, gap=7)
        pdf_document.ln(1)

        if not answers.exists():
            write_line("No answers submitted for this form.")
        else:
            for index, answer in enumerate(answers, start=1):
                question_text = answer.question.question_text or "Question"
                write_line(f"{index}. {question_text}", bold=True, gap=5)

                if answer.question.question_type == "clo_percentage":
                    clo_rows = _extract_clo_rows_for_pdf(answer)
                    if clo_rows:
                        _write_clo_table_to_pdf(
                            pdf_document,
                            clo_rows,
                            title_font if use_unicode_font else body_font,
                        )
                    else:
                        write_line("No CLO data provided", gap=5)
                else:
                    answer_text = _format_submission_answer_for_pdf(answer)
                    write_line(answer_text, gap=5)
                pdf_document.ln(2)

        pdf_bytes = bytes(pdf_document.output())
    except Exception as pdf_error:
        print(f"Submission PDF generation failed: {pdf_error}")
        return JsonResponse(
            {"error": f"Failed to generate PDF: {pdf_error}"},
            status=500,
        )

    course_code = (submission.course.code or "course").replace(" ", "_")
    form_type = submission.dynamic_form.form_type.upper()
    filename = f"submission_{course_code}_{form_type}_{submission.id}.pdf"
    response = HttpResponse(pdf_bytes, content_type="application/pdf")
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


@login_required
@require_http_methods(["GET"])
def api_outline_pdf(request, outline_id):
    """Download a course outline as a PDF (faculty: own only; CRC/admin: any)."""
    try:
        outline = CourseOutline.objects.select_related(
            "course", "faculty", "course__department"
        ).get(id=outline_id)
    except CourseOutline.DoesNotExist:
        return JsonResponse({"error": "Course outline not found"}, status=404)

    if request.user.role == User.ROLE_FACULTY and outline.faculty_id != request.user.id:
        return JsonResponse({"error": "Access denied"}, status=403)
    if request.user.role not in (
        User.ROLE_FACULTY,
        User.ROLE_ADMIN,
        User.ROLE_CRC_MEMBER,
    ):
        return JsonResponse({"error": "Access denied"}, status=403)

    course = outline.course
    faculty = outline.faculty
    department_name = course.department.name if course.department else "N/A"
    submitted_at = (
        outline.submitted_at.strftime("%B %d, %Y")
        if outline.submitted_at
        else "Not submitted"
    )
    status_label = (outline.status or "").replace("_", " ").title()

    meta_html = (
        f"<h1>Course Outline</h1>"
        f"<p><b>Title:</b> {html_module.escape(outline.title or 'Course Outline')}</p>"
        f"<p><b>Course:</b> {html_module.escape(course.code)} — "
        f"{html_module.escape(course.title)}</p>"
        f"<p><b>Department:</b> {html_module.escape(department_name)}</p>"
        f"<p><b>Faculty:</b> {html_module.escape(faculty.username)}</p>"
        f"<p><b>Version:</b> {outline.version}</p>"
        f"<p><b>Status:</b> {html_module.escape(status_label)}</p>"
        f"<p><b>Submitted:</b> {html_module.escape(submitted_at)}</p>"
        "<hr/>"
    )
    if outline.notes:
        meta_html += (
            f"<p><b>CRC Notes:</b> {html_module.escape(outline.notes)}</p><hr/>"
        )

    content_html = outline.content or "<p><i>No outline content.</i></p>"
    # Content is already HTML from the editor; wrap for layout.
    combined_html = meta_html + f"<div>{content_html}</div>"

    pdf_document = FPDF(orientation="P", unit="mm", format="A4")
    pdf_document.set_auto_page_break(auto=True, margin=15)
    pdf_document.set_left_margin(14)
    pdf_document.set_right_margin(14)

    try:
        if _register_dejavu_fonts_if_available(pdf_document):
            pdf_document.add_page()
            pdf_document.set_font("DejaVuSans", size=11)
        else:
            combined_html = _normalize_html_for_core_pdf_fonts(combined_html)
            pdf_document.add_page()
            pdf_document.set_font("Helvetica", size=11)

        pdf_document.write_html(combined_html)
        pdf_bytes = bytes(pdf_document.output())
    except Exception as pdf_error:
        print(f"Outline PDF generation failed: {pdf_error}")
        return JsonResponse(
            {"error": f"Failed to generate PDF: {pdf_error}"},
            status=500,
        )

    course_code = (course.code or "course").replace(" ", "_")
    filename = f"outline_{course_code}_v{outline.version}_{outline.id}.pdf"
    response = HttpResponse(pdf_bytes, content_type="application/pdf")
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


# CRC Dashboard Statistics API
@login_required
@user_passes_test(is_admin_or_crc)
@require_http_methods(["GET"])
def api_crc_dashboard_stats(request):
    """Get comprehensive statistics for CRC dashboard (UNIVERSAL FORMS ONLY)"""
    try:
        # Total faculty count
        total_faculty = User.objects.filter(role=User.ROLE_FACULTY).count()
        
        # Total courses count
        total_courses = Course.objects.count()
        
        # Pending course outlines (submitted but not approved)
        pending_outlines = CourseOutline.objects.filter(status='submitted').count()
        
        # Total form submissions (UNIVERSAL FORMS ONLY)
        total_submissions = DynamicFormSubmission.objects.filter(
            status='submitted',
            dynamic_form__form_type__in=['ccr', 'crr']
        ).count()
        
        # Recent submissions (last 7 days) - UNIVERSAL FORMS ONLY
        week_ago = datetime.now() - timedelta(days=7)
        recent_submissions = DynamicFormSubmission.objects.filter(
            status='submitted',
            submission_date__gte=week_ago,
            dynamic_form__form_type__in=['ccr', 'crr']  # Only universal forms
        ).count()
        
        # Faculty with most submissions (UNIVERSAL FORMS ONLY)
        faculty_submission_counts = DynamicFormSubmission.objects.filter(
            status='submitted',
            dynamic_form__form_type__in=['ccr', 'crr']  # Only universal forms
        ).values('faculty__username').annotate(
            count=Count('id')
        ).order_by('-count')[:5]
        
        # Forms status (active/inactive) - ONLY CCR and CRR
        form_status = {
            'ccr_active': DynamicForm.objects.filter(form_type='ccr', status='active').exists(),
            'crr_active': DynamicForm.objects.filter(form_type='crr', status='active').exists(),
        }
        
        # Departments with most courses
        department_stats = Department.objects.annotate(
            course_count=Count('courses')
        ).order_by('-course_count').values('name', 'code', 'course_count')[:5]
        
        return JsonResponse({
            'total_faculty': total_faculty,
            'total_courses': total_courses,
            'pending_outlines': pending_outlines,
            'total_submissions': total_submissions,
            'recent_submissions': recent_submissions,
            'top_faculty_submissions': list(faculty_submission_counts),
            'form_status': form_status,
            'department_stats': list(department_stats),
            'generated_at': datetime.now().isoformat()
        })
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)


# Submission Approval APIs
@login_required
@user_passes_test(is_admin_or_crc)
@csrf_exempt
@require_http_methods(["POST"])
def api_approve_submission(request, submission_id):
    """Approve a form submission"""
    try:
        submission = DynamicFormSubmission.objects.get(id=submission_id)
        submission.status = 'approved'
        submission.save()
        
        return JsonResponse({
            'message': 'Submission approved successfully',
            'submission_id': submission.id,
            'status': submission.status
        })
    except DynamicFormSubmission.DoesNotExist:
        return JsonResponse({'error': 'Submission not found'}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)

@login_required
@user_passes_test(is_admin_or_crc)
@csrf_exempt
@require_http_methods(["POST"])
def api_reject_submission(request, submission_id):
    """Reject a form submission"""
    try:
        submission = DynamicFormSubmission.objects.select_related(
            'faculty', 'course', 'dynamic_form'
        ).get(id=submission_id)
        submission.status = 'revision_requested'
        submission.save()

        email_sent = notify_faculty_form_revision(
            submission=submission,
            notes='',
            request=request,
            requested_by=request.user,
        )
        
        return JsonResponse({
            'message': 'Submission rejected and revision requested',
            'submission_id': submission.id,
            'status': submission.status,
            'faculty_notified': email_sent,
        })
    except DynamicFormSubmission.DoesNotExist:
        return JsonResponse({'error': 'Submission not found'}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)

@login_required
@user_passes_test(is_admin_or_crc)
@csrf_exempt
@require_http_methods(["POST"])
def api_request_revision_submission(request, submission_id):
    """Request revision for a submission"""
    try:
        data = json.loads(request.body)
        notes = data.get('notes', '')
        
        submission = DynamicFormSubmission.objects.select_related(
            'faculty', 'course', 'dynamic_form'
        ).get(id=submission_id)
        submission.status = 'revision_requested'
        
        submission.save()

        email_sent = notify_faculty_form_revision(
            submission=submission,
            notes=notes,
            request=request,
            requested_by=request.user,
        )
        
        return JsonResponse({
            'message': 'Revision requested for submission',
            'submission_id': submission.id,
            'status': submission.status,
            'notes': notes,
            'faculty_notified': email_sent,
        })
    except DynamicFormSubmission.DoesNotExist:
        return JsonResponse({'error': 'Submission not found'}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)

# Admin Faculty Management - Password Reset
@login_required
@user_passes_test(is_admin)
@csrf_exempt
@require_http_methods(["POST"])
def api_admin_reset_faculty_password(request, user_id):
    """Reset faculty password (admin only)"""
    try:
        data = json.loads(request.body)
        new_password = data.get('new_password')
        
        if not new_password:
            return JsonResponse({'error': 'New password is required'}, status=400)
        
        user = User.objects.get(id=user_id, role=User.ROLE_FACULTY)
        user.set_password(new_password)
        user.save()
        
        return JsonResponse({
            'message': f'Password reset successfully for {user.username}',
            'user_id': user.id,
            'username': user.username
        })
    except User.DoesNotExist:
        return JsonResponse({'error': 'Faculty user not found'}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)

# Faculty Submissions API
@login_required
@require_http_methods(["GET"])
def api_faculty_submissions(request):
    """Get all submissions for the current faculty member (UNIVERSAL FORMS ONLY)"""
    if request.user.role != User.ROLE_FACULTY:
        return JsonResponse({'error': 'Access denied'}, status=403)
    
    try:
        submissions = DynamicFormSubmission.objects.filter(
            faculty=request.user,
            dynamic_form__form_type__in=['ccr', 'crr']  # Only universal forms
        ).select_related('course', 'dynamic_form').order_by('-submission_date')
        
        submissions_list = []
        for submission in submissions:
            # Get answer count for this submission
            answer_count = FormAnswer.objects.filter(submission=submission).count()
            
            submissions_list.append({
                'id': submission.id,
                'course': {
                    'id': submission.course.id,
                    'code': submission.course.code,
                    'title': submission.course.title
                },
                'dynamic_form': {
                    'id': submission.dynamic_form.id,
                    'name': submission.dynamic_form.name,
                    'form_type': submission.dynamic_form.form_type
                },
                'status': submission.status,
                'is_coordinator': submission.is_coordinator,
                'submission_date': submission.submission_date.isoformat() if submission.submission_date else None,
                'answer_count': answer_count,
                'section': submission.section
            })
        
        return JsonResponse(submissions_list, safe=False)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)

# Faculty Course Outlines API
@login_required
@require_http_methods(["GET"])
def api_faculty_course_outlines(request):
    """Faculty: outlines for assigned courses. CRC: all outlines that have entered review (not drafts)."""
    if request.user.role not in (User.ROLE_FACULTY, User.ROLE_CRC_MEMBER):
        return JsonResponse({'error': 'Access denied'}, status=403)
    
    try:
        if request.user.role == User.ROLE_CRC_MEMBER:
            outlines = CourseOutline.objects.filter(
                status__in=[
                    CourseOutline.STATUS_SUBMITTED,
                    CourseOutline.STATUS_REVISION,
                    CourseOutline.STATUS_APPROVED,
                ]
            ).select_related('course', 'faculty').order_by('-submitted_at', '-approved_at', '-updated_at', '-id')
        else:
            assigned_course_ids = CourseFaculty.objects.filter(
                faculty=request.user
            ).values_list('course_id', flat=True)

            outlines = CourseOutline.objects.filter(
                course_id__in=assigned_course_ids
            ).select_related('course', 'faculty').order_by('course_id', '-version', '-created_at')
        
        outlines_list = []
        for outline in outlines:
            is_author = outline.faculty_id == request.user.id
            can_edit = is_author and outline.status in [
                CourseOutline.STATUS_DRAFT,
                CourseOutline.STATUS_REVISION,
            ]
            can_submit = can_edit
            outlines_list.append({
                'id': outline.id,
                'course': {
                    'id': outline.course.id,
                    'code': outline.course.code,
                    'title': outline.course.title
                },
                'version': outline.version,
                'title': outline.title,
                'description': outline.description,
                'status': outline.status,
                'is_current': outline.is_current,
                'notes': outline.notes,
                'faculty_author_id': outline.faculty_id,
                'faculty_author_username': outline.faculty.username,
                'can_edit': can_edit,
                'can_submit': can_submit,
                'created_at': outline.created_at.isoformat() if outline.created_at else None,
                'submitted_at': outline.submitted_at.isoformat() if outline.submitted_at else None,
                'approved_at': outline.approved_at.isoformat() if outline.approved_at else None
            })
        
        return JsonResponse(outlines_list, safe=False)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)

# Faculty Profile Update API
@login_required
@csrf_exempt
@require_http_methods(["POST"])
def api_faculty_profile_update(request):
    """Update faculty profile information"""
    if request.user.role != User.ROLE_FACULTY:
        return JsonResponse({'error': 'Access denied'}, status=403)
    
    try:
        data = json.loads(request.body)
        user = request.user
        
        if 'email' in data:
            user.email = data['email']
        
        if 'department' in data:
            user.department = data['department']
        
        if 'designation' in data:
            user.designation = data['designation']
        
        user.save()
        
        return JsonResponse({
            'message': 'Profile updated successfully',
            'user': {
                'id': user.id,
                'username': user.username,
                'email': user.email,
                'department': user.department,
                'designation': user.designation
            }
        })
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)

# Faculty Submissions API (simple version)
@login_required
@require_http_methods(["GET"])
def api_faculty_submissions_list(request):
    """Get all form submissions for the faculty member"""
    if request.user.role != User.ROLE_FACULTY:
        return JsonResponse({'error': 'Access denied'}, status=403)
    
    try:
        submissions = DynamicFormSubmission.objects.filter(
            faculty=request.user
        ).select_related('course', 'dynamic_form').order_by('-submission_date')
        
        submissions_list = []
        for submission in submissions:
            answer_count = FormAnswer.objects.filter(submission=submission).count()
            submissions_list.append({
                'id': submission.id,
                'course': {
                    'id': submission.course.id,
                    'title': submission.course.title,
                    'code': submission.course.code
                },
                'dynamic_form': {
                    'id': submission.dynamic_form.id,
                    'name': submission.dynamic_form.name,
                    'form_type': submission.dynamic_form.form_type
                },
                'status': submission.status,
                'is_coordinator': submission.is_coordinator,
                'section': submission.section,
                'submission_date': submission.submission_date.isoformat() if submission.submission_date else None,
                'answer_count': answer_count
            })
        
        return JsonResponse(submissions_list, safe=False)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)

# Faculty Course Outline Structure API
@login_required
@require_http_methods(["GET"])
def api_faculty_course_outline_structure(request):
    """Outline starter structure: faculty coordinators or CRC (any course)."""
    if request.user.role not in (User.ROLE_FACULTY, User.ROLE_CRC_MEMBER):
        return JsonResponse({'error': 'Access denied'}, status=403)
    
    try:
        course_id = request.GET.get('course_id')
        
        if not course_id:
            return JsonResponse({'error': 'course_id is required'}, status=400)
        
        if request.user.role == User.ROLE_FACULTY:
            try:
                CourseFaculty.objects.get(
                    faculty=request.user,
                    course_id=course_id,
                    is_coordinator=True
                )
            except CourseFaculty.DoesNotExist:
                return JsonResponse({
                    'error': 'Only course coordinators can create course outlines',
                    'has_access': False
                }, status=403)
        
        course = Course.objects.get(id=course_id)
        
        # Check for existing outline
        existing_outline = CourseOutline.objects.filter(
            course_id=course_id,
            faculty=request.user
        ).order_by('-version').first()
        
        # Default course outline structure
        default_structure = {
            "course_info": {
                "course_name": course.title,
                "course_code": course.code,
                "career_degree": "Undergraduate",
                "obe_enabled": True,
                "academic_term": "Jul 2025",
                "credits": course.credits,
                "department": course.department.name if course.department else ""
            },
            "sections": [
                {
                    "id": "course_information",
                    "title": "Course Information",
                    "type": "table",
                    "headers": ["Item", "Details"],
                    "rows": [
                        {"item": "Course Title", "details": course.title},
                        {"item": "Course Code", "details": course.code},
                        {"item": "Credit Hours", "details": str(course.credits)},
                        {"item": "Prerequisites", "details": ""},
                        {"item": "Corequisites", "details": ""}
                    ]
                },
                {
                    "id": "additional_info",
                    "title": "Additional Information",
                    "type": "text",
                    "content": "Enter additional course information here..."
                },
                {
                    "id": "calendar_activities",
                    "title": "Calendar of Activities",
                    "type": "table",
                    "headers": ["Week", "Topic", "Learning Outcomes", "Readings", "Assessment"],
                    "rows": []
                },
                {
                    "id": "course_books",
                    "title": "Course Books",
                    "type": "table",
                    "headers": ["Title", "Author", "Publisher", "Year", "ISBN"],
                    "rows": []
                },
                {
                    "id": "web_resources",
                    "title": "Web Resources",
                    "type": "list",
                    "items": []
                },
                {
                    "id": "assessment",
                    "title": "Assessment and Grading",
                    "type": "table",
                    "headers": ["Assessment Type", "Weight", "Due Date", "Description"],
                    "rows": []
                },
                {
                    "id": "learning_outcomes",
                    "title": "Course Learning Outcomes (CLOs)",
                    "type": "table",
                    "headers": ["CLO", "Description", "Bloom's Level", "Assessment Methods"],
                    "rows": []
                },
                {
                    "id": "grading_policy",
                    "title": "Grading Policy",
                    "type": "table",
                    "headers": ["Grade", "Range", "Points"],
                    "rows": [
                        {"grade": "A", "range": "90-100", "points": "4.0"},
                        {"grade": "B", "range": "80-89", "points": "3.0"},
                        {"grade": "C", "range": "70-79", "points": "2.0"},
                        {"grade": "D", "range": "60-69", "points": "1.0"},
                        {"grade": "F", "range": "Below 60", "points": "0.0"}
                    ]
                }
            ]
        }
        
        # If there's an existing outline, use its content
        if existing_outline and existing_outline.content:
            try:
                # Try to parse existing content as JSON
                existing_content = json.loads(existing_outline.content) if isinstance(existing_outline.content, str) else existing_outline.content
                
                # Merge with default structure
                if 'course_info' in existing_content:
                    default_structure['course_info'].update(existing_content['course_info'])
                
                if 'sections' in existing_content:
                    # Update sections from existing content
                    for existing_section in existing_content['sections']:
                        section_id = existing_section.get('id')
                        if section_id:
                            # Find and update the corresponding section in default structure
                            for i, default_section in enumerate(default_structure['sections']):
                                if default_section['id'] == section_id:
                                    default_structure['sections'][i].update(existing_section)
                                    break
            except:
                # If content is not valid JSON, use default structure
                pass
        
        return JsonResponse({
            'has_access': True,
            'template_available': False,
            'message': 'Course outlines are created directly',
            'structure': default_structure,
            'existing_outline': {
                'id': existing_outline.id if existing_outline else None,
                'version': existing_outline.version if existing_outline else 0,
                'title': existing_outline.title if existing_outline else '',
                'status': existing_outline.status if existing_outline else 'draft',
                'content': existing_outline.content if existing_outline else None
            } if existing_outline else None
        })
    except Course.DoesNotExist:
        return JsonResponse({'error': 'Course not found'}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)


@login_required
@require_http_methods(["GET"])
def api_faculty_form_availability(request):
    """Check which forms are available for faculty for each course"""
    if request.user.role != User.ROLE_FACULTY:
        return JsonResponse({'error': 'Access denied'}, status=403)
    
    try:
        course_id = request.GET.get('course_id')
        
        # Check GLOBAL form availability
        ccr_active = DynamicForm.objects.filter(
            status=DynamicForm.STATUS_ACTIVE,
            form_type='ccr'
        ).exists()
        
        crr_active = DynamicForm.objects.filter(
            status=DynamicForm.STATUS_ACTIVE,
            form_type='crr'
        ).exists()
        
        if not course_id:
            # Return all courses with form availability
            course_assignments = CourseFaculty.objects.filter(
                faculty=request.user
            ).select_related('course', 'course__department')
            
            courses_data = []
            for assignment in course_assignments:
                courses_data.append({
                    'course_id': assignment.course.id,
                    'course_code': assignment.course.code,
                    'course_title': assignment.course.title,
                    'is_coordinator': assignment.is_coordinator,
                    'section': assignment.section_display(),
                    'forms_available': {
                        'ccr': ccr_active and assignment.is_coordinator,
                        'crr': crr_active
                    }
                })
            
            return JsonResponse({
                'global_availability': {
                    'ccr': ccr_active,
                    'crr': crr_active
                },
                'courses': courses_data
            }, safe=False)
        else:
            # Check specific course
            try:
                assignment = CourseFaculty.objects.get(
                    faculty=request.user,
                    course_id=course_id
                )
                
                return JsonResponse({
                    'course_id': course_id,
                    'is_coordinator': assignment.is_coordinator,
                    'forms_available': {
                        'ccr': ccr_active and assignment.is_coordinator,
                        'crr': crr_active
                    }
                })
                
            except CourseFaculty.DoesNotExist:
                return JsonResponse({'error': 'Course not assigned to you'}, status=400)
                
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)


@login_required
@user_passes_test(is_admin_or_crc)
@require_http_methods(["GET"])
def api_get_all_outlines(request):
    """Get ALL outlines - both approved and submitted"""
    try:
        outlines = CourseOutline.objects.filter(
            status__in=['submitted', 'approved']
        ).select_related('course', 'faculty').order_by('-created_at')
        
        outlines_list = []
        for outline in outlines:
            outlines_list.append({
                'id': outline.id,
                'course_code': outline.course.code,
                'course_title': outline.course.title,
                'faculty': outline.faculty.username,
                'version': outline.version,
                'title': outline.title,
                'status': outline.status,
                'created_at': outline.created_at.strftime('%Y-%m-%d') if outline.created_at else None
            })
        
        return JsonResponse(outlines_list, safe=False)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)


@login_required
@user_passes_test(is_admin)
@csrf_exempt
@require_http_methods(["PUT"])
def api_department_update(request, department_id):
    try:
        data = json.loads(request.body)
        department = Department.objects.get(id=department_id)
        department.name = data.get('name', department.name)
        department.code = data.get('code', department.code)
        department.description = data.get('description', department.description)
        department.save()
        return JsonResponse({
            'id': department.id,
            'name': department.name,
            'code': department.code,
            'description': department.description
        })
    except Department.DoesNotExist:
        return JsonResponse({'error': 'Department not found'}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)


# new apii

# Analysis Endpoints
@login_required
@user_passes_test(is_admin_or_crc)
@require_http_methods(["GET"])
def api_analysis_submissions_over_time(request):
    """Get form submissions over time (weekly periods).

    Always returns a contiguous timeline of the configured number of weeks so
    that the chart shows the full window even when most weeks have no data.
    Optional query params:
        - ``course_id``: restrict the chart to a single course.
        - ``weeks``: number of weeks of history to include (default 12, capped
          between 4 and 52).
    """
    try:
        selected_course_id = request.GET.get('course_id')
        try:
            weeks_window = int(request.GET.get('weeks', 12))
        except (TypeError, ValueError):
            weeks_window = 12
        weeks_window = max(4, min(weeks_window, 52))

        # Anchor the timeline on the Monday of the current week so each bucket
        # represents a full ISO week (Mon-Sun). Use timezone-aware "now" when
        # USE_TZ is enabled to avoid Django's naive-datetime warning.
        if getattr(settings, 'USE_TZ', False):
            from django.utils import timezone
            today = timezone.localtime(timezone.now())
        else:
            today = datetime.now()
        current_week_start = (today - timedelta(days=today.weekday())).replace(
            hour=0, minute=0, second=0, microsecond=0
        )
        earliest_week_start = current_week_start - timedelta(weeks=weeks_window - 1)

        # Pre-seed every week in the window with zeroes so the chart always
        # renders a continuous x-axis instead of a single floating point.
        ordered_week_keys = []
        week_data = {}
        for week_offset in range(weeks_window):
            week_start = earliest_week_start + timedelta(weeks=week_offset)
            week_key = week_start.strftime('%Y-%m-%d')
            ordered_week_keys.append(week_key)
            week_data[week_key] = {'ccr': 0, 'crr': 0, 'total': 0}

        submissions = DynamicFormSubmission.objects.filter(
            submission_date__gte=earliest_week_start,
            dynamic_form__form_type__in=['ccr', 'crr'],
        ).select_related('dynamic_form').order_by('submission_date')

        if selected_course_id:
            submissions = submissions.filter(course_id=selected_course_id)

        for submission in submissions:
            submission_week_start = submission.submission_date - timedelta(
                days=submission.submission_date.weekday()
            )
            submission_week_start = submission_week_start.replace(
                hour=0, minute=0, second=0, microsecond=0
            )
            week_key = submission_week_start.strftime('%Y-%m-%d')
            # Submissions slightly older than the window can be safely ignored.
            bucket = week_data.get(week_key)
            if bucket is None:
                continue
            bucket['total'] += 1
            if submission.dynamic_form.form_type == 'ccr':
                bucket['ccr'] += 1
            else:
                bucket['crr'] += 1

        result = {
            'weeks': ordered_week_keys,
            'ccr_data': [week_data[week_key]['ccr'] for week_key in ordered_week_keys],
            'crr_data': [week_data[week_key]['crr'] for week_key in ordered_week_keys],
            'total_data': [week_data[week_key]['total'] for week_key in ordered_week_keys],
        }

        return JsonResponse(result)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)

@login_required
@user_passes_test(is_admin_or_crc)
@require_http_methods(["GET"])
def api_analysis_form_status(request):
    """Get form status distribution.

    Optional query param ``course_id`` restricts the chart to a single course.
    """
    try:
        selected_course_id = request.GET.get('course_id')

        # Get status counts for universal forms
        status_queryset = DynamicFormSubmission.objects.filter(
            dynamic_form__form_type__in=['ccr', 'crr']
        )

        if selected_course_id:
            status_queryset = status_queryset.filter(course_id=selected_course_id)

        status_counts = status_queryset.values('status').annotate(count=Count('id'))
        
        # Get totals
        total_submissions = sum(item['count'] for item in status_counts)
        
        # Format for pie chart
        result = {
            'labels': [],
            'data': [],
            'colors': []
        }
        
        status_colors = {
            'submitted': '#FFC107',  # Yellow
            'approved': '#4CAF50',    # Green
            'draft': '#9E9E9E',       # Grey
            'revision_requested': '#F44336',  # Red
        }
        
        for item in status_counts:
            result['labels'].append(item['status'].title())
            result['data'].append(item['count'])
            result['colors'].append(status_colors.get(item['status'], '#2196F3'))
        
        return JsonResponse(result)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)


# Update the CLO achievement analysis function
@login_required
@user_passes_test(is_admin_or_crc)
@require_http_methods(["GET"])
def api_analysis_clo_achievement(request):
    """Get CLO achievement rates by analyzing actual form answers AND course outlines.

    Optional query param ``course_id`` restricts the analysis to a single course.
    """
    try:
        selected_course_id = request.GET.get('course_id')

        # Initialize CLO data for 4 CLOs
        clo_data = {
            'CLO1': {'scores': [], 'count': 0, 'achieved_count': 0, 'sources': []},
            'CLO2': {'scores': [], 'count': 0, 'achieved_count': 0, 'sources': []},
            'CLO3': {'scores': [], 'count': 0, 'achieved_count': 0, 'sources': []},
            'CLO4': {'scores': [], 'count': 0, 'achieved_count': 0, 'sources': []},
        }

        # PART 1: Analyze form submissions (CCR/CRR forms)
        submissions = DynamicFormSubmission.objects.filter(
            dynamic_form__form_type__in=['ccr', 'crr'],
            status__in=['submitted', 'approved']  # Only analyze submitted/approved forms
        ).prefetch_related('answers__question')

        if selected_course_id:
            submissions = submissions.filter(course_id=selected_course_id)
        
        form_count = 0
        for submission in submissions:
            answers = submission.answers.all()
            
            for answer in answers:
                question_text = answer.question.question_text.lower()
                
                # Method 1: Check for CLO percentage dictionary answers
                if isinstance(answer.answer_data, dict):
                    for clo_num in [1, 2, 3, 4]:
                        clo_key = f'CLO{clo_num}'
                        # Try different key formats
                        for key in [f'clo{clo_num}', f'clo_{clo_num}', f'CLO{clo_num}', f'CLO_{clo_num}']:
                            if key in answer.answer_data:
                                score = answer.answer_data[key]
                                if isinstance(score, (int, float)):
                                    clo_data[clo_key]['scores'].append(score)
                                    clo_data[clo_key]['count'] += 1
                                    clo_data[clo_key]['sources'].append(f"Form: {submission.course.code}")
                                    if score >= 70:
                                        clo_data[clo_key]['achieved_count'] += 1
                                elif isinstance(score, str):
                                    try:
                                        # Handle percentage strings
                                        if '%' in score:
                                            score_val = float(score.replace('%', '').strip())
                                        else:
                                            score_val = float(score)
                                        
                                        clo_data[clo_key]['scores'].append(score_val)
                                        clo_data[clo_key]['count'] += 1
                                        clo_data[clo_key]['sources'].append(f"Form: {submission.course.code}")
                                        if score_val >= 70:
                                            clo_data[clo_key]['achieved_count'] += 1
                                    except:
                                        pass
                
                # Method 2: Check question text for CLO references
                for clo_num in [1, 2, 3, 4]:
                    clo_patterns = [
                        f'clo{clo_num}',
                        f'clo {clo_num}',
                        f'course learning outcome {clo_num}',
                        f'learning outcome {clo_num}',
                        f'clo-{clo_num}',
                        f'clo_{clo_num}',
                    ]
                    
                    if any(pattern in question_text for pattern in clo_patterns):
                        clo_key = f'CLO{clo_num}'
                        
                        # Extract score from answer
                        score = extract_clo_score_from_answer(answer)
                        if score is not None:
                            clo_data[clo_key]['scores'].append(score)
                            clo_data[clo_key]['count'] += 1
                            clo_data[clo_key]['sources'].append(f"Form: {submission.course.code}")
                            
                            # Check if achieved (score >= 70%)
                            if score >= 70:
                                clo_data[clo_key]['achieved_count'] += 1
        
        # PART 2: Analyze course outlines for CLO data
        # Get approved course outlines
        approved_outlines = CourseOutline.objects.filter(
            status='approved',
            is_current=True
        )

        if selected_course_id:
            approved_outlines = approved_outlines.filter(course_id=selected_course_id)
        
        outline_count = 0
        for outline in approved_outlines:
            try:
                # Parse outline content (assuming JSON structure)
                if outline.content:
                    content = outline.content
                    if isinstance(content, str):
                        try:
                            content_data = json.loads(content)
                        except:
                            content_data = None
                    else:
                        content_data = content
                    
                    if content_data and isinstance(content_data, dict):
                        # Look for CLO data in outline structure
                        if 'sections' in content_data:
                            for section in content_data['sections']:
                                if isinstance(section, dict):
                                    # Check for CLO tables or sections
                                    section_title = section.get('title', '').lower()
                                    section_id = section.get('id', '').lower()
                                    
                                    # Check if this is a CLO-related section
                                    if any(clo_term in section_title or clo_term in section_id 
                                           for clo_term in ['clo', 'course learning outcome', 'learning outcome']):
                                        
                                        # Extract rows from table sections
                                        if section.get('type') == 'table' and 'rows' in section:
                                            for row in section['rows']:
                                                if isinstance(row, dict):
                                                    # Look for CLO data in row
                                                    row_text = str(row).lower()
                                                    for clo_num in [1, 2, 3, 4]:
                                                        if f'clo{clo_num}' in row_text or f'clo {clo_num}' in row_text:
                                                            clo_key = f'CLO{clo_num}'
                                                            
                                                            # Try to extract percentage from row
                                                            row_str = str(row)
                                                            import re
                                                            
                                                            # Look for percentages in row
                                                            percent_pattern = r'(\d+\.?\d*)%'
                                                            percentages = re.findall(percent_pattern, row_str)
                                                            if percentages:
                                                                try:
                                                                    score = float(percentages[0])
                                                                    clo_data[clo_key]['scores'].append(score)
                                                                    clo_data[clo_key]['count'] += 1
                                                                    clo_data[clo_key]['sources'].append(f"Outline: {outline.course.code}")
                                                                    if score >= 70:
                                                                        clo_data[clo_key]['achieved_count'] += 1
                                                                except:
                                                                    pass
            except Exception as e:
                print(f"Error parsing outline {outline.id}: {str(e)}")
                continue
        
        # Calculate achievement rates and prepare result
        result = {
            'clos': ['CLO1', 'CLO2', 'CLO3', 'CLO4'],
            'achievement_rates': [],
            'average_scores': [],
            'total_responses': [],
            'achieved_counts': [],
            'details': [],
            'sources_summary': {
                'form_submissions': form_count,
                'course_outlines': outline_count,
                'total_sources': form_count + outline_count
            }
        }
        
        for clo in result['clos']:
            scores = clo_data[clo]['scores']
            count = clo_data[clo]['count']
            achieved = clo_data[clo]['achieved_count']
            
            if count > 0:
                avg_score = sum(scores) / len(scores)
                achievement_rate = (achieved / count) * 100
            else:
                avg_score = 0
                achievement_rate = 0
            
            result['average_scores'].append(round(avg_score, 1))
            result['achievement_rates'].append(round(achievement_rate, 1))
            result['total_responses'].append(count)
            result['achieved_counts'].append(achieved)
            
            # Add detailed breakdown
            result['details'].append({
                'clo': clo,
                'average_score': round(avg_score, 1),
                'achievement_rate': round(achievement_rate, 1),
                'total_responses': count,
                'achieved_responses': achieved,
                'score_distribution': get_score_distribution(scores) if scores else {},
                'data_sources': list(set(clo_data[clo]['sources']))[:5]  # Unique sources, limit to 5
            })
        
        return JsonResponse(result)
    except Exception as e:
        print(f"Error in CLO achievement analysis: {str(e)}")
        import traceback
        traceback.print_exc()
        return JsonResponse({'error': str(e)}, status=400)


def _parse_clo_percentage_dict(answer_data):
    """Extract clo1..clo4 numeric scores from a clo_percentage answer_data dict."""
    scores = {}
    if not isinstance(answer_data, dict):
        return scores
    for clo_num in [1, 2, 3, 4]:
        for key in (
            f'clo{clo_num}',
            f'clo_{clo_num}',
            f'CLO{clo_num}',
            f'CLO_{clo_num}',
        ):
            if key not in answer_data:
                continue
            raw_value = answer_data[key]
            try:
                if isinstance(raw_value, str):
                    raw_value = raw_value.replace('%', '').strip()
                score_value = float(raw_value)
                scores[clo_num] = score_value
                break
            except (TypeError, ValueError):
                continue
    return scores


def _short_ccr_criterion_label(question_text, fallback_index):
    """Map CCR Q6-Q9 question text to a short chart legend label."""
    text = (question_text or '').lower()
    if 'student-centered' in text or 'student centered' in text:
        return 'Student-centered'
    if 'measurable' in text:
        return 'Measurable'
    if 'achievable' in text or 'realistic' in text:
        return 'Achievable'
    if 'bloom' in text or 'action verb' in text or "dave" in text:
        return "Bloom's verb"
    cleaned = (question_text or '').strip()
    if cleaned:
        return cleaned[:42] + ('…' if len(cleaned) > 42 else '')
    return f'Criterion {fallback_index}'


def _clo_number_from_mapping_question(question):
    """Return 1-4 for CCR Q10-Q13 style questions labeled CLO-1 … CLO-4."""
    text = (question.question_text or '').strip().lower()
    match = re.search(r'clo[-\s]?([1-4])\b', text)
    if match:
        return int(match.group(1))
    return None


def _normalize_mapping_option(option_value):
    """Normalize checkbox option labels for Domain / Level / GA Mapping."""
    text = str(option_value or '').strip().lower()
    if 'domain' in text:
        return 'Learning Domain'
    if text == 'level' or text.startswith('level'):
        return 'Level'
    if 'ga' in text or 'graduate' in text:
        return 'GA Mapping'
    return None


@login_required
@user_passes_test(is_admin_or_crc)
@require_http_methods(["GET"])
def api_analysis_ccr_clo_graphs(request):
    """
    Build Analysis Dashboard graphs from real CCR form answers:

    - Q6–Q9 (clo_percentage): CLO achievement / quality criteria
    - Q10–Q13 (checkbox CLO-1..CLO-4): Learning Domain / Level / GA Mapping
    - Pie: healthy vs problematic CLO instances (avg Q6–Q9 score < 70)
    """
    try:
        selected_course_id = request.GET.get('course_id')
        achievement_threshold = 70.0

        submissions = (
            DynamicFormSubmission.objects.filter(
                dynamic_form__form_type='ccr',
                status__in=['submitted', 'approved'],
            )
            .select_related('course', 'dynamic_form', 'faculty')
            .prefetch_related('answers__question')
        )
        if selected_course_id:
            submissions = submissions.filter(course_id=selected_course_id)

        submissions_list = list(submissions)
        submission_count = len(submissions_list)

        # criterion_key -> {clo_num -> [scores]}
        criterion_scores = {}
        criterion_labels = {}
        # clo_num -> list of per-submission averages (across Q6-Q9)
        clo_submission_averages = {1: [], 2: [], 3: [], 4: []}
        # mapping: clo_num -> dimension -> count
        mapping_dimensions = ['Learning Domain', 'Level', 'GA Mapping']
        mapping_counts = {
            clo_num: {dimension: 0 for dimension in mapping_dimensions}
            for clo_num in [1, 2, 3, 4]
        }
        mapping_response_counts = {1: 0, 2: 0, 3: 0, 4: 0}

        healthy_clo_instances = 0
        problematic_clo_instances = 0
        clo_health_details = {
            clo_num: {'healthy': 0, 'problematic': 0, 'averages': []}
            for clo_num in [1, 2, 3, 4]
        }

        for submission in submissions_list:
            quality_answers = []
            for answer in submission.answers.all():
                question = answer.question
                if question.question_type == 'clo_percentage':
                    quality_answers.append(answer)
                    continue

                clo_num = _clo_number_from_mapping_question(question)
                if clo_num is None or question.question_type != 'checkbox':
                    continue

                mapping_response_counts[clo_num] += 1
                selected_options = []
                if isinstance(answer.answer_data, list):
                    selected_options = answer.answer_data
                elif answer.answer_text:
                    selected_options = [
                        part.strip()
                        for part in answer.answer_text.split(',')
                        if part.strip()
                    ]

                for option in selected_options:
                    dimension = _normalize_mapping_option(option)
                    if dimension:
                        mapping_counts[clo_num][dimension] += 1

            # Sort Q6-Q9 by question order so criteria stay stable.
            quality_answers.sort(key=lambda item: (item.question.order, item.question.id))

            per_clo_scores_for_submission = {1: [], 2: [], 3: [], 4: []}
            for criterion_index, answer in enumerate(quality_answers, start=1):
                question = answer.question
                criterion_key = f'q_{question.id}'
                if criterion_key not in criterion_labels:
                    criterion_labels[criterion_key] = _short_ccr_criterion_label(
                        question.question_text,
                        criterion_index,
                    )
                    criterion_scores[criterion_key] = {1: [], 2: [], 3: [], 4: []}

                parsed_scores = _parse_clo_percentage_dict(answer.answer_data)
                for clo_num, score_value in parsed_scores.items():
                    criterion_scores[criterion_key][clo_num].append(score_value)
                    per_clo_scores_for_submission[clo_num].append(score_value)

            for clo_num, score_list in per_clo_scores_for_submission.items():
                if not score_list:
                    continue
                average_score = sum(score_list) / len(score_list)
                clo_submission_averages[clo_num].append(average_score)
                clo_health_details[clo_num]['averages'].append(average_score)
                if average_score >= achievement_threshold:
                    healthy_clo_instances += 1
                    clo_health_details[clo_num]['healthy'] += 1
                else:
                    problematic_clo_instances += 1
                    clo_health_details[clo_num]['problematic'] += 1

        # Preserve criterion order by first-seen insertion (already sorted per submission).
        ordered_criterion_keys = list(criterion_scores.keys())
        criteria_labels = [
            criterion_labels[key] for key in ordered_criterion_keys
        ]

        criteria_averages = []
        for criterion_key in ordered_criterion_keys:
            row = []
            for clo_num in [1, 2, 3, 4]:
                values = criterion_scores[criterion_key][clo_num]
                row.append(round(sum(values) / len(values), 1) if values else 0)
            criteria_averages.append(row)

        clo_labels = ['CLO1', 'CLO2', 'CLO3', 'CLO4']
        average_scores = []
        achievement_rates = []
        total_score_counts = []
        achieved_counts = []

        for clo_num in [1, 2, 3, 4]:
            all_scores = []
            for criterion_key in ordered_criterion_keys:
                all_scores.extend(criterion_scores[criterion_key][clo_num])

            total_score_counts.append(len(all_scores))
            if all_scores:
                average_scores.append(round(sum(all_scores) / len(all_scores), 1))
                achieved = sum(1 for score in all_scores if score >= achievement_threshold)
                achieved_counts.append(achieved)
                achievement_rates.append(round((achieved / len(all_scores)) * 100, 1))
            else:
                average_scores.append(0)
                achieved_counts.append(0)
                achievement_rates.append(0)

        mapping_count_series = {
            dimension: [
                mapping_counts[clo_num][dimension] for clo_num in [1, 2, 3, 4]
            ]
            for dimension in mapping_dimensions
        }
        mapping_coverage_rates = {
            dimension: [
                round(
                    (mapping_counts[clo_num][dimension] / mapping_response_counts[clo_num]) * 100,
                    1,
                )
                if mapping_response_counts[clo_num]
                else 0
                for clo_num in [1, 2, 3, 4]
            ]
            for dimension in mapping_dimensions
        }

        problematic_details = []
        for clo_num in [1, 2, 3, 4]:
            averages = clo_health_details[clo_num]['averages']
            overall_average = (
                round(sum(averages) / len(averages), 1) if averages else 0
            )
            problematic_details.append({
                'clo': f'CLO{clo_num}',
                'average_score': overall_average,
                'healthy_instances': clo_health_details[clo_num]['healthy'],
                'problematic_instances': clo_health_details[clo_num]['problematic'],
                'status': (
                    'problematic'
                    if overall_average and overall_average < achievement_threshold
                    else ('healthy' if averages else 'no_data')
                ),
            })

        return JsonResponse({
            'submission_count': submission_count,
            'threshold': achievement_threshold,
            'source': {
                'form_type': 'ccr',
                'achievement_questions': 'Q6-Q9 (clo_percentage criteria)',
                'mapping_questions': 'Q10-Q13 (CLO Domain / Level / GA Mapping)',
            },
            'achievement': {
                'clos': clo_labels,
                'criteria_labels': criteria_labels,
                'criteria_averages': criteria_averages,
                'average_scores': average_scores,
                'achievement_rates': achievement_rates,
                'total_responses': total_score_counts,
                'achieved_counts': achieved_counts,
            },
            'mapping': {
                'clos': clo_labels,
                'dimensions': mapping_dimensions,
                'counts': mapping_count_series,
                'coverage_rates': mapping_coverage_rates,
                'response_counts': [
                    mapping_response_counts[clo_num] for clo_num in [1, 2, 3, 4]
                ],
            },
            'problematic': {
                'labels': ['Healthy CLOs', 'Problematic CLOs'],
                'data': [healthy_clo_instances, problematic_clo_instances],
                'colors': ['#10b981', '#ef4444'],
                'threshold': achievement_threshold,
                'details': problematic_details,
                'total_instances': healthy_clo_instances + problematic_clo_instances,
            },
        })
    except Exception as e:
        print(f"Error in CCR CLO graphs analysis: {str(e)}")
        import traceback
        traceback.print_exc()
        return JsonResponse({'error': str(e)}, status=400)


def extract_clo_score_from_answer(answer):
    """Extract a numerical score (0-100) from an answer"""
    try:
        # If the question is CLO percentage type, check answer_data first
        if answer.question.question_type == 'clo_percentage' and answer.answer_data:
            if isinstance(answer.answer_data, dict):
                # For CLO percentage type, average all CLO values
                clo_values = []
                for clo_num in [1, 2, 3, 4]:
                    for key in [f'clo{clo_num}', f'clo_{clo_num}', f'CLO{clo_num}']:
                        if key in answer.answer_data:
                            value = answer.answer_data[key]
                            if isinstance(value, (int, float)):
                                clo_values.append(value)
                            elif isinstance(value, str):
                                try:
                                    if '%' in value:
                                        clo_values.append(float(value.replace('%', '').strip()))
                                    else:
                                        clo_values.append(float(value))
                                except:
                                    pass
                if clo_values:
                    return sum(clo_values) / len(clo_values)
                
                # Look for general score fields
                for key in ['score', 'percentage', 'rating', 'value', 'achievement']:
                    if key in answer.answer_data:
                        try:
                            value = answer.answer_data[key]
                            if isinstance(value, (int, float)):
                                return float(value)
                            elif isinstance(value, str):
                                if '%' in value:
                                    return float(value.replace('%', '').strip())
                                else:
                                    return float(value)
                        except:
                            continue
            
            # If it's a list, try to extract numeric values
            elif isinstance(answer.answer_data, list):
                numeric_values = []
                for item in answer.answer_data:
                    if isinstance(item, (int, float)):
                        numeric_values.append(item)
                    elif isinstance(item, str):
                        try:
                            # Try to convert to float
                            val = float(item)
                            numeric_values.append(val)
                        except:
                            # Try to extract number from string
                            import re
                            numbers = re.findall(r'\d+\.?\d*', item)
                            if numbers:
                                try:
                                    numeric_values.append(float(numbers[0]))
                                except:
                                    pass
                
                if numeric_values:
                    return sum(numeric_values) / len(numeric_values)
        
        # Case 4: Try to extract number from text
        if answer.answer_text:
            import re
            # Look for percentages first
            percent_pattern = r'(\d+\.?\d*)%'
            percent_matches = re.findall(percent_pattern, answer.answer_text)
            if percent_matches:
                try:
                    return float(percent_matches[0])
                except:
                    pass
            
            # Look for any numbers
            numbers = re.findall(r'\d+\.?\d*', answer.answer_text)
            if numbers:
                try:
                    # Take the first number that looks like a percentage (0-100)
                    for num in numbers:
                        val = float(num)
                        if 0 <= val <= 100:
                            return val
                    # If no number in 0-100 range, take the first one
                    return float(numbers[0])
                except:
                    pass
        
        # Case 5: Map textual answers to scores
        if answer.answer_text:
            text = answer.answer_text.lower().strip()
            
            # Check for percentage phrases
            if '100%' in text or 'hundred percent' in text:
                return 100
            elif '90%' in text or 'ninety percent' in text:
                return 90
            elif '80%' in text or 'eighty percent' in text:
                return 80
            elif '70%' in text or 'seventy percent' in text:
                return 70
            elif '60%' in text or 'sixty percent' in text:
                return 60
            elif '50%' in text or 'fifty percent' in text:
                return 50
            elif '40%' in text or 'forty percent' in text:
                return 40
            elif '30%' in text or 'thirty percent' in text:
                return 30
            elif '20%' in text or 'twenty percent' in text:
                return 20
            elif '10%' in text or 'ten percent' in text:
                return 10
            elif '0%' in text or 'zero percent' in text:
                return 0
            
            # Map textual ratings to scores
            mapping = {
                'excellent': 90, 'outstanding': 95, 'exceptional': 95,
                'very good': 85, 'very good': 85,
                'good': 80, 'well achieved': 80,
                'satisfactory': 75, 'adequate': 75, 'moderate': 75,
                'average': 70, 'medium': 70,
                'fair': 65, 'acceptable': 65,
                'poor': 60, 'below average': 60,
                'very poor': 50, 'inadequate': 50,
                'not achieved': 40, 'failed': 40, 'unsatisfactory': 40,
                'yes': 80, 'achieved': 80,
                'no': 40, 'not achieved': 40,
                'partially': 60, 'partially achieved': 60,
                'fully': 90, 'fully achieved': 90,
                'substantially': 85, 'mostly': 80,
            }
            
            for key, value in mapping.items():
                if key in text:
                    return value
            
            # Check for Likert scale
            likert_map = {
                '5': 90, '5/5': 90, '5 out of 5': 90, 'five': 90,
                '4': 75, '4/5': 75, '4 out of 5': 75, 'four': 75,
                '3': 60, '3/5': 60, '3 out of 5': 60, 'three': 60,
                '2': 45, '2/5': 45, '2 out of 5': 45, 'two': 45,
                '1': 30, '1/5': 30, '1 out of 5': 30, 'one': 30,
            }
            
            for key, value in likert_map.items():
                if key in text:
                    return value
        
        return None
    except:
        return None

def get_score_distribution(scores):
    """Get distribution of scores in categories"""
    if not scores:
        return {}
    
    distribution = {
        'Excellent (90-100)': 0,
        'Good (80-89)': 0,
        'Satisfactory (70-79)': 0,
        'Needs Improvement (60-69)': 0,
        'Poor (Below 60)': 0,
    }
    
    for score in scores:
        if score >= 90:
            distribution['Excellent (90-100)'] += 1
        elif score >= 80:
            distribution['Good (80-89)'] += 1
        elif score >= 70:
            distribution['Satisfactory (70-79)'] += 1
        elif score >= 60:
            distribution['Needs Improvement (60-69)'] += 1
        else:
            distribution['Poor (Below 60)'] += 1
    
    # Convert to percentages
    total = len(scores)
    if total > 0:
        for key in distribution:
            distribution[key] = round((distribution[key] / total) * 100, 1)
    
    return distribution

# Add this new endpoint for detailed CLO analysis
@login_required
@user_passes_test(is_admin_or_crc)
@require_http_methods(["GET"])
def api_analysis_detailed_clo(request, clo_number):
    """Get detailed analysis for a specific CLO"""
    try:
        if clo_number not in [1, 2, 3, 4]:
            return JsonResponse({'error': 'Invalid CLO number. Must be 1-4'}, status=400)
        
        clo_key = f'CLO{clo_number}'
        
        # Get all answers related to this CLO
        answers = FormAnswer.objects.filter(
            question__question_text__icontains=f'clo{clo_number}',
            submission__dynamic_form__form_type__in=['ccr', 'crr']
        ).select_related('submission', 'question', 'submission__course', 'submission__faculty')
        
        # Prepare detailed data
        detailed_data = []
        courses_data = {}
        faculty_data = {}
        
        for answer in answers:
            score = extract_clo_score_from_answer(answer)
            if score is None:
                continue
            
            # Course analysis
            course_code = answer.submission.course.code
            if course_code not in courses_data:
                courses_data[course_code] = {'scores': [], 'count': 0}
            courses_data[course_code]['scores'].append(score)
            courses_data[course_code]['count'] += 1
            
            # Faculty analysis
            faculty_name = answer.submission.faculty.username
            if faculty_name not in faculty_data:
                faculty_data[faculty_name] = {'scores': [], 'count': 0}
            faculty_data[faculty_name]['scores'].append(score)
            faculty_data[faculty_name]['count'] += 1
            
            # Add to detailed list
            detailed_data.append({
                'id': answer.id,
                'course_code': course_code,
                'course_title': answer.submission.course.title,
                'faculty': faculty_name,
                'question': answer.question.question_text[:100] + ('...' if len(answer.question.question_text) > 100 else ''),
                'answer': answer.answer_text[:200] if answer.answer_text else str(answer.answer_data)[:200],
                'score': score,
                'achieved': score >= 70,
                'submission_date': answer.submission.submission_date.isoformat() if answer.submission.submission_date else None,
                'form_type': answer.submission.dynamic_form.form_type,
            })
        
        # Calculate course averages
        course_analysis = []
        for course_code, data in courses_data.items():
            avg_score = sum(data['scores']) / len(data['scores'])
            achievement_rate = sum(1 for s in data['scores'] if s >= 70) / len(data['scores']) * 100
            course_analysis.append({
                'course_code': course_code,
                'average_score': round(avg_score, 1),
                'achievement_rate': round(achievement_rate, 1),
                'response_count': data['count']
            })
        
        # Calculate faculty averages
        faculty_analysis = []
        for faculty_name, data in faculty_data.items():
            avg_score = sum(data['scores']) / len(data['scores'])
            achievement_rate = sum(1 for s in data['scores'] if s >= 70) / len(data['scores']) * 100
            faculty_analysis.append({
                'faculty': faculty_name,
                'average_score': round(avg_score, 1),
                'achievement_rate': round(achievement_rate, 1),
                'response_count': data['count']
            })
        
        # Sort analyses
        course_analysis.sort(key=lambda x: x['achievement_rate'], reverse=True)
        faculty_analysis.sort(key=lambda x: x['achievement_rate'], reverse=True)
        
        # Calculate overall statistics
        all_scores = []
        for answer in detailed_data:
            all_scores.append(answer['score'])
        
        if all_scores:
            overall_avg = sum(all_scores) / len(all_scores)
            overall_achievement = sum(1 for s in all_scores if s >= 70) / len(all_scores) * 100
        else:
            overall_avg = 0
            overall_achievement = 0
        
        return JsonResponse({
            'clo': clo_key,
            'overall_statistics': {
                'average_score': round(overall_avg, 1),
                'achievement_rate': round(overall_achievement, 1),
                'total_responses': len(all_scores),
                'achieved_responses': sum(1 for s in all_scores if s >= 70),
                'score_range': {
                    'min': min(all_scores) if all_scores else 0,
                    'max': max(all_scores) if all_scores else 0,
                    'median': get_median(all_scores) if all_scores else 0
                }
            },
            'course_analysis': course_analysis,
            'faculty_analysis': faculty_analysis[:10],  # Top 10 faculty
            'detailed_responses': detailed_data[:50],  # Limit to 50 most recent
            'generated_at': datetime.now().isoformat()
        })
        
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)

def get_median(scores):
    """Calculate median of scores"""
    sorted_scores = sorted(scores)
    n = len(sorted_scores)
    if n % 2 == 0:
        return (sorted_scores[n//2 - 1] + sorted_scores[n//2]) / 2
    else:
        return sorted_scores[n//2]

# Add endpoint for trend analysis
@login_required
@user_passes_test(is_admin_or_crc)
@require_http_methods(["GET"])
def api_analysis_clo_trends(request):
    """Get CLO achievement trends over time"""
    try:
        # Get data for last 4 quarters
        end_date = datetime.now()
        start_date = end_date - timedelta(days=365)  # Last year
        
        # Initialize data structure
        quarters = []
        clo_trends = {
            'CLO1': [],
            'CLO2': [],
            'CLO3': [],
            'CLO4': []
        }
        
        # Divide into quarters
        for i in range(4):
            quarter_start = start_date + timedelta(days=i*90)
            quarter_end = start_date + timedelta(days=(i+1)*90)
            quarter_name = f"Q{i+1} {quarter_start.year}"
            quarters.append(quarter_name)
            
            # Get answers for this quarter
            for clo_num in [1, 2, 3, 4]:
                answers = FormAnswer.objects.filter(
                    question__question_text__icontains=f'clo{clo_num}',
                    submission__submission_date__gte=quarter_start,
                    submission__submission_date__lt=quarter_end,
                    submission__dynamic_form__form_type__in=['ccr', 'crr']
                )
                
                scores = []
                for answer in answers:
                    score = extract_clo_score_from_answer(answer)
                    if score is not None:
                        scores.append(score)
                
                if scores:
                    achievement_rate = sum(1 for s in scores if s >= 70) / len(scores) * 100
                else:
                    achievement_rate = 0
                
                clo_trends[f'CLO{clo_num}'].append(round(achievement_rate, 1))
        
        return JsonResponse({
            'quarters': quarters,
            'trends': clo_trends,
            'time_period': f"{start_date.strftime('%b %Y')} - {end_date.strftime('%b %Y')}"
        })
        
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)

@login_required
@user_passes_test(is_admin_or_crc)
@csrf_exempt
@require_http_methods(["POST"])
def api_compare_outlines(request):
    """Compare two course outline versions using an LLM.

    The configured AI provider (via LiteLLM) is used to produce a rich
    qualitative comparison: an executive summary, a similarity score, the
    added / modified / removed sections with natural-language explanations,
    and improvement recommendations.

    When the AI is unavailable (no API key, network error, malformed
    response, ...) the function falls back to the structural diff helpers
    so the UI keeps working.
    """
    try:
        data = json.loads(request.body)
        outline1_id = data.get('outline1_id')
        outline2_id = data.get('outline2_id')

        if not outline1_id or not outline2_id:
            return JsonResponse({'error': 'Both outline IDs are required'}, status=400)

        first_outline = CourseOutline.objects.select_related('course', 'faculty').get(id=outline1_id)
        second_outline = CourseOutline.objects.select_related('course', 'faculty').get(id=outline2_id)

        if first_outline.course_id != second_outline.course_id:
            return JsonResponse({'error': 'Outlines must be from the same course'}, status=400)

        parsed_first_content = _outline_content_for_comparison(first_outline.content)
        parsed_second_content = _outline_content_for_comparison(second_outline.content)

        outline_meta = {
            'outline1': _serialize_outline_meta(first_outline),
            'outline2': _serialize_outline_meta(second_outline),
            'course': {
                'code': first_outline.course.code,
                'title': first_outline.course.title,
            },
            'compared_at': datetime.now().isoformat(),
        }

        ai_config = get_ai_provider_config()
        ai_payload = None
        ai_error_message = None

        if ai_config.get('available'):
            try:
                ai_payload = _generate_ai_outline_comparison(
                    first_outline,
                    second_outline,
                    parsed_first_content,
                    parsed_second_content,
                    ai_config,
                )
            except Exception as ai_exc:
                ai_error_message = str(ai_exc)
                print(f"AI outline comparison failed: {ai_error_message}")
        else:
            ai_error_message = (
                f"AI provider '{ai_config.get('provider', 'unknown')}' is not configured."
            )

        fallback_differences = _build_fallback_differences(
            parsed_first_content,
            parsed_second_content,
        )

        if ai_payload is not None:
            response_body = {
                **outline_meta,
                'comparison_mode': 'ai',
                'ai_provider': ai_config.get('provider'),
                'ai_model': ai_config.get('model'),
                'ai': ai_payload,
                'differences': fallback_differences,
            }
        else:
            response_body = {
                **outline_meta,
                'comparison_mode': 'fallback',
                'ai_provider': ai_config.get('provider'),
                'ai_model': ai_config.get('model'),
                'ai_unavailable_reason': ai_error_message,
                'differences': fallback_differences,
            }

        return JsonResponse(response_body)

    except CourseOutline.DoesNotExist:
        return JsonResponse({'error': 'Outline not found'}, status=404)
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON body'}, status=400)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)


def _serialize_outline_meta(outline):
    return {
        'id': outline.id,
        'version': outline.version,
        'title': outline.title,
        'status': outline.status,
        'is_current': outline.is_current,
        'faculty': outline.faculty.username if outline.faculty_id else '',
        'created_at': outline.created_at.isoformat() if outline.created_at else None,
        'updated_at': outline.updated_at.isoformat() if outline.updated_at else None,
        'submitted_at': outline.submitted_at.isoformat() if outline.submitted_at else None,
        'approved_at': outline.approved_at.isoformat() if outline.approved_at else None,
    }


def _parse_outline_content(raw_content):
    """Return outline content as a dict if it looks like JSON, otherwise as text."""
    if raw_content in (None, ''):
        return ''
    if isinstance(raw_content, dict) or isinstance(raw_content, list):
        return raw_content
    if isinstance(raw_content, str):
        candidate = raw_content.strip()
        if candidate.startswith('{') or candidate.startswith('['):
            try:
                return json.loads(candidate)
            except json.JSONDecodeError:
                return raw_content
        return raw_content
    return str(raw_content)


def _html_to_comparison_text(raw_html):
    """Convert HTML outline markup into readable plain text for full-document diffs."""
    if not raw_html:
        return ''

    text = str(raw_html)
    # Preserve structure by turning block tags into line breaks.
    text = re.sub(
        r'<\s*br\s*/?\s*>',
        '\n',
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(
        r'</\s*(?:p|div|h[1-6]|li|tr|section|article|table|ul|ol|blockquote|pre)\s*>',
        '\n',
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(
        r'<\s*(?:p|div|h[1-6]|li|tr|section|article|table|ul|ol|blockquote|pre)(?:\s[^>]*)?>',
        '\n',
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(r'<[^>]+>', ' ', text)
    text = html_module.unescape(text)

    lines = []
    for line in text.splitlines():
        normalized_line = re.sub(r'[ \t\u00a0]+', ' ', line).strip()
        if normalized_line:
            lines.append(normalized_line)
    return '\n'.join(lines)


def _outline_content_for_comparison(raw_content):
    """
    Normalize outline content for AI + structural comparison.

    HTML editor content is converted to plain text so the whole document
    (not just the first few thousand HTML characters) is compared.
    """
    parsed = _parse_outline_content(raw_content)
    if isinstance(parsed, (dict, list)):
        return parsed

    text = str(parsed or '')
    if re.search(r'<[a-zA-Z][^>]*>', text):
        return _html_to_comparison_text(text)
    return text.strip()


def _build_fallback_differences(parsed_first_content, parsed_second_content):
    """Always compute structural/textual diffs so we have data even without AI."""
    try:
        if isinstance(parsed_first_content, dict) and isinstance(parsed_second_content, dict):
            return compare_structured_content(parsed_first_content, parsed_second_content)
        first_text = (
            parsed_first_content
            if isinstance(parsed_first_content, str)
            else json.dumps(parsed_first_content, indent=2, ensure_ascii=False)
        )
        second_text = (
            parsed_second_content
            if isinstance(parsed_second_content, str)
            else json.dumps(parsed_second_content, indent=2, ensure_ascii=False)
        )
        return compare_text_content(first_text, second_text)
    except Exception as exc:
        print(f"Fallback diff failed: {exc}")
        return {'added': [], 'modified': [], 'deleted': []}


def _build_outline_excerpt_for_ai(parsed_content, character_limit=120000):
    """Format full outline content for the LLM prompt (plain text preferred)."""
    if isinstance(parsed_content, (dict, list)):
        try:
            text = json.dumps(parsed_content, indent=2, ensure_ascii=False)
        except (TypeError, ValueError):
            text = str(parsed_content)
    else:
        text = str(parsed_content or '')

    # Soft safety cap only — large enough for full course outlines.
    if len(text) > character_limit:
        return (
            text[:character_limit]
            + "\n\n... [content truncated after "
            + str(character_limit)
            + " characters; remaining sections omitted due to size] ..."
        )
    return text


def _extract_json_from_ai_response(raw_text):
    """Extract the first valid JSON object from a LLM response.

    LLMs occasionally wrap JSON in markdown fences or add commentary. This
    helper tolerates both.
    """
    if raw_text is None:
        raise ValueError("Empty AI response")

    text = raw_text.strip()
    if not text:
        raise ValueError("Empty AI response")

    fence_match = re.search(r"```(?:json)?\s*(\{[\s\S]*?\})\s*```", text, re.IGNORECASE)
    if fence_match:
        return json.loads(fence_match.group(1))

    first_brace = text.find('{')
    last_brace = text.rfind('}')
    if first_brace != -1 and last_brace != -1 and last_brace > first_brace:
        candidate = text[first_brace:last_brace + 1]
        return json.loads(candidate)

    return json.loads(text)


def _normalize_ai_comparison_payload(payload):
    """Coerce the AI response into the shape the frontend expects."""
    if not isinstance(payload, dict):
        raise ValueError("AI response is not a JSON object")

    def _coerce_string(value, default=""):
        if value is None:
            return default
        if isinstance(value, (dict, list)):
            return json.dumps(value, ensure_ascii=False)
        return str(value)

    def _coerce_list_of_dicts(items, required_keys):
        normalized_items = []
        if not isinstance(items, list):
            return normalized_items
        for entry in items:
            if not isinstance(entry, dict):
                continue
            normalized_entry = {}
            for required_key in required_keys:
                normalized_entry[required_key] = _coerce_string(entry.get(required_key, ""))
            optional_severity = entry.get('severity') or entry.get('impact')
            if optional_severity:
                normalized_entry['severity'] = _coerce_string(optional_severity)
            normalized_items.append(normalized_entry)
        return normalized_items

    try:
        similarity_value = float(payload.get('similarity_score', 0))
    except (TypeError, ValueError):
        similarity_value = 0.0
    similarity_value = max(0.0, min(100.0, similarity_value))

    overall_change_level = _coerce_string(
        payload.get('overall_change_level') or payload.get('change_level') or 'moderate'
    ).lower()
    if overall_change_level not in {'minimal', 'minor', 'moderate', 'major', 'significant'}:
        overall_change_level = 'moderate'

    key_themes = payload.get('key_themes') or payload.get('themes') or []
    if not isinstance(key_themes, list):
        key_themes = [key_themes]
    key_themes = [_coerce_string(theme) for theme in key_themes if theme]

    recommendations = payload.get('recommendations') or []
    if not isinstance(recommendations, list):
        recommendations = [recommendations]
    recommendations = [_coerce_string(item) for item in recommendations if item]

    return {
        'summary': _coerce_string(payload.get('summary'), default=""),
        'similarity_score': round(similarity_value, 1),
        'overall_change_level': overall_change_level,
        'key_themes': key_themes,
        'added_sections': _coerce_list_of_dicts(
            payload.get('added_sections') or payload.get('added') or [],
            ['title', 'description'],
        ),
        'removed_sections': _coerce_list_of_dicts(
            payload.get('removed_sections') or payload.get('deleted') or payload.get('removed') or [],
            ['title', 'description'],
        ),
        'modified_sections': _coerce_list_of_dicts(
            payload.get('modified_sections') or payload.get('modified') or payload.get('changed') or [],
            ['title', 'description', 'old_summary', 'new_summary'],
        ),
        'recommendations': recommendations,
    }


def _generate_ai_outline_comparison(
    first_outline,
    second_outline,
    parsed_first_content,
    parsed_second_content,
    ai_config,
):
    """Call the LLM to produce a structured comparison of two outlines."""
    first_excerpt = _build_outline_excerpt_for_ai(parsed_first_content)
    second_excerpt = _build_outline_excerpt_for_ai(parsed_second_content)

    prompt = (
        "You are an expert academic curriculum reviewer comparing two versions of a "
        "course outline. Produce a precise, professional analysis.\n\n"
        f"Course: {first_outline.course.code} - {first_outline.course.title}\n\n"
        "IMPORTANT: Both outlines below are the FULL document text (converted from HTML "
        "when needed). Compare the entire document from beginning to end. Do not ignore "
        "later sections such as weekly plan, assessment, textbooks, or policies.\n\n"
        "=== OUTLINE A (older / left side) ===\n"
        f"Version: {first_outline.version}\n"
        f"Title: {first_outline.title}\n"
        f"Status: {first_outline.status}\n"
        f"Author: {first_outline.faculty.username if first_outline.faculty_id else 'Unknown'}\n"
        "Full content:\n"
        f"{first_excerpt}\n\n"
        "=== OUTLINE B (newer / right side) ===\n"
        f"Version: {second_outline.version}\n"
        f"Title: {second_outline.title}\n"
        f"Status: {second_outline.status}\n"
        f"Author: {second_outline.faculty.username if second_outline.faculty_id else 'Unknown'}\n"
        "Full content:\n"
        f"{second_excerpt}\n\n"
        "Compare Outline B against Outline A and respond with ONLY valid JSON, no "
        "Markdown fences, with the following shape:\n"
        "{\n"
        '  "summary": "2-4 sentence executive summary of how Outline B differs from Outline A",\n'
        '  "similarity_score": 0-100 (higher means more similar),\n'
        '  "overall_change_level": "minimal" | "minor" | "moderate" | "major",\n'
        '  "key_themes": ["short bullet phrases of the main themes of change"],\n'
        '  "added_sections": [ { "title": "...", "description": "what was added and why it matters", "severity": "low|medium|high" } ],\n'
        '  "removed_sections": [ { "title": "...", "description": "what was removed and the impact", "severity": "low|medium|high" } ],\n'
        '  "modified_sections": [ { "title": "...", "description": "nature of the change", "old_summary": "what Outline A said", "new_summary": "what Outline B says", "severity": "low|medium|high" } ],\n'
        '  "recommendations": ["actionable improvement recommendations for the new outline"]\n'
        "}\n"
        "Cover differences across the WHOLE outline (course info, CLOs, weekly topics, "
        "assessment, textbooks, policies, etc.). Ignore trivial formatting/whitespace "
        "changes. If a category has no items, return an empty list. Keep descriptions "
        "concise (<= 2 sentences). Output only the JSON object."
    )

    # Longer outlines need more time / output budget.
    comparison_config = dict(ai_config)
    comparison_config['timeout'] = max(float(ai_config.get('timeout') or 30), 90.0)
    comparison_config['max_tokens'] = max(int(ai_config.get('max_tokens') or 4000), 6000)

    raw_response = call_llm_api(prompt, comparison_config)
    parsed_response = _extract_json_from_ai_response(raw_response)
    normalized_response = _normalize_ai_comparison_payload(parsed_response)
    return normalized_response


def compare_structured_content(content1, content2):
    """Compare two structured JSON contents"""
    differences = {
        'added': [],
        'modified': [],
        'deleted': []
    }

    # Compare sections if they exist
    if 'sections' in content1 and 'sections' in content2:
        sections1 = {section.get('id', str(i)): section for i, section in enumerate(content1['sections'])}
        sections2 = {section.get('id', str(i)): section for i, section in enumerate(content2['sections'])}

        # Find added sections
        for section_id, section in sections2.items():
            if section_id not in sections1:
                differences['added'].append({
                    'id': section_id,
                    'title': section.get('title', 'Untitled Section'),
                    'content': section
                })

        # Find deleted sections
        for section_id, section in sections1.items():
            if section_id not in sections2:
                differences['deleted'].append({
                    'id': section_id,
                    'title': section.get('title', 'Untitled Section'),
                    'content': section
                })

        # Find modified sections
        for section_id in set(sections1.keys()) & set(sections2.keys()):
            if sections1[section_id] != sections2[section_id]:
                differences['modified'].append({
                    'id': section_id,
                    'title': sections2[section_id].get('title', 'Untitled Section'),
                    'old_content': sections1[section_id],
                    'new_content': sections2[section_id]
                })
    else:
        # Fall back to full-document text comparison for non-section JSON.
        return compare_text_content(
            json.dumps(content1, indent=2, ensure_ascii=False),
            json.dumps(content2, indent=2, ensure_ascii=False),
        )

    return differences


def compare_text_content(text1, text2):
    """Compare two full-document texts using sequence matching on blocks."""
    blocks1 = [block.strip() for block in re.split(r'\n+', text1 or '') if block.strip()]
    blocks2 = [block.strip() for block in re.split(r'\n+', text2 or '') if block.strip()]

    differences = {
        'added': [],
        'modified': [],
        'deleted': []
    }

    matcher = SequenceMatcher(a=blocks1, b=blocks2, autojunk=False)
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'equal':
            continue

        if tag == 'insert':
            for offset, block in enumerate(blocks2[j1:j2]):
                differences['added'].append({
                    'line': j1 + offset + 1,
                    'title': f'Added block {j1 + offset + 1}',
                    'content': block,
                })
            continue

        if tag == 'delete':
            for offset, block in enumerate(blocks1[i1:i2]):
                differences['deleted'].append({
                    'line': i1 + offset + 1,
                    'title': f'Removed block {i1 + offset + 1}',
                    'content': block,
                })
            continue

        if tag == 'replace':
            old_blocks = blocks1[i1:i2]
            new_blocks = blocks2[j1:j2]
            paired_count = min(len(old_blocks), len(new_blocks))
            for offset in range(paired_count):
                differences['modified'].append({
                    'line': i1 + offset + 1,
                    'title': f'Changed block {i1 + offset + 1}',
                    'old_content': old_blocks[offset],
                    'new_content': new_blocks[offset],
                })
            for offset, block in enumerate(new_blocks[paired_count:]):
                differences['added'].append({
                    'line': j1 + paired_count + offset + 1,
                    'title': f'Added block {j1 + paired_count + offset + 1}',
                    'content': block,
                })
            for offset, block in enumerate(old_blocks[paired_count:]):
                differences['deleted'].append({
                    'line': i1 + paired_count + offset + 1,
                    'title': f'Removed block {i1 + paired_count + offset + 1}',
                    'content': block,
                })

    return differences


@login_required
@user_passes_test(is_admin_or_crc)
@require_http_methods(["GET"])
def api_analysis_course_submissions(request):
    """Return CCR/CRR submissions for a single course along with all answers.

    This powers the per-course detail panel on the Analysis Dashboard so the
    CRC member can drill into every dynamically created form submission for
    the selected course.
    """
    try:
        selected_course_id = request.GET.get('course_id')
        if not selected_course_id:
            return JsonResponse({'error': 'course_id is required'}, status=400)

        try:
            course = Course.objects.select_related('department').get(id=selected_course_id)
        except Course.DoesNotExist:
            return JsonResponse({'error': 'Course not found'}, status=404)

        submissions = DynamicFormSubmission.objects.filter(
            course_id=selected_course_id,
            dynamic_form__form_type__in=['ccr', 'crr'],
        ).select_related(
            'faculty', 'dynamic_form'
        ).prefetch_related('answers__question').order_by('-submission_date')

        ccr_count = 0
        crr_count = 0
        approved_count = 0
        submitted_count = 0
        draft_count = 0
        revision_count = 0

        submissions_payload = []
        for submission in submissions:
            form_type = submission.dynamic_form.form_type
            if form_type == 'ccr':
                ccr_count += 1
            elif form_type == 'crr':
                crr_count += 1

            if submission.status == 'approved':
                approved_count += 1
            elif submission.status == 'submitted':
                submitted_count += 1
            elif submission.status == 'draft':
                draft_count += 1
            elif submission.status == 'revision_requested':
                revision_count += 1

            answers_payload = []
            for answer in submission.answers.all():
                answers_payload.append({
                    'question_id': answer.question.id,
                    'question_text': answer.question.question_text,
                    'question_type': answer.question.question_type,
                    'order': answer.question.order,
                    'answer_text': answer.answer_text,
                    'answer_data': answer.answer_data,
                    'has_file': bool(answer.file_upload),
                    'file_url': answer.file_upload.url if answer.file_upload else None,
                })

            answers_payload.sort(key=lambda item: item.get('order') or 0)

            submissions_payload.append({
                'submission_id': submission.id,
                'form_id': submission.dynamic_form.id,
                'form_name': submission.dynamic_form.name,
                'form_type': form_type,
                'form_type_label': submission.dynamic_form.get_form_type_display(),
                'faculty_username': submission.faculty.username,
                'faculty_email': submission.faculty.email,
                'faculty_department': submission.faculty.department,
                'is_coordinator': submission.is_coordinator,
                'section': submission.section,
                'status': submission.status,
                'status_label': submission.get_status_display(),
                'submission_date': submission.submission_date.isoformat() if submission.submission_date else None,
                'updated_at': submission.updated_at.isoformat() if submission.updated_at else None,
                'answers': answers_payload,
            })

        total_submissions = len(submissions_payload)
        approval_rate = round((approved_count / total_submissions) * 100, 1) if total_submissions else 0

        return JsonResponse({
            'course': {
                'id': course.id,
                'code': course.code,
                'title': course.title,
                'description': course.description,
                'credits': course.credits,
                'department': course.department.name if course.department else '',
            },
            'summary': {
                'total_submissions': total_submissions,
                'ccr_submissions': ccr_count,
                'crr_submissions': crr_count,
                'approved': approved_count,
                'submitted': submitted_count,
                'draft': draft_count,
                'revision_requested': revision_count,
                'approval_rate': approval_rate,
            },
            'submissions': submissions_payload,
        })
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)


@login_required
@user_passes_test(is_admin_or_crc)
@require_http_methods(["GET"])
def api_analysis_clo_by_course(request):
    """Get CLO achievement data broken down by course"""
    try:
        courses = Course.objects.all()
        course_clo_data = []
        
        for course in courses:
            course_data = {
                'course_id': course.id,
                'course_code': course.code,
                'course_title': course.title,
                'department': course.department.name if course.department else '',
                'clo_data': {
                    'CLO1': {'scores': [], 'count': 0, 'achieved': 0},
                    'CLO2': {'scores': [], 'count': 0, 'achieved': 0},
                    'CLO3': {'scores': [], 'count': 0, 'achieved': 0},
                    'CLO4': {'scores': [], 'count': 0, 'achieved': 0},
                }
            }
            
            # Get form submissions for this course
            submissions = DynamicFormSubmission.objects.filter(
                course=course,
                dynamic_form__form_type__in=['ccr', 'crr'],
                status__in=['submitted', 'approved']
            ).prefetch_related('answers__question')
            
            for submission in submissions:
                answers = submission.answers.all()
                
                for answer in answers:
                    # Check for CLO percentage dictionary
                    if isinstance(answer.answer_data, dict):
                        for clo_num in [1, 2, 3, 4]:
                            clo_key = f'CLO{clo_num}'
                            for key in [f'clo{clo_num}', f'clo_{clo_num}', f'CLO{clo_num}']:
                                if key in answer.answer_data:
                                    score = answer.answer_data[key]
                                    if isinstance(score, (int, float)):
                                        course_data['clo_data'][clo_key]['scores'].append(score)
                                        course_data['clo_data'][clo_key]['count'] += 1
                                        if score >= 70:
                                            course_data['clo_data'][clo_key]['achieved'] += 1
                                    elif isinstance(score, str):
                                        try:
                                            if '%' in score:
                                                score_val = float(score.replace('%', '').strip())
                                            else:
                                                score_val = float(score)
                                            
                                            course_data['clo_data'][clo_key]['scores'].append(score_val)
                                            course_data['clo_data'][clo_key]['count'] += 1
                                            if score_val >= 70:
                                                course_data['clo_data'][clo_key]['achieved'] += 1
                                        except:
                                            pass
            
            # Calculate averages and rates for this course
            for clo in ['CLO1', 'CLO2', 'CLO3', 'CLO4']:
                scores = course_data['clo_data'][clo]['scores']
                count = course_data['clo_data'][clo]['count']
                achieved = course_data['clo_data'][clo]['achieved']
                
                if count > 0:
                    avg_score = sum(scores) / len(scores)
                    achievement_rate = (achieved / count) * 100
                else:
                    avg_score = 0
                    achievement_rate = 0
                
                course_data['clo_data'][clo]['average_score'] = round(avg_score, 1)
                course_data['clo_data'][clo]['achievement_rate'] = round(achievement_rate, 1)
            
            course_clo_data.append(course_data)
        
        return JsonResponse(course_clo_data, safe=False)
        
    except Exception as e:
        print(f"Error in CLO by course analysis: {str(e)}")
        return JsonResponse({'error': str(e)}, status=400)



def get_ai_provider_config():
    """Get AI provider configuration from settings"""
    return {
        'provider': settings.AI_CONFIG.get('provider', 'openai'),
        'model': settings.AI_CONFIG.get('model', 'gpt-4'),
        'api_key': settings.AI_CONFIG.get('api_key'),
        'api_base': settings.AI_CONFIG.get('api_base'),
        'available': bool(settings.AI_CONFIG.get('api_key')),
        'timeout': 30.0,
        'max_tokens': 4000,
        'temperature': 0.3
    }

def get_model_string(provider, model_name):
    """Convert provider and model name to LiteLLM compatible string"""
    provider_mappings = {
        'openai': model_name,  # gpt-4, gpt-3.5-turbo, etc.
        'openrouter': f'openrouter/{model_name}',
        'deepseek': f'deepseek/{model_name}',
        'anthropic': f'claude-{model_name}',
        'groq': f'groq/{model_name}',
        'ollama': model_name,
        'together': f'together_ai/{model_name}',
        'huggingface': f'huggingface/{model_name}',
    }
    
    # Return the mapping or default to model_name
    return provider_mappings.get(provider, model_name)

def call_llm_api(prompt, config=None):
    """Generic function to call any LLM provider via LiteLLM"""
    if config is None:
        config = get_ai_provider_config()
    
    if not config['available']:
        raise Exception(f"AI API key not configured for provider: {config['provider']}")
    
    try:
        # Construct model string
        model_string = get_model_string(config['provider'], config['model'])
        
        # Call the API via LiteLLM
        response = litellm.completion(
            model=model_string,
            messages=[
                {"role": "system", "content": "You are an expert CQI analyst for higher education institutions."},
                {"role": "user", "content": prompt}
            ],
            api_key=config['api_key'],
            api_base=config.get('api_base'),
            temperature=config.get('temperature', 0.3),
            max_tokens=config.get('max_tokens', 4000),
            timeout=config.get('timeout', 30.0)
        )
        
        return response.choices[0].message.content
        
    except Exception as e:
        print(f"AI API Error ({config['provider']}): {str(e)}")
        raise Exception(f"AI service error ({config['provider']}): {str(e)}")



#  CQI report function

def markdown_to_cqi_html(markdown_text):
    """Convert CQI Markdown (tables, lists) to HTML for the CRC dashboard."""
    if not markdown_text:
        return ""
    return markdown.markdown(
        markdown_text,
        extensions=[
            "markdown.extensions.tables",
            "markdown.extensions.nl2br",
            "markdown.extensions.sane_lists",
        ],
    )


def build_cqi_chart_data(context_data):
    """Build chart payloads for CQI report visualization and Word export."""
    statistics = context_data.get("statistics") or {}
    clo_analysis = context_data.get("clo_analysis") or {}
    course_review_summary = context_data.get("course_review_summary") or []
    cqi_tables = context_data.get("cqi_tables") or {}

    approved = int(statistics.get("approved_submissions") or 0)
    pending = int(statistics.get("pending_submissions") or 0)
    revision = int(statistics.get("revision_requests") or 0)
    total = int(statistics.get("total_submissions") or 0)
    other = max(total - approved - pending - revision, 0)

    status_labels = []
    status_values = []
    status_colors = []
    status_items = [
        ("Approved", approved, "#16a34a"),
        ("Pending Review", pending, "#ca8a04"),
        ("Revision Requested", revision, "#dc2626"),
        ("Other", other, "#64748b"),
    ]
    for label, value, color in status_items:
        if value > 0:
            status_labels.append(label)
            status_values.append(value)
            status_colors.append(color)

    clo_labels = []
    clo_rates = []
    clo_averages = []
    for clo_name in ("CLO1", "CLO2", "CLO3", "CLO4"):
        payload = (clo_analysis.get("per_clo") or {}).get(clo_name) or {}
        clo_labels.append(clo_name)
        clo_rates.append(payload.get("achievement_rate_percent") or 0)
        clo_averages.append(payload.get("average_score") or 0)

    ccr_total = sum(int(row.get("ccr_submissions") or 0) for row in course_review_summary)
    crr_total = sum(int(row.get("crr_submissions") or 0) for row in course_review_summary)

    table_01_chart = (cqi_tables.get("table_01") or {}).get("chart") or {}
    table_02_chart = (cqi_tables.get("table_02") or {}).get("chart") or {}
    table_03_chart = (cqi_tables.get("table_03") or {}).get("chart") or {}

    # Required corrections = CLOs below 70% achievement (negative/attention signal)
    required_corrections = {
        "labels": clo_labels,
        "data": [max(0.0, round(70.0 - float(rate or 0), 1)) for rate in clo_rates],
        "colors": ["#dc2626", "#ea580c", "#ca8a04", "#b91c1c"],
    }

    return {
        "status_distribution": {
            "labels": status_labels,
            "data": status_values,
            "colors": status_colors,
        },
        "clo_achievement": {
            "labels": clo_labels,
            "achievement_rates": clo_rates,
            "average_scores": clo_averages,
        },
        "required_clo_corrections": required_corrections,
        "form_type_distribution": {
            "labels": ["CCR", "CRR"],
            "data": [ccr_total, crr_total],
            "colors": ["#7c3aed", "#2563eb"],
        },
        "table_01_recommendations": table_01_chart,
        "table_02_hec": table_02_chart,
        "table_03_clo_gaps": table_03_chart,
        "metrics": {
            "total_submissions": total,
            "approved_submissions": approved,
            "pending_submissions": pending,
            "revision_requests": revision,
            "approval_rate": round((approved / total) * 100, 1) if total else 0,
            "average_clo_rate": (
                round(sum(clo_rates) / len(clo_rates), 1) if clo_rates else 0
            ),
            "courses_reviewed": len(course_review_summary),
            "outlines_analyzed": len(
                context_data.get("course_outlines")
                or context_data.get("outlines")
                or []
            ),
            "table_01_rows": len((cqi_tables.get("table_01") or {}).get("rows") or []),
            "table_02_rows": len((cqi_tables.get("table_02") or {}).get("rows") or []),
            "table_03_rows": len((cqi_tables.get("table_03") or {}).get("rows") or []),
        },
    }


# PDF core fonts (Times, Helvetica) only cover Latin-1; normalize or use TTF for Unicode.
_CQI_PDF_UNICODE_REPLACEMENTS = (
    ("\u2014", "--"),  # em dash
    ("\u2013", "-"),  # en dash
    ("\u2212", "-"),  # minus sign
    ("\u2018", "'"),
    ("\u2019", "'"),
    ("\u201c", '"'),
    ("\u201d", '"'),
    ("\u2026", "..."),
    ("\u00a0", " "),
    ("\u200b", ""),
    ("\ufeff", ""),
    ("\u221a", "Y"),  # square root / checkmark-like in reports
    ("\u2713", "Y"),
    ("\u2714", "Y"),
    ("\u2717", "N"),
    ("\u2718", "N"),
)


def _normalize_html_for_core_pdf_fonts(html_text):
    out = html_text
    for before_char, after_char in _CQI_PDF_UNICODE_REPLACEMENTS:
        out = out.replace(before_char, after_char)
    return out


def _register_dejavu_fonts_if_available(pdf_document):
    """
    Register DejaVu TrueType fonts when present on the server so write_html supports
    full Unicode (em dash, check marks, etc.).
    """
    font_pairs = (
        (
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
            "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
        ),
        (
            "/usr/share/fonts/TTF/DejaVuSans.ttf",
            "/usr/share/fonts/TTF/DejaVuSans-Bold.ttf",
        ),
    )
    for regular_font_path, bold_font_path in font_pairs:
        if not os.path.isfile(regular_font_path):
            continue
        bold_path = (
            bold_font_path if os.path.isfile(bold_font_path) else regular_font_path
        )
        try:
            pdf_document.add_font("DejaVuSans", "", regular_font_path)
            pdf_document.add_font("DejaVuSans", "B", bold_path)
            return True
        except Exception:
            continue
    return False


@login_required
@user_passes_test(is_admin_or_crc)
@csrf_exempt
@require_http_methods(["POST"])
def api_generate_cqi_report(request):
    """Generate CQI report using any AI provider via LiteLLM"""
    try:
        data = json.loads(request.body)
        course_id = data.get('course_id')
        time_period = data.get('time_period', 'quarter')
        report_type = data.get('report_type', 'summary')
        
        print(f"CQI Report Request: course_id={course_id}, time_period={time_period}, report_type={report_type}")
        
        # Check if AI is configured
        ai_config = get_ai_provider_config()
        if not ai_config['available']:
            return JsonResponse({
                'error': 'AI service not configured. Please set AI_API_KEY in environment variables.',
                'success': False,
                'fallback_report': "CQI Report Generation Failed: AI service not configured."
            }, status=400)
        
        # Collect comprehensive data for AI analysis
        context_data = collect_data_for_ai(course_id, time_period)
        
        # Add metadata to context
        context_data['metadata'] = {
            'course_id': course_id,
            'time_period': time_period,
            'report_type': report_type,
            'generated_at': datetime.now().isoformat(),
            'ai_provider': ai_config['provider'],
            'ai_model': ai_config['model'],
            'data_summary': {
                'form_submissions_count': len(context_data.get('form_submissions', [])),
                'course_outlines_count': len(context_data.get('course_outlines', [])),
                'has_clo_data': bool(context_data.get('clo_analysis', {}))
            }
        }
        
        # Generate report using AI
        try:
            report = generate_ai_report(context_data, report_type, ai_config)
            tables_html = context_data.get("cqi_tables_html") or ""
            report_html = tables_html + markdown_to_cqi_html(report)

            return JsonResponse({
                'report': report,
                'report_html': report_html,
                'tables_html': tables_html,
                'generated_at': datetime.now().isoformat(),
                'course_id': course_id,
                'time_period': time_period,
                'report_type': report_type,
                'ai_provider': ai_config['provider'],
                'ai_model': ai_config['model'],
                'cover': context_data.get('cover'),
                'chart_data': build_cqi_chart_data(context_data),
                'cqi_tables': context_data.get('cqi_tables'),
                'context_summary': {
                    'form_submissions_analyzed': len(context_data.get('form_submissions', [])),
                    'course_outlines_analyzed': len(context_data.get('course_outlines', [])),
                    'clo_analysis_included': bool(context_data.get('clo_analysis', {})),
                    'statistics': context_data.get('statistics', {}),
                    'evidence_tables': {
                        'table_01_rows': len(
                            ((context_data.get('cqi_tables') or {}).get('table_01') or {}).get('rows') or []
                        ),
                        'table_02_rows': len(
                            ((context_data.get('cqi_tables') or {}).get('table_02') or {}).get('rows') or []
                        ),
                        'table_03_rows': len(
                            ((context_data.get('cqi_tables') or {}).get('table_03') or {}).get('rows') or []
                        ),
                    },
                },
                'success': True
            })
        except Exception as ai_error:
            print(f"AI generation error: {ai_error}")
            # Fallback to manual report
            fallback_report = generate_fallback_report(context_data, report_type)
            tables_html = context_data.get("cqi_tables_html") or ""
            fallback_html = tables_html + markdown_to_cqi_html(fallback_report)
            return JsonResponse({
                'report': fallback_report,
                'report_html': fallback_html,
                'tables_html': tables_html,
                'generated_at': datetime.now().isoformat(),
                'course_id': course_id,
                'time_period': time_period,
                'report_type': report_type,
                'note': f'Generated using fallback method due to AI error: {str(ai_error)}',
                'success': False,
                'error': str(ai_error),
                'cover': context_data.get('cover'),
                'chart_data': build_cqi_chart_data(context_data),
                'cqi_tables': context_data.get('cqi_tables'),
            })
        
    except Exception as e:
        print(f"Error in CQI report generation: {str(e)}")
        import traceback
        traceback.print_exc()
        return JsonResponse({
            'error': str(e),
            'message': 'Failed to generate report. Please try again.',
            'success': False,
            'fallback_report': f"CQI Report Generation Failed\n\nError: {str(e)}"
        }, status=400)


def build_clo_analysis_for_cqi_submissions(filtered_submissions_queryset):
    """Summarize CLO-related answers for CQI context (submissions in the selected window)."""
    clo_data = {
        f"CLO{i}": {"scores": [], "achieved_count": 0, "count": 0}
        for i in range(1, 5)
    }
    submissions_scoped = filtered_submissions_queryset.filter(
        status__in=["submitted", "approved"]
    ).prefetch_related("answers__question")

    for submission in submissions_scoped:
        for answer in submission.answers.all():
            if isinstance(answer.answer_data, dict):
                for clo_num in [1, 2, 3, 4]:
                    clo_key = f"CLO{clo_num}"
                    matched = False
                    for key in (
                        f"clo{clo_num}",
                        f"clo_{clo_num}",
                        f"CLO{clo_num}",
                        f"CLO_{clo_num}",
                    ):
                        if key not in answer.answer_data:
                            continue
                        score = answer.answer_data[key]
                        if isinstance(score, (int, float)):
                            clo_data[clo_key]["scores"].append(float(score))
                            clo_data[clo_key]["count"] += 1
                            if score >= 70:
                                clo_data[clo_key]["achieved_count"] += 1
                            matched = True
                            break
                        if isinstance(score, str):
                            try:
                                score_val = (
                                    float(score.replace("%", "").strip())
                                    if "%" in score
                                    else float(score)
                                )
                                clo_data[clo_key]["scores"].append(score_val)
                                clo_data[clo_key]["count"] += 1
                                if score_val >= 70:
                                    clo_data[clo_key]["achieved_count"] += 1
                                matched = True
                            except (ValueError, TypeError):
                                pass
                            break
                    if matched:
                        continue

            question_text = answer.question.question_text.lower()
            for clo_num in [1, 2, 3, 4]:
                clo_patterns = [
                    f"clo{clo_num}",
                    f"clo {clo_num}",
                    f"course learning outcome {clo_num}",
                    f"learning outcome {clo_num}",
                    f"clo-{clo_num}",
                    f"clo_{clo_num}",
                ]
                if not any(pattern in question_text for pattern in clo_patterns):
                    continue
                clo_key = f"CLO{clo_num}"
                score = extract_clo_score_from_answer(answer)
                if score is None:
                    continue
                clo_data[clo_key]["scores"].append(score)
                clo_data[clo_key]["count"] += 1
                if score >= 70:
                    clo_data[clo_key]["achieved_count"] += 1

    per_clo = {}
    for clo_key, bucket in clo_data.items():
        count = bucket["count"]
        if count > 0:
            average_score = sum(bucket["scores"]) / len(bucket["scores"])
            achievement_rate = (bucket["achieved_count"] / count) * 100
        else:
            average_score = None
            achievement_rate = None
        per_clo[clo_key] = {
            "average_score": round(average_score, 1) if average_score is not None else None,
            "achievement_rate_percent": (
                round(achievement_rate, 1) if achievement_rate is not None else None
            ),
            "response_count": count,
        }
    return {
        "per_clo": per_clo,
        "submissions_with_clo_related_answers": submissions_scoped.count(),
    }


def build_course_review_summary_for_cqi(submissions_queryset):
    """One row per course for Table 1-style CQI summaries."""
    per_course = {}
    queryset = submissions_queryset.select_related("course", "dynamic_form").order_by(
        "-submission_date"
    )
    for submission in queryset:
        course_id_value = submission.course_id
        if course_id_value not in per_course:
            per_course[course_id_value] = {
                "course_id": course_id_value,
                "course_code": submission.course.code,
                "course_title": submission.course.title,
                "coordinator_name": (submission.course_coordinator or "").strip(),
                "ccr_submissions": 0,
                "crr_submissions": 0,
                "has_revision_requested": False,
                "has_pending_submitted": False,
            }
        row = per_course[course_id_value]
        if not row["coordinator_name"] and submission.course_coordinator:
            row["coordinator_name"] = submission.course_coordinator.strip()
        form_type = submission.dynamic_form.form_type
        if form_type == "ccr":
            row["ccr_submissions"] += 1
        elif form_type == "crr":
            row["crr_submissions"] += 1
        if submission.status == "revision_requested":
            row["has_revision_requested"] = True
        if submission.status == "submitted":
            row["has_pending_submitted"] = True

    for course_id_value, row in per_course.items():
        if row["coordinator_name"]:
            continue
        coordinator_assignment = (
            CourseFaculty.objects.filter(course_id=course_id_value, is_coordinator=True)
            .select_related("faculty")
            .first()
        )
        if not coordinator_assignment:
            continue
        full_name = (coordinator_assignment.faculty.get_full_name() or "").strip()
        row["coordinator_name"] = full_name or coordinator_assignment.faculty.username

    return sorted(per_course.values(), key=lambda item: (item["course_code"] or "").lower())


def _answer_text_value(answer):
    if answer is None:
        return ""
    if answer.answer_text and str(answer.answer_text).strip():
        return str(answer.answer_text).strip()
    if isinstance(answer.answer_data, str):
        return answer.answer_data.strip()
    return ""


def _classify_hec_accordance(answer_text):
    """Map CCR Q1 free text to Matched / Not Matched / Not Available."""
    text = (answer_text or "").strip()
    lowered = text.lower()
    if not text or lowered in {"n/a", "na", "none", "-", "nil"}:
        return "not_available"
    if re.match(r"^(no|n)\b", lowered) or any(
        phrase in lowered[:160]
        for phrase in (
            "not covered",
            "not included",
            "does not include",
            "missing",
            "not match",
        )
    ):
        return "not_matched"
    if re.match(r"^(yes|y)\b", lowered) or "covers" in lowered[:120] or "matched" in lowered[:120]:
        return "matched"
    if "partial" in lowered:
        return "not_matched"
    return "matched" if len(text) > 20 else "not_available"


def _recommendation_update_from_content_tools(answer_text):
    """CCR Q2: Yes = recommendation to update content/tools."""
    text = (answer_text or "").strip()
    lowered = text.lower()
    if not text or lowered in {"n/a", "na", "none", "-", "nil", "no"}:
        return "No"
    if re.match(r"^(no|none|n/?a)\b", lowered) and not any(
        word in lowered for word in ("add", "remove", "update", "replace", "include")
    ):
        return "No"
    if any(
        word in lowered
        for word in ("add", "remove", "update", "replace", "reduce", "include", "should", "need")
    ):
        return "Yes"
    return "Yes" if len(text) > 35 else "No"


def _recommendation_update_from_week_or_books(answer_text):
    """CCR Q3/Q4: leading Yes usually means current state is OK (no update)."""
    text = (answer_text or "").strip()
    lowered = text.lower()
    if not text or lowered in {"n/a", "na", "none", "-", "nil"}:
        return "N/A"
    if re.match(r"^(no|n)\b", lowered) or lowered.startswith("partial"):
        return "Yes"
    if any(
        phrase in lowered
        for phrase in (
            "not appropriate",
            "not relevant",
            "outdated",
            "should be replaced",
            "need to",
            "needs to",
            "better balance is required",
        )
    ):
        return "Yes"
    if re.match(r"^(yes|y)\b", lowered):
        return "No"
    return "Yes" if "update" in lowered or "change" in lowered else "No"


def _crr_indicates_clo_not_attained(answer_text):
    """CCR/CRR Q2-style course outcomes: True when CLOs were not fully met."""
    text = (answer_text or "").strip()
    lowered = text.lower()
    if not text:
        return False
    negative_signals = (
        "not met",
        "not achieved",
        "not fully",
        "partially",
        "partial",
        "weak",
        "not satisfied",
        "unsatisfied",
        "only partially",
        "were only partially",
        "few clos",
        "some clos",
    )
    if any(signal in lowered for signal in negative_signals):
        return True
    if re.match(r"^(no|n)\b", lowered):
        return True
    if re.match(r"^(yes|y)\b", lowered) and not any(
        signal in lowered for signal in ("partial", "however", "weak", "except")
    ):
        return False
    return False


def _get_submission_answers_by_order(submission):
    """Map question.order -> FormAnswer for a submission."""
    by_order = {}
    for answer in submission.answers.all():
        by_order[answer.question.order] = answer
    return by_order


def _course_clo_not_attained_flags(ccr_submission):
    """
    From CCR clo_percentage answers (Q6-Q9), mark CLOs with average score < 70
    as not attained for Table 03.
    """
    per_clo_scores = {1: [], 2: [], 3: [], 4: []}
    for answer in ccr_submission.answers.all():
        if answer.question.question_type != "clo_percentage":
            continue
        parsed = _parse_clo_percentage_dict(answer.answer_data)
        for clo_num, score in parsed.items():
            per_clo_scores[clo_num].append(score)

    flags = {}
    for clo_num in [1, 2, 3, 4]:
        scores = per_clo_scores[clo_num]
        if not scores:
            flags[clo_num] = False
            continue
        average_score = sum(scores) / len(scores)
        flags[clo_num] = average_score < 70.0
    return flags


def build_cqi_evidence_tables(filtered_submissions_queryset):
    """
    Build Table 01–03 from real CCR/CRR answers:

    - Table 01: CCR Q2/Q3/Q4 recommendation-to-update flags
    - Table 02: CCR Q1 HEC accordance
    - Table 03: CRR outcomes where CLOs not attained (+ CCR CLO flags)
    """
    submissions = list(
        filtered_submissions_queryset.filter(status__in=["submitted", "approved"])
        .select_related("course", "faculty", "dynamic_form")
        .prefetch_related("answers__question")
        .order_by("course__code", "-submission_date")
    )

    # Latest CCR per course for Tables 01/02 and CLO flags
    latest_ccr_by_course = {}
    for submission in submissions:
        if submission.dynamic_form.form_type != "ccr":
            continue
        course_id_value = submission.course_id
        if course_id_value not in latest_ccr_by_course:
            latest_ccr_by_course[course_id_value] = submission

    table_01_rows = []
    table_02_rows = []
    for index, (course_id_value, submission) in enumerate(
        sorted(
            latest_ccr_by_course.items(),
            key=lambda item: (item[1].course.code or "").lower(),
        ),
        start=1,
    ):
        answers_by_order = _get_submission_answers_by_order(submission)
        q1_text = _answer_text_value(answers_by_order.get(0))
        q2_text = _answer_text_value(answers_by_order.get(1))
        q3_text = _answer_text_value(answers_by_order.get(2))
        q4_text = _answer_text_value(answers_by_order.get(3))

        coordinator_name = (submission.course_coordinator or "").strip()
        if not coordinator_name:
            coordinator_assignment = (
                CourseFaculty.objects.filter(
                    course_id=course_id_value, is_coordinator=True
                )
                .select_related("faculty")
                .first()
            )
            if coordinator_assignment:
                coordinator_name = (
                    coordinator_assignment.faculty.get_full_name() or ""
                ).strip() or coordinator_assignment.faculty.username

        content_tools = _recommendation_update_from_content_tools(q2_text)
        week_wise = _recommendation_update_from_week_or_books(q3_text)
        textbook = _recommendation_update_from_week_or_books(q4_text)

        table_01_rows.append(
            {
                "sr": index,
                "course_code": submission.course.code,
                "course_title": submission.course.title,
                "course_label": f"{submission.course.code} — {submission.course.title}",
                "coordinator_name": coordinator_name or "N/A",
                "content_tools_update": content_tools,
                "week_wise_update": week_wise,
                "textbook_update": textbook,
                "q2_excerpt": (q2_text[:180] + "…") if len(q2_text) > 180 else q2_text,
                "q3_excerpt": (q3_text[:180] + "…") if len(q3_text) > 180 else q3_text,
                "q4_excerpt": (q4_text[:180] + "…") if len(q4_text) > 180 else q4_text,
            }
        )

        hec_status = _classify_hec_accordance(q1_text)
        table_02_rows.append(
            {
                "sr": index,
                "course_code": submission.course.code,
                "course_title": submission.course.title,
                "course_label": f"{submission.course.code} — {submission.course.title}",
                "hec_status": hec_status,
                "not_matched": "✓" if hec_status == "not_matched" else "",
                "matched": "✓" if hec_status == "matched" else "",
                "not_available": "✓" if hec_status == "not_available" else "",
                "q1_excerpt": (q1_text[:220] + "…") if len(q1_text) > 220 else q1_text,
            }
        )

    # Precompute CLO not-attained flags per course from CCR
    clo_flags_by_course = {}
    for course_id_value, submission in latest_ccr_by_course.items():
        clo_flags_by_course[course_id_value] = _course_clo_not_attained_flags(submission)

    table_03_rows = []
    table_03_index = 0
    for submission in submissions:
        if submission.dynamic_form.form_type != "crr":
            continue
        answers_by_order = _get_submission_answers_by_order(submission)
        # CRR Q2 = order 1 (course outcomes / CLO attainment)
        outcomes_text = _answer_text_value(answers_by_order.get(1))
        if not _crr_indicates_clo_not_attained(outcomes_text):
            continue

        clo_flags = clo_flags_by_course.get(submission.course_id) or {
            1: False,
            2: False,
            3: False,
            4: False,
        }
        # If CCR has no CLO scores, still include the row and mark unknown as blank,
        # but keep at least one mark if text clearly indicates gap.
        if not any(clo_flags.values()):
            # Fall back: mark all as needing review with "•"
            clo_display = {1: "•", 2: "•", 3: "•", 4: "•"}
        else:
            clo_display = {
                clo_num: ("✓" if clo_flags[clo_num] else "") for clo_num in [1, 2, 3, 4]
            }

        if not any(clo_display.values()):
            continue

        table_03_index += 1
        table_03_rows.append(
            {
                "sr": table_03_index,
                "course_code": submission.course.code,
                "course_title": submission.course.title,
                "course_label": f"{submission.course.code} — {submission.course.title}",
                "section": submission.section or "N/A",
                "faculty": submission.faculty.username,
                "clo1": clo_display[1],
                "clo2": clo_display[2],
                "clo3": clo_display[3],
                "clo4": clo_display[4],
                "outcomes_excerpt": (
                    (outcomes_text[:200] + "…")
                    if len(outcomes_text) > 200
                    else outcomes_text
                ),
            }
        )

    def _count_yes(rows, key):
        return sum(1 for row in rows if row.get(key) == "Yes")

    table_01_chart = {
        "labels": [
            "Content / Tools & Tech",
            "Week-wise Distribution",
            "Textbook",
        ],
        "data": [
            _count_yes(table_01_rows, "content_tools_update"),
            _count_yes(table_01_rows, "week_wise_update"),
            _count_yes(table_01_rows, "textbook_update"),
        ],
        "colors": ["#2563eb", "#d97706", "#7c3aed"],
    }
    table_02_chart = {
        "labels": ["Not Matched", "Matched", "Not Available"],
        "data": [
            sum(1 for row in table_02_rows if row["hec_status"] == "not_matched"),
            sum(1 for row in table_02_rows if row["hec_status"] == "matched"),
            sum(1 for row in table_02_rows if row["hec_status"] == "not_available"),
        ],
        "colors": ["#dc2626", "#16a34a", "#94a3b8"],
    }
    table_03_chart = {
        "labels": ["CLO1", "CLO2", "CLO3", "CLO4"],
        "data": [
            sum(1 for row in table_03_rows if row.get(f"clo{clo_num}"))
            for clo_num in [1, 2, 3, 4]
        ],
        "colors": ["#ef4444", "#f97316", "#eab308", "#dc2626"],
    }

    return {
        "table_01": {
            "title": "Table 01: Recommendation to Update Course Materials (CCR Q2–Q4)",
            "source": "CCR Form Q2, Q3, Q4",
            "rows": table_01_rows,
            "chart": table_01_chart,
        },
        "table_02": {
            "title": "Table 02: Accordance with HEC Content (CCR Q1)",
            "source": "CCR Form Q1",
            "rows": table_02_rows,
            "chart": table_02_chart,
        },
        "table_03": {
            "title": "Table 03: Courses / Sections where CLOs were Not Attained (CRR Q2)",
            "source": "CRR Form Q2 (shown only when CLOs were not fully attained)",
            "rows": table_03_rows,
            "chart": table_03_chart,
            "included": bool(table_03_rows),
        },
    }


def build_cqi_tables_markdown(cqi_tables):
    """Markdown pipe tables for AI/fallback narrative."""
    table_01 = cqi_tables.get("table_01") or {}
    table_02 = cqi_tables.get("table_02") or {}
    table_03 = cqi_tables.get("table_03") or {}

    lines = []
    lines.append(f"### {table_01.get('title', 'Table 01')}")
    lines.append("")
    lines.append(
        "| Sr | Course Title | Course Coordinator | Content / Tools & Tech | Week-wise Distribution | Textbook |"
    )
    lines.append("| ---: | --- | --- | :---: | :---: | :---: |")
    for row in table_01.get("rows") or []:
        lines.append(
            f"| {row['sr']} | {row['course_label']} | {row['coordinator_name']} | "
            f"{row['content_tools_update']} | {row['week_wise_update']} | {row['textbook_update']} |"
        )
    if not table_01.get("rows"):
        lines.append("| — | _No CCR submissions with Q2–Q4 answers in this window._ | | | | |")
    lines.append("")

    lines.append(f"### {table_02.get('title', 'Table 02')}")
    lines.append("")
    lines.append(
        "| Sr | Course Title | Not Matched with HEC Content | Matched with HEC Content | Not Available |"
    )
    lines.append("| ---: | --- | :---: | :---: | :---: |")
    for row in table_02.get("rows") or []:
        lines.append(
            f"| {row['sr']} | {row['course_label']} | {row['not_matched'] or '—'} | "
            f"{row['matched'] or '—'} | {row['not_available'] or '—'} |"
        )
    if not table_02.get("rows"):
        lines.append("| — | _No CCR Q1 answers in this window._ | | | |")
    lines.append("")

    if table_03.get("included"):
        lines.append(f"### {table_03.get('title', 'Table 03')}")
        lines.append("")
        lines.append(
            "| Sr | Course | Section | CLO1 | CLO2 | CLO3 | CLO4 |"
        )
        lines.append("| ---: | --- | --- | :---: | :---: | :---: | :---: |")
        for row in table_03.get("rows") or []:
            lines.append(
                f"| {row['sr']} | {row['course_label']} | {row['section']} | "
                f"{row['clo1'] or '—'} | {row['clo2'] or '—'} | {row['clo3'] or '—'} | {row['clo4'] or '—'} |"
            )
        lines.append("")
    else:
        lines.append("### Table 03: Courses / Sections where CLOs were Not Attained (CRR Q2)")
        lines.append("")
        lines.append(
            "_No CRR submissions in this window indicated that CLOs were not attained; Table 03 is omitted._"
        )
        lines.append("")

    return "\n".join(lines)


def build_cqi_tables_html(cqi_tables):
    """Professional HTML tables for report preview / PDF."""
    table_01 = cqi_tables.get("table_01") or {}
    table_02 = cqi_tables.get("table_02") or {}
    table_03 = cqi_tables.get("table_03") or {}

    def badge(value):
        text = html_module.escape(str(value or ""))
        if value == "Yes":
            return f'<span style="display:inline-block;padding:2px 8px;border-radius:999px;background:#fef3c7;color:#92400e;font-weight:600;">{text}</span>'
        if value == "No":
            return f'<span style="display:inline-block;padding:2px 8px;border-radius:999px;background:#dcfce7;color:#166534;font-weight:600;">{text}</span>'
        return f'<span style="color:#64748b;">{text or "—"}</span>'

    def check_cell(value):
        if value == "✓":
            return '<span style="color:#dc2626;font-weight:700;">✓</span>'
        if value == "•":
            return '<span style="color:#ca8a04;font-weight:700;">•</span>'
        return '<span style="color:#cbd5e1;">—</span>'

    parts = [
        '<div class="cqi-evidence-tables">',
        '<h2>Evidence Tables (from CCR / CRR form answers)</h2>',
        '<p style="color:#64748b;font-size:0.95rem;">The following tables are generated from real ACQIP form submissions for the selected period.</p>',
    ]

    # Table 01
    parts.append(f"<h3>{html_module.escape(table_01.get('title', 'Table 01'))}</h3>")
    parts.append(
        f'<p style="color:#64748b;font-size:0.85rem;"><em>Source: {html_module.escape(table_01.get("source", ""))}</em></p>'
    )
    parts.append(
        '<table border="1" cellpadding="8" cellspacing="0" width="100%" '
        'style="border-collapse:collapse;width:100%;font-size:13px;margin:12px 0 24px;">'
        "<thead>"
        '<tr style="background:#f8fafc;">'
        '<th rowspan="2" style="text-align:left;">Sr.</th>'
        '<th rowspan="2" style="text-align:left;">Course Title</th>'
        '<th rowspan="2" style="text-align:left;">Course Coordinator</th>'
        '<th colspan="3" style="text-align:center;background:#eef2ff;">Recommendation to Update</th>'
        "</tr>"
        '<tr style="background:#eef2ff;">'
        "<th>Content / Tools &amp; Tech</th>"
        "<th>Week-wise Distribution</th>"
        "<th>Textbook</th>"
        "</tr>"
        "</thead><tbody>"
    )
    for row in table_01.get("rows") or []:
        parts.append(
            "<tr>"
            f"<td>{row['sr']}</td>"
            f"<td><strong>{html_module.escape(row['course_code'])}</strong><br/>"
            f"{html_module.escape(row['course_title'])}</td>"
            f"<td>{html_module.escape(row['coordinator_name'])}</td>"
            f"<td style='text-align:center;'>{badge(row['content_tools_update'])}</td>"
            f"<td style='text-align:center;'>{badge(row['week_wise_update'])}</td>"
            f"<td style='text-align:center;'>{badge(row['textbook_update'])}</td>"
            "</tr>"
        )
    if not table_01.get("rows"):
        parts.append(
            '<tr><td colspan="6" style="text-align:center;color:#64748b;">'
            "No CCR Q2–Q4 answers available for this period.</td></tr>"
        )
    parts.append("</tbody></table>")

    # Table 02
    parts.append(f"<h3>{html_module.escape(table_02.get('title', 'Table 02'))}</h3>")
    parts.append(
        f'<p style="color:#64748b;font-size:0.85rem;"><em>Source: {html_module.escape(table_02.get("source", ""))}</em></p>'
    )
    parts.append(
        '<table border="1" cellpadding="8" cellspacing="0" width="100%" '
        'style="border-collapse:collapse;width:100%;font-size:13px;margin:12px 0 24px;">'
        "<thead>"
        '<tr style="background:#f8fafc;">'
        '<th style="text-align:left;">Sr.</th>'
        '<th style="text-align:left;">Course Title</th>'
        '<th style="text-align:center;">Not Matched with HEC Content</th>'
        '<th style="text-align:center;">Matched with HEC Content</th>'
        '<th style="text-align:center;">Not Available</th>'
        "</tr></thead><tbody>"
    )
    for row in table_02.get("rows") or []:
        parts.append(
            "<tr>"
            f"<td>{row['sr']}</td>"
            f"<td><strong>{html_module.escape(row['course_code'])}</strong><br/>"
            f"{html_module.escape(row['course_title'])}</td>"
            f"<td style='text-align:center;'>{check_cell(row['not_matched'])}</td>"
            f"<td style='text-align:center;'>{check_cell(row['matched'])}</td>"
            f"<td style='text-align:center;'>{check_cell(row['not_available'])}</td>"
            "</tr>"
        )
    if not table_02.get("rows"):
        parts.append(
            '<tr><td colspan="5" style="text-align:center;color:#64748b;">'
            "No CCR Q1 answers available for this period.</td></tr>"
        )
    parts.append("</tbody></table>")

    # Table 03 conditional
    if table_03.get("included"):
        parts.append(f"<h3>{html_module.escape(table_03.get('title', 'Table 03'))}</h3>")
        parts.append(
            f'<p style="color:#64748b;font-size:0.85rem;"><em>Source: {html_module.escape(table_03.get("source", ""))}. '
            "✓ = CLO not attained (CCR quality average &lt; 70%).</em></p>"
        )
        parts.append(
            '<table border="1" cellpadding="8" cellspacing="0" width="100%" '
            'style="border-collapse:collapse;width:100%;font-size:13px;margin:12px 0 24px;">'
            "<thead>"
            '<tr style="background:#f8fafc;">'
            '<th rowspan="2" style="text-align:left;">Sr.</th>'
            '<th rowspan="2" style="text-align:left;">Course</th>'
            '<th rowspan="2" style="text-align:left;">Section</th>'
            '<th colspan="4" style="text-align:center;background:#fef2f2;">CLO Not Attained</th>'
            "</tr>"
            '<tr style="background:#fef2f2;">'
            "<th>CLO1</th><th>CLO2</th><th>CLO3</th><th>CLO4</th>"
            "</tr></thead><tbody>"
        )
        for row in table_03.get("rows") or []:
            parts.append(
                "<tr>"
                f"<td>{row['sr']}</td>"
                f"<td><strong>{html_module.escape(row['course_code'])}</strong><br/>"
                f"{html_module.escape(row['course_title'])}</td>"
                f"<td>{html_module.escape(row['section'])}</td>"
                f"<td style='text-align:center;'>{check_cell(row['clo1'])}</td>"
                f"<td style='text-align:center;'>{check_cell(row['clo2'])}</td>"
                f"<td style='text-align:center;'>{check_cell(row['clo3'])}</td>"
                f"<td style='text-align:center;'>{check_cell(row['clo4'])}</td>"
                "</tr>"
            )
        parts.append("</tbody></table>")
    else:
        parts.append(
            "<h3>Table 03: Courses / Sections where CLOs were Not Attained (CRR Q2)</h3>"
            '<p style="color:#64748b;"><em>No CRR submissions indicated unattained CLOs in this period — table omitted.</em></p>'
        )

    parts.append("</div>")
    return "\n".join(parts)


def cqi_new_courses_detected(course_id, created_since):
    """Courses first created in the system within the window (proxy for newly offered)."""
    course_queryset = Course.objects.select_related("department").filter(
        created_at__gte=created_since
    )
    if course_id:
        course_queryset = course_queryset.filter(id=course_id)
    return [
        {
            "code": course.code,
            "title": course.title,
            "department": course.department.name if course.department else None,
        }
        for course in course_queryset.order_by("code")[:40]
    ]


def collect_data_for_ai(course_id, time_period):
    """Collect structured data for CQI report generation."""
    data = {}
    cover_department = None

    if course_id:
        try:
            course = Course.objects.select_related("department").get(id=course_id)
            cover_department = course.department.name if course.department else None
            data["course"] = {
                "code": course.code,
                "title": course.title,
                "department": cover_department,
                "credits": course.credits,
            }
            outline_rows = CourseOutline.objects.filter(course=course).order_by("-version")[:8]
            data["outlines"] = [
                {
                    "version": outline.version,
                    "status": outline.status,
                    "title": outline.title,
                    "created_at": outline.created_at.isoformat() if outline.created_at else None,
                    "notes": (outline.notes or "")[:500],
                }
                for outline in outline_rows
            ]
        except Course.DoesNotExist:
            data["outlines"] = []
    else:
        data["outlines"] = []

    if course_id:
        submissions_queryset = DynamicFormSubmission.objects.filter(
            course_id=course_id,
            dynamic_form__form_type__in=["ccr", "crr"],
        ).select_related("faculty", "dynamic_form", "course")
    else:
        submissions_queryset = DynamicFormSubmission.objects.filter(
            dynamic_form__form_type__in=["ccr", "crr"]
        ).select_related("faculty", "dynamic_form", "course")

    if time_period == "week":
        period_start = datetime.now() - timedelta(days=7)
    elif time_period == "month":
        period_start = datetime.now() - timedelta(days=30)
    elif time_period == "quarter":
        period_start = datetime.now() - timedelta(days=90)
    elif time_period == "all":
        period_start = None
    else:
        period_start = datetime.now() - timedelta(days=7)

    filtered_submissions = submissions_queryset
    if period_start is not None:
        filtered_submissions = submissions_queryset.filter(
            submission_date__gte=period_start
        )

    data["statistics"] = {
        "total_submissions": filtered_submissions.count(),
        "approved_submissions": filtered_submissions.filter(status="approved").count(),
        "pending_submissions": filtered_submissions.filter(status="submitted").count(),
        "revision_requests": filtered_submissions.filter(
            status="revision_requested"
        ).count(),
    }

    course_review_summary = build_course_review_summary_for_cqi(filtered_submissions)
    if not cover_department and course_review_summary:
        first_course = (
            Course.objects.select_related("department")
            .filter(id=course_review_summary[0]["course_id"])
            .first()
        )
        if first_course and first_course.department:
            cover_department = first_course.department.name
    if data.get("course", {}).get("department"):
        cover_department = data["course"]["department"]

    data["cover"] = {
        "institution": "Capital University of Science and Technology",
        "department": cover_department or "Academic Department",
        "report_year": datetime.now().year,
    }

    new_courses_since = (
        datetime.now() - timedelta(days=730)
        if time_period == "all"
        else (period_start or datetime.now() - timedelta(days=90))
    )
    data["new_courses_detected"] = cqi_new_courses_detected(course_id, new_courses_since)

    course_ids_in_scope = [row["course_id"] for row in course_review_summary]
    if course_id and not course_ids_in_scope:
        course_ids_in_scope = [int(course_id)]

    data["course_outlines"] = []
    if course_ids_in_scope:
        for outline in (
            CourseOutline.objects.filter(course_id__in=course_ids_in_scope)
            .select_related("course")
            .order_by("-updated_at")[:100]
        ):
            data["course_outlines"].append(
                {
                    "course_code": outline.course.code,
                    "course_title": outline.course.title,
                    "version": outline.version,
                    "status": outline.status,
                    "title": outline.title,
                    "notes": (outline.notes or "")[:400],
                    "updated_at": outline.updated_at.isoformat()
                    if outline.updated_at
                    else None,
                }
            )

    if not data.get("outlines") and data["course_outlines"]:
        data["outlines"] = data["course_outlines"][:15]

    data["course_review_summary"] = course_review_summary
    data["clo_analysis"] = build_clo_analysis_for_cqi_submissions(filtered_submissions)
    data["cqi_tables"] = build_cqi_evidence_tables(filtered_submissions)
    data["cqi_tables_markdown"] = build_cqi_tables_markdown(data["cqi_tables"])
    data["cqi_tables_html"] = build_cqi_tables_html(data["cqi_tables"])

    data["submissions"] = []
    submissions_list = list(
        filtered_submissions.order_by("-submission_date")[:60]
    )
    for submission in submissions_list:
        data["submissions"].append(
            {
                "form_type": submission.dynamic_form.form_type,
                "form_name": submission.dynamic_form.name,
                "faculty": submission.faculty.username,
                "status": submission.status,
                "submission_date": submission.submission_date.isoformat()
                if submission.submission_date
                else None,
                "course_code": submission.course.code,
                "course_title": submission.course.title,
                "section": submission.section or "",
                "coordinator_on_form": submission.course_coordinator or "",
            }
        )
    data["form_submissions"] = data["submissions"]
    return data


def generate_ai_report(context_data, report_type="summary", ai_config=None):
    """Generate a CQI report using the institutional CUST-style section layout."""
    if ai_config is None:
        ai_config = get_ai_provider_config()

    depth_notes = {
        "summary": (
            "Keep narrative paragraphs relatively brief (3–5 sentences each), but still include "
            "every required section and every markdown table that has supporting rows in the JSON."
        ),
        "detailed": (
            "Write full analytical narrative under each section; expand on evidence from "
            "`form_submissions`, `course_outlines`, and answers implied by statuses."
        ),
        "recommendations": (
            "Shorten the Introduction and HEC-alignment narratives; keep Annex-A, Table 1, and "
            "end with a clearly prioritized list of CQI / OBE recommendations (numbered)."
        ),
    }.get(
        report_type,
        "Use professional narrative length similar to `detailed` but stay within token limits.",
    )

    prompt = f"""
ROLE: You are a CQI (Continuous Quality Improvement) analyst for Capital University of Science
and Technology, writing in the same structure and tone as the department's annual CQI report
(CRC curriculum review, OBE, CCR/CRR forms, CLO review).

TASK: Analyze ONLY the JSON data below and produce a single CQI report in **Markdown**.

COVER / LETTERHEAD (important):
- Do **not** repeat the university name or department title as a markdown heading; the web UI and
  PDF exporter add the official letterhead from the `cover` object.
- Start your Markdown with exactly: `## Introduction` (first line of your output).

REPORT TYPE: {report_type.upper()}
DEPTH: {depth_notes}

DATA (JSON):
{json.dumps(context_data, indent=2)}

REQUIRED MARKDOWN STRUCTURE (headings must match):

## Introduction
- Opening paragraph: CRC curriculum review and program quality, scoped to this dataset (single
  course vs department-wide using `course`, `statistics`, and `course_review_summary`).
- Numbered list (1–6) of inputs considered. Map to ACQIP as follows (state clearly when data is
  not in the system):
  1. Exit survey (only if referenced in submissions; otherwise say it is outside this dataset)
  2. Alumni survey (same)
  3. Employer survey (same)
  4. Course feedback (CRR and related answers)
  5. CLO attainment and instructor feedback (`clo_analysis`, CCR/CRR)
  6. Industrial Advisory Board or external recommendations (only if present; otherwise note N/A)
- One paragraph on CQI + OBE, CRC review groups, and use of CCR and CRR forms in ACQIP; cite
  numeric counts from `statistics`.

## Annex-A — Course Content Review Process
- Short intro sentence on the CCR process.
- Bullet list with these labels (verbatim), each followed by a short clause tied to the data
  when possible:
  - Compliance of course contents with the HEC recommended contents
  - Instructor's feedback on course contents
  - Week-wise distribution of course contents
  - Relevance and recency of Text and Recommended books
  - Appropriateness of pre-requisite courses
  - Verification of CLOs statements according to SMART criteria
  - Learning domains of CLOs and their level
  - Correctness of CLOs to GAs mapping
- State how many courses appear in `course_review_summary` (and labs if distinguishable).

### Evidence tables (already computed — do not invent rows)
- Professional Tables 01–03 are pre-built from real CCR/CRR answers in `cqi_tables` and
  `cqi_tables_markdown`. The system injects them into the report HTML automatically.
- Discuss findings from:
  - `cqi_tables.table_01` (CCR Q2–Q4 recommendation to update content/tools, week-wise, textbook)
  - `cqi_tables.table_02` (CCR Q1 HEC matched / not matched / not available)
  - `cqi_tables.table_03` (CRR Q2 CLO not attained; omit discussion if `included` is false)
- If `new_courses_detected` is non-empty, add a short markdown table: Sr | Courses.
  Otherwise state that no new courses were flagged.

## Summary of course and lab reviews
- Narrative comparing CCR vs CRR coverage using `statistics` and `course_review_summary`.

## Accordance with HEC curriculum / course materials
- Narrative on HEC alignment using `cqi_tables.table_02` and outline notes; do not claim
  external HEC verification beyond form answers.

## Status of CLOs
- Interpret `clo_analysis`; highlight CLOs below 70% as requiring correction.

## Course Review Report (CRR)
- Introduce the CRR form with bullets: Overall performance of students; Course Outcomes; Coverage
  of course contents; Strategy to support underperforming students; Suggested improvements for
  effective course conduct.
- Use `statistics` and submission list for response counts where possible.
- Reference Table 03 when `cqi_tables.table_03.included` is true.

## Closing and next steps
- Short conclusion and 3–5 concrete next steps for the CRC.

RULES:
- Professional, analytical, constructive tone.
- Do not fabricate survey results or external bodies not present in the JSON.
- Do not invent table rows that contradict `cqi_tables`.
- Use markdown `##` / `###` headings only (no HTML).

"""
    print(f"Generating {report_type} report with {ai_config['provider']} ({ai_config['model']})...")
    return call_llm_api(prompt, ai_config)


def generate_fallback_report(context_data, report_type):
    """Structured Markdown fallback when AI is unavailable."""
    stats = context_data.get("statistics", {})
    cover = context_data.get("cover", {})
    institution = cover.get("institution", "Capital University of Science and Technology")
    department = cover.get("department", "Academic Department")
    year = cover.get("report_year", datetime.now().year)
    total = stats.get("total_submissions") or 0
    approved = stats.get("approved_submissions") or 0
    approval_pct = (approved / total * 100) if total else 0.0
    rows = context_data.get("course_review_summary") or []
    clo = context_data.get("clo_analysis") or {}
    course = context_data.get("course") or {}
    outline_count = len(context_data.get("course_outlines") or context_data.get("outlines") or [])
    # Professional HTML tables are prepended separately in the API response.
    evidence_pointer = (
        "Professional **Tables 01–03** are generated from live CCR/CRR answers and are shown "
        "in the Evidence Tables section at the top of this report "
        "(Table 01: recommendation to update; Table 02: HEC accordance; "
        "Table 03: CLO not attained when applicable)."
    )

    new_courses_list = context_data.get("new_courses_detected", [])
    if new_courses_list:
        table2_lines = ["| Sr | Courses |", "| --- | --- |"]
        for index, item in enumerate(new_courses_list, start=1):
            label = f"{item.get('code', '')} — {item.get('title', '')}"
            table2_lines.append(f"| {index} | {label} |")
        table2 = "\n".join(table2_lines)
    else:
        table2 = "_No new courses were flagged in the system for this window._"

    clo_lines = []
    for clo_name, payload in (clo.get("per_clo") or {}).items():
        avg = payload.get("average_score")
        rate = payload.get("achievement_rate_percent")
        count = payload.get("response_count")
        clo_lines.append(
            f"- **{clo_name}**: responses={count}, avg={avg}, achievement_rate={rate}%"
        )
    clo_section = "\n".join(clo_lines) if clo_lines else "_No CLO quantitative fields were parsed from forms in this window._"

    scope_note = ""
    if course.get("code"):
        scope_note = f"Scope: single course **{course.get('code')}** — {course.get('title', '')}."

    depth = report_type.upper()
    return f"""## Introduction

Automated fallback CQI narrative ({depth}) generated **{datetime.now().strftime("%Y-%m-%d %H:%M")}** because the AI provider was unavailable or returned an error.

{scope_note}

The Curriculum Review Committee (CRC) monitors continuous quality improvement using ACQIP CCR and CRR workflows. In the selected period the system recorded **{total}** CCR/CRR submissions (**{approved}** approved, **{approval_pct:.1f}%** approval rate), **{stats.get("pending_submissions", 0)}** pending CRC review, and **{stats.get("revision_requests", 0)}** revision requests.

1. Exit survey — *Not available unless referenced in form text.*
2. Alumni survey — *Not available unless referenced in form text.*
3. Employer survey — *Not available unless referenced in form text.*
4. Course feedback — *Partially captured through CRR submissions.*
5. CLO attainment and instructor feedback — *Summarized from parsed CLO fields below.*
6. Industrial Advisory Board — *N/A in structured data unless noted in outlines.*

## Annex-A — Course Content Review Process

The Course Contents Review (CCR) form supports structured review of materials across HEC compliance, instructor feedback, week-wise distribution, textbooks, prerequisites, SMART CLO criteria, learning domains/levels, and GA mapping.

**Courses with activity in this window:** {len(rows)}.

{evidence_pointer}

### New courses

{table2}

## Summary of course and lab reviews

CCR and CRR coverage for this window is summarized in the evidence tables and statistics above. **Outline records available:** {outline_count}.

## Status of CLOs

{clo_section}

## Closing and next steps

1. Follow up on courses marked **Yes** for recommendation-to-update in Table 01.
2. Review HEC gaps marked in Table 02 with course coordinators.
3. Address unattained CLOs listed in Table 03 (when present) through targeted remediation.
4. Resolve pending and revision-requested submissions in ACQIP.

---
_Cover metadata: **{institution}**, **{department}**, **{year}**._
"""


@login_required
@user_passes_test(is_admin_or_crc)
@csrf_exempt
@require_http_methods(["POST"])
def api_cqi_report_pdf(request):
    """Render the latest Markdown CQI body into a PDF for download."""
    try:
        payload = json.loads(request.body)
    except json.JSONDecodeError:
        return JsonResponse({"error": "Invalid JSON"}, status=400)

    report_markdown = (
        payload.get("report_markdown") or payload.get("report_body") or ""
    ).strip()
    if not report_markdown:
        return JsonResponse({"error": "report_markdown is required"}, status=400)

    cover = payload.get("cover") or {}
    institution = cover.get("institution", "Capital University of Science and Technology")
    department = cover.get("department", "Academic Department")
    year = cover.get("report_year", datetime.now().year)

    try:
        body_html = markdown_to_cqi_html(report_markdown)
    except Exception as exc:
        return JsonResponse({"error": f"Markdown conversion failed: {exc}"}, status=400)

    tables_html = (payload.get("tables_html") or "").strip()
    if not tables_html and payload.get("cqi_tables"):
        tables_html = build_cqi_tables_html(payload.get("cqi_tables") or {})

    cover_html = (
        '<div align="center">'
        f"<p><b>{html_module.escape(institution)}</b></p>"
        f"<p>{html_module.escape(department)}</p>"
        f"<p><b>CQI Report {html_module.escape(str(year))}</b></p>"
        "</div><hr/>"
    )

    combined_html = cover_html + tables_html + body_html

    pdf_document = FPDF(orientation="P", unit="mm", format="A4")
    pdf_document.set_auto_page_break(auto=True, margin=15)
    pdf_document.set_left_margin(14)
    pdf_document.set_right_margin(14)

    if _register_dejavu_fonts_if_available(pdf_document):
        pdf_document.add_page()
        pdf_document.set_font("DejaVuSans", size=11)
    else:
        combined_html = _normalize_html_for_core_pdf_fonts(combined_html)
        pdf_document.add_page()
        pdf_document.set_font("Helvetica", size=11)

    try:
        pdf_document.write_html(combined_html)
    except Exception as exc:
        return JsonResponse({"error": f"PDF layout failed: {exc}"}, status=500)

    pdf_bytes = bytes(pdf_document.output())
    filename = f"CQI_Report_{year}.pdf"
    response = HttpResponse(pdf_bytes, content_type="application/pdf")
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


def _add_markdown_to_docx(document, markdown_text):
    """Append a lightweight Markdown conversion into a python-docx Document."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if not markdown_text:
        return

    lines = markdown_text.replace("\r\n", "\n").split("\n")
    table_buffer = []

    def flush_table():
        nonlocal table_buffer
        if not table_buffer:
            return
        rows = []
        for table_line in table_buffer:
            cells = [cell.strip() for cell in table_line.strip().strip("|").split("|")]
            if cells and all(set(cell) <= set("-: ") for cell in cells):
                continue
            rows.append(cells)
        table_buffer = []
        if not rows:
            return
        column_count = max(len(row) for row in rows)
        table = document.add_table(rows=len(rows), cols=column_count)
        table.style = "Table Grid"
        for row_index, row_cells in enumerate(rows):
            for column_index in range(column_count):
                cell_text = row_cells[column_index] if column_index < len(row_cells) else ""
                cell = table.rows[row_index].cells[column_index]
                cell.text = cell_text
                if row_index == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
        document.add_paragraph("")

    for raw_line in lines:
        line = raw_line.rstrip()
        stripped = line.strip()

        if stripped.startswith("|") and stripped.endswith("|"):
            table_buffer.append(stripped)
            continue

        flush_table()

        if not stripped:
            continue

        if stripped.startswith("### "):
            heading = document.add_heading(stripped[4:].strip(), level=3)
            continue
        if stripped.startswith("## "):
            heading = document.add_heading(stripped[3:].strip(), level=2)
            continue
        if stripped.startswith("# "):
            heading = document.add_heading(stripped[2:].strip(), level=1)
            continue

        if stripped.startswith("- ") or stripped.startswith("* "):
            paragraph = document.add_paragraph(stripped[2:].strip(), style="List Bullet")
            continue

        numbered_match = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if numbered_match:
            paragraph = document.add_paragraph(
                numbered_match.group(2).strip(), style="List Number"
            )
            continue

        paragraph = document.add_paragraph()
        # Simple bold/italic handling for **text** and *text*
        remaining = stripped
        while remaining:
            bold_match = re.search(r"\*\*(.+?)\*\*", remaining)
            italic_match = re.search(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", remaining)
            next_match = None
            match_type = None
            for candidate_type, candidate in (("bold", bold_match), ("italic", italic_match)):
                if candidate and (
                    next_match is None or candidate.start() < next_match.start()
                ):
                    next_match = candidate
                    match_type = candidate_type
            if not next_match:
                paragraph.add_run(remaining)
                break
            if next_match.start() > 0:
                paragraph.add_run(remaining[: next_match.start()])
            run = paragraph.add_run(next_match.group(1))
            if match_type == "bold":
                run.bold = True
            else:
                run.italic = True
            remaining = remaining[next_match.end() :]

    flush_table()


def _decode_chart_image_bytes(data_url_or_base64):
    """Decode a canvas data URL or raw base64 PNG into bytes."""
    import base64

    if not data_url_or_base64:
        return None
    payload = data_url_or_base64.strip()
    if "," in payload and payload.startswith("data:"):
        payload = payload.split(",", 1)[1]
    try:
        return base64.b64decode(payload)
    except Exception:
        return None


def _add_cqi_evidence_tables_to_docx(document, cqi_tables):
    """Write professional Table 01–03 into a Word document."""
    from docx.shared import Pt

    if not cqi_tables:
        return

    document.add_heading("Evidence Tables (CCR / CRR Form Answers)", level=2)

    table_01 = cqi_tables.get("table_01") or {}
    document.add_heading(table_01.get("title", "Table 01"), level=3)
    source_paragraph = document.add_paragraph(
        f"Source: {table_01.get('source', 'CCR Form')}"
    )
    source_paragraph.runs[0].italic = True
    source_paragraph.runs[0].font.size = Pt(9)

    rows = table_01.get("rows") or []
    word_table = document.add_table(rows=2 + max(len(rows), 1), cols=6)
    word_table.style = "Table Grid"
    headers_top = [
        "Sr.",
        "Course Title",
        "Course Coordinator",
        "Recommendation to Update",
        "",
        "",
    ]
    headers_sub = [
        "",
        "",
        "",
        "Content / Tools & Tech",
        "Week-wise Distribution",
        "Textbook",
    ]
    for index, label in enumerate(headers_top):
        word_table.rows[0].cells[index].text = label
    for index, label in enumerate(headers_sub):
        word_table.rows[1].cells[index].text = label
    # Merge header group for Recommendation to Update
    word_table.rows[0].cells[3].merge(word_table.rows[0].cells[5])
    word_table.rows[0].cells[3].text = "Recommendation to Update"
    if rows:
        for row_offset, row in enumerate(rows):
            cells = word_table.rows[row_offset + 2].cells
            cells[0].text = str(row.get("sr", ""))
            cells[1].text = row.get("course_label") or row.get("course_title") or ""
            cells[2].text = row.get("coordinator_name") or ""
            cells[3].text = row.get("content_tools_update") or ""
            cells[4].text = row.get("week_wise_update") or ""
            cells[5].text = row.get("textbook_update") or ""
    else:
        word_table.rows[2].cells[0].text = "No CCR Q2–Q4 answers in this period."
        word_table.rows[2].cells[0].merge(word_table.rows[2].cells[5])
    document.add_paragraph("")

    table_02 = cqi_tables.get("table_02") or {}
    document.add_heading(table_02.get("title", "Table 02"), level=3)
    source_paragraph = document.add_paragraph(
        f"Source: {table_02.get('source', 'CCR Form Q1')}"
    )
    source_paragraph.runs[0].italic = True
    source_paragraph.runs[0].font.size = Pt(9)
    rows = table_02.get("rows") or []
    word_table = document.add_table(rows=1 + max(len(rows), 1), cols=5)
    word_table.style = "Table Grid"
    for index, label in enumerate(
        [
            "Sr.",
            "Course Title",
            "Not Matched with HEC Content",
            "Matched with HEC Content",
            "Not Available",
        ]
    ):
        word_table.rows[0].cells[index].text = label
    if rows:
        for row_offset, row in enumerate(rows):
            cells = word_table.rows[row_offset + 1].cells
            cells[0].text = str(row.get("sr", ""))
            cells[1].text = row.get("course_label") or row.get("course_title") or ""
            cells[2].text = row.get("not_matched") or "—"
            cells[3].text = row.get("matched") or "—"
            cells[4].text = row.get("not_available") or "—"
    else:
        word_table.rows[1].cells[0].text = "No CCR Q1 answers in this period."
        word_table.rows[1].cells[0].merge(word_table.rows[1].cells[4])
    document.add_paragraph("")

    table_03 = cqi_tables.get("table_03") or {}
    document.add_heading(table_03.get("title", "Table 03"), level=3)
    if table_03.get("included"):
        source_paragraph = document.add_paragraph(
            f"Source: {table_03.get('source', 'CRR Form Q2')}. ✓ = CLO not attained."
        )
        source_paragraph.runs[0].italic = True
        source_paragraph.runs[0].font.size = Pt(9)
        rows = table_03.get("rows") or []
        word_table = document.add_table(rows=2 + len(rows), cols=7)
        word_table.style = "Table Grid"
        for index, label in enumerate(
            ["Sr.", "Course", "Section", "CLO Not Attained", "", "", ""]
        ):
            word_table.rows[0].cells[index].text = label
        word_table.rows[0].cells[3].merge(word_table.rows[0].cells[6])
        word_table.rows[0].cells[3].text = "CLO Not Attained"
        for index, label in enumerate(
            ["", "", "", "CLO1", "CLO2", "CLO3", "CLO4"]
        ):
            word_table.rows[1].cells[index].text = label
        for row_offset, row in enumerate(rows):
            cells = word_table.rows[row_offset + 2].cells
            cells[0].text = str(row.get("sr", ""))
            cells[1].text = row.get("course_label") or ""
            cells[2].text = row.get("section") or "N/A"
            cells[3].text = row.get("clo1") or "—"
            cells[4].text = row.get("clo2") or "—"
            cells[5].text = row.get("clo3") or "—"
            cells[6].text = row.get("clo4") or "—"
    else:
        note = document.add_paragraph(
            "No CRR submissions indicated unattained CLOs in this period — table omitted."
        )
        note.runs[0].italic = True
    document.add_paragraph("")


@login_required
@user_passes_test(is_admin_or_crc)
@csrf_exempt
@require_http_methods(["POST"])
def api_cqi_report_docx(request):
    """Build a professional Word (.docx) CQI report with embedded charts."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from io import BytesIO
    except ImportError:
        return JsonResponse(
            {
                "error": "python-docx is not installed. Run: pip install python-docx",
            },
            status=500,
        )

    try:
        payload = json.loads(request.body)
    except json.JSONDecodeError:
        return JsonResponse({"error": "Invalid JSON"}, status=400)

    report_markdown = (
        payload.get("report_markdown") or payload.get("report_body") or ""
    ).strip()
    if not report_markdown:
        return JsonResponse({"error": "report_markdown is required"}, status=400)

    cover = payload.get("cover") or {}
    chart_images = payload.get("chart_images") or {}
    metrics = payload.get("metrics") or {}
    cqi_tables = payload.get("cqi_tables") or {}
    institution = cover.get(
        "institution", "Capital University of Science and Technology"
    )
    department = cover.get("department", "Academic Department")
    year = cover.get("report_year", datetime.now().year)

    document = Document()

    # Page margins
    for section in document.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # Cover / letterhead
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(institution)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(88, 28, 135)

    department_paragraph = document.add_paragraph()
    department_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    department_run = department_paragraph.add_run(department)
    department_run.font.size = Pt(12)

    report_title = document.add_paragraph()
    report_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    report_run = report_title.add_run(f"Continuous Quality Improvement (CQI) Report — {year}")
    report_run.bold = True
    report_run.font.size = Pt(14)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(
        f"Generated on {datetime.now().strftime('%d %B %Y, %H:%M')}"
    )
    subtitle_run.font.size = Pt(10)
    subtitle_run.font.color.rgb = RGBColor(100, 116, 139)

    document.add_paragraph("")

    # Metrics snapshot
    if metrics:
        document.add_heading("Executive Snapshot", level=2)
        metric_items = [
            ("Total Submissions", metrics.get("total_submissions", 0)),
            ("Approval Rate", f"{metrics.get('approval_rate', 0)}%"),
            ("Pending Review", metrics.get("pending_submissions", 0)),
            ("Avg CLO Rate", f"{metrics.get('average_clo_rate', 0)}%"),
            ("Courses Reviewed", metrics.get("courses_reviewed", 0)),
            ("Outlines Analyzed", metrics.get("outlines_analyzed", 0)),
        ]
        metrics_table = document.add_table(rows=2, cols=3)
        metrics_table.style = "Table Grid"
        for index, (label, value) in enumerate(metric_items):
            row_index = index // 3
            col_index = index % 3
            if row_index > 1:
                break
            cell = metrics_table.rows[row_index].cells[col_index]
            cell.text = f"{label}\n{value}"
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if label in run.text:
                        run.bold = True
        document.add_paragraph("")

    _add_cqi_evidence_tables_to_docx(document, cqi_tables)

    # Charts
    status_image = _decode_chart_image_bytes(chart_images.get("status_distribution"))
    clo_image = _decode_chart_image_bytes(chart_images.get("clo_achievement"))
    form_type_image = _decode_chart_image_bytes(chart_images.get("form_type_distribution"))
    table_01_image = _decode_chart_image_bytes(chart_images.get("table_01_recommendations"))
    table_02_image = _decode_chart_image_bytes(chart_images.get("table_02_hec"))
    table_03_image = _decode_chart_image_bytes(chart_images.get("table_03_clo_gaps"))
    corrections_image = _decode_chart_image_bytes(chart_images.get("required_clo_corrections"))

    if any(
        [
            status_image,
            clo_image,
            form_type_image,
            table_01_image,
            table_02_image,
            table_03_image,
            corrections_image,
        ]
    ):
        document.add_heading("Analytical Figures", level=2)

    chart_figure_specs = [
        (status_image, "Figure 1 — Form Status Distribution"),
        (clo_image, "Figure 2 — CLO Achievement Rates (CCR Q6–Q9)"),
        (corrections_image, "Figure 3 — Required Correction in CLOs (gap to 70%)"),
        (form_type_image, "Figure 4 — CCR vs CRR Submissions"),
        (table_01_image, "Figure 5 — Graph of Table 01 (Recommendation to Update)"),
        (table_02_image, "Figure 6 — Graph of Table 02 (HEC Accordance)"),
        (table_03_image, "Figure 7 — Graph of Table 03 (CLO Not Attained)"),
    ]
    for image_bytes, caption_text in chart_figure_specs:
        if not image_bytes:
            continue
        caption = document.add_paragraph(caption_text)
        caption.runs[0].italic = True
        document.add_picture(BytesIO(image_bytes), width=Inches(5.8))
        document.add_paragraph("")

    document.add_heading("Detailed CQI Narrative", level=2)
    _add_markdown_to_docx(document, report_markdown)

    footer = document.add_paragraph()
    footer_run = footer.add_run(
        "Prepared through ACQIP — Academic Continuous Quality Improvement Platform"
    )
    footer_run.italic = True
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(100, 116, 139)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    filename = f"CQI_Report_{year}.docx"
    response = HttpResponse(
        buffer.getvalue(),
        content_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
    )
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response

